import streamlit as st
import os
import sqlite3
import uuid
import json
import hashlib
import bcrypt
import smtplib
import shutil
import time
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import streamlit.components.v1 as components
from datetime import datetime, timedelta
from google import genai
from google.genai import types
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.colors import HexColor
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.pdfgen import canvas

# ==========================================
# EMAIL CONFIGURATION (SMTP)
# ==========================================
# Configure these values in Streamlit secrets or environment variables
# Go to: Secrets tab -> Add the following keys:
#   SMTP_HOST = "smtp.gmail.com"  (or your email provider)
#   SMTP_PORT = "587"
#   SMTP_USER = "your-email@gmail.com"
#   SMTP_PASSWORD = "your-app-password"
#   SMTP_FROM_NAME = "Liderança Psicanalítica"
#   SMTP_FROM_EMAIL = "your-email@gmail.com"
# ==========================================

def get_secret(key, default=""):
    """Safely get a secret from st.secrets or environment variables."""
    try:
        val = st.secrets.get(key, None)
        if val is not None:
            return val
    except Exception:
        pass
    return os.environ.get(key, default)

def get_smtp_config():
    """Get SMTP configuration from secrets or environment"""
    try:
        return {
            'host': get_secret("SMTP_HOST", ""),
            'port': int(get_secret("SMTP_PORT", "587")),
            'user': get_secret("SMTP_USER", ""),
            'password': get_secret("SMTP_PASSWORD", ""),
            'from_name': get_secret("SMTP_FROM_NAME", "Liderança Psicanalítica"),
            'from_email': get_secret("SMTP_FROM_EMAIL", "")
        }
    except:
        return None

def is_email_configured():
    """Check if email is properly configured"""
    config = get_smtp_config()
    if not config:
        return False
    return bool(config['host'] and config['user'] and config['password'] and config['from_email'])

def send_email(to_email, subject, html_content):
    """Send email using SMTP configuration"""
    if not is_email_configured():
        print(f"[EMAIL] SMTP not configured. Would send to: {to_email}")
        print(f"[EMAIL] Subject: {subject}")
        return False, "SMTP não configurado"
    
    config = get_smtp_config()
    
    try:
        msg = MIMEMultipart('alternative')
        msg['Subject'] = subject
        msg['From'] = f"{config['from_name']} <{config['from_email']}>"
        msg['To'] = to_email
        
        html_part = MIMEText(html_content, 'html', 'utf-8')
        msg.attach(html_part)
        
        with smtplib.SMTP(config['host'], config['port']) as server:
            server.starttls()
            server.login(config['user'], config['password'])
            server.sendmail(config['from_email'], to_email, msg.as_string())
        
        return True, "Email enviado com sucesso"
    except Exception as e:
        print(f"[EMAIL ERROR] {str(e)}")
        return False, str(e)

def send_employee_result_email(employee_name, employee_email, dominant_profile, secondary_profile, bion_role, manager_name):
    """Send assessment result to employee"""
    subject = "Seu Perfil LPS: Insights sobre sua Liderança e Comportamento"
    
    result_text = f"{dominant_profile} + {secondary_profile} ({bion_role})"
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: Arial, sans-serif; background-color: #f5f5f5; margin: 0; padding: 20px; }}
            .container {{ max-width: 600px; margin: 0 auto; background: white; border-radius: 10px; overflow: hidden; }}
            .header {{ background-color: #18738c; color: white; padding: 30px; text-align: center; }}
            .header h1 {{ margin: 0; color: #d19f09; }}
            .content {{ padding: 30px; line-height: 1.7; color: #333; }}
            .result-box {{ background: linear-gradient(135deg, #18738c, #1a5490); color: white; padding: 25px; border-radius: 10px; text-align: center; margin: 25px 0; }}
            .result-box h2 {{ margin: 0; color: #d19f09; font-size: 1.6rem; }}
            .signature {{ margin-top: 30px; padding-top: 20px; border-top: 1px solid #eee; }}
            .footer {{ background-color: #f5f5f5; padding: 20px; text-align: center; font-size: 0.9rem; color: #666; }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>Liderança Psicanalítica</h1>
                <p>Seu Perfil LPS</p>
            </div>
            <div class="content">
                <p>Ola, <strong>{employee_name}</strong>!</p>
                
                <p>Voce acaba de concluir o LPTest, uma etapa fundamental na sua jornada de desenvolvimento dentro da Plataforma LPS.</p>
                
                <p>A analise do seu perfil combina os fundamentos da Psicanálise com as descobertas da Neurociência para mapear nao apenas suas habilidades tecnicas, mas as forcas invisiveis e os mecanismos de defesa que moldam como voce se relaciona com sua equipe e seus lideres.</p>
                
                <div class="result-box">
                    <p style="margin: 0 0 10px 0; font-size: 0.9rem; color: #ccc;">Seu Resultado Principal:</p>
                    <h2>{result_text}</h2>
                </div>
                
                <p>Compreender o funcionamento do seu "eu" profissional e o primeiro passo para uma liderança consciente e um ambiente psicologicamente seguro.</p>
                
                <p>O seu gestor ja recebeu o mapeamento completo e, em breve, voces poderao discutir estrategias de alocacao e desenvolvimento baseadas nesses dados estrategicos.</p>
                
                <div class="signature">
                    <p>Atenciosamente,<br><strong>Viviane Nishiura & Equipe LPS</strong></p>
                </div>
            </div>
            <div class="footer">
                <p>Este e um e-mail automatico da Plataforma LPS.</p>
                <p>Liderança Psicanalítica - Transformando gestores em lideres conscientes.</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    return send_email(employee_email, subject, html_content)

def send_manager_notification_email(manager_email, manager_name, employee_name, employee_profile, bion_role):
    """Notify manager when an employee completes assessment"""
    subject = f"Novo Assessment Concluído - {employee_name}"
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: Arial, sans-serif; background-color: #f5f5f5; margin: 0; padding: 20px; }}
            .container {{ max-width: 600px; margin: 0 auto; background: white; border-radius: 10px; overflow: hidden; }}
            .header {{ background-color: #18738c; color: white; padding: 30px; text-align: center; }}
            .header h1 {{ margin: 0; color: #d19f09; }}
            .content {{ padding: 30px; }}
            .alert-box {{ background-color: #d4edda; border-left: 4px solid #28a745; padding: 15px; border-radius: 5px; margin: 20px 0; }}
            .employee-card {{ background: #f8f9fa; padding: 20px; border-radius: 10px; margin: 20px 0; }}
            .bion-badge {{ display: inline-block; background-color: #d19f09; color: #18738c; padding: 8px 16px; border-radius: 15px; font-weight: bold; }}
            .cta-button {{ display: inline-block; background-color: #18738c; color: white; padding: 15px 30px; border-radius: 25px; text-decoration: none; font-weight: bold; margin-top: 20px; }}
            .footer {{ background-color: #f5f5f5; padding: 20px; text-align: center; font-size: 0.9rem; color: #666; }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>Liderança Psicanalítica</h1>
                <p>Notificação de Assessment</p>
            </div>
            <div class="content">
                <p>Olá <strong>{manager_name}</strong>,</p>
                
                <div class="alert-box">
                    <strong>Novo membro mapeado!</strong> {employee_name} completou o LPTest.
                </div>
                
                <div class="employee-card">
                    <h3 style="margin-top: 0; color: #18738c;">{employee_name}</h3>
                    <p><strong>Perfil:</strong> {employee_profile}</p>
                    <span class="bion-badge">{bion_role}</span>
                </div>
                
                <p>Acesse sua área de gestor para ver o mapeamento completo da equipe e os insights da IA sobre a dinâmica do grupo.</p>
                
                <p style="text-align: center;">
                    <a href="#" class="cta-button">Acessar Dashboard</a>
                </p>
            </div>
            <div class="footer">
                <p>Este é um e-mail automático da Plataforma LPS.</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    return send_email(manager_email, subject, html_content)

def send_welcome_email(user_email, user_name, password):
    """Send welcome email with login credentials (triggered manually after payment confirmation)"""
    subject = "Acesso Liberado - Liderança Psicanalítica"
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: Arial, sans-serif; background-color: #f5f5f5; margin: 0; padding: 20px; }}
            .container {{ max-width: 600px; margin: 0 auto; background: white; border-radius: 10px; overflow: hidden; }}
            .header {{ background-color: #18738c; color: white; padding: 30px; text-align: center; }}
            .header h1 {{ margin: 0; color: #d19f09; }}
            .content {{ padding: 30px; }}
            .credentials-box {{ background: linear-gradient(135deg, #18738c, #1a5490); color: white; padding: 25px; border-radius: 10px; margin: 20px 0; }}
            .credentials-box h3 {{ margin-top: 0; color: #d19f09; }}
            .credential-item {{ background: rgba(255,255,255,0.1); padding: 10px 15px; border-radius: 5px; margin: 10px 0; }}
            .credential-label {{ font-size: 0.9rem; opacity: 0.8; }}
            .credential-value {{ font-size: 1.1rem; font-weight: bold; }}
            .cta-button {{ display: inline-block; background-color: #d19f09; color: #18738c; padding: 15px 30px; border-radius: 25px; text-decoration: none; font-weight: bold; margin-top: 20px; }}
            .features {{ background: #f8f9fa; padding: 20px; border-radius: 10px; margin: 20px 0; }}
            .feature-item {{ padding: 8px 0; border-bottom: 1px solid #e0e0e0; }}
            .feature-item:last-child {{ border-bottom: none; }}
            .footer {{ background-color: #f5f5f5; padding: 20px; text-align: center; font-size: 0.9rem; color: #666; }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>Bem-vindo(a) ao LPS!</h1>
                <p>Seu acesso foi liberado</p>
            </div>
            <div class="content">
                <p>Olá <strong>{user_name}</strong>,</p>
                <p>Parabéns! Seu pagamento foi confirmado e seu acesso à plataforma Liderança Psicanalítica está liberado.</p>
                
                <div class="credentials-box">
                    <h3>Seus Dados de Acesso</h3>
                    <div class="credential-item">
                        <div class="credential-label">E-mail:</div>
                        <div class="credential-value">{user_email}</div>
                    </div>
                    <div class="credential-item">
                        <div class="credential-label">Senha:</div>
                        <div class="credential-value">{password}</div>
                    </div>
                </div>
                
                <p><strong>Importante:</strong> Recomendamos que você altere sua senha após o primeiro acesso.</p>
                
                <div class="features">
                    <h4 style="margin-top: 0; color: #18738c;">O que você terá acesso:</h4>
                    <div class="feature-item">Curso completo com 8 módulos de Liderança Psicanalítica</div>
                    <div class="feature-item">LPTest - Assessment de perfil de liderança</div>
                    <div class="feature-item">Gestão de Equipe - Mapeie até 4 colaboradores</div>
                    <div class="feature-item">LPChat - Consultor de IA especializado</div>
                    <div class="feature-item">Acesso à Mentoria Executiva LPS</div>
                </div>
                
                <p style="text-align: center;">
                    <a href="#" class="cta-button">Acessar a Plataforma</a>
                </p>
            </div>
            <div class="footer">
                <p>Dúvidas? Entre em contato via WhatsApp.</p>
                <p>Liderança Psicanalítica - Transformando gestores em líderes conscientes.</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    return send_email(user_email, subject, html_content)

# Definir diretório de trabalho como local do script
os.chdir(os.path.dirname(os.path.abspath(__file__)))

# Configuração da Página - Tema LPS
st.set_page_config(
    page_title="Plataforma LPS",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

if 'cache_cleared' not in st.session_state:
    st.cache_data.clear()
    st.cache_resource.clear()
    st.session_state.cache_cleared = True

st.markdown("""
<style>
/* ══════════════════════════════════════════════════════════════
   EXPANDER TITLE — force visible text with LPS colours.
   Icons render correctly via Material Symbols font (see fix
   in main CSS block). No icon hiding needed here.
   ══════════════════════════════════════════════════════════════ */
[data-testid="stExpander"] summary p,
[data-testid="stExpander"] summary [data-testid="stMarkdownContainer"] p,
[data-testid="stExpander"] summary div p,
[data-testid="stExpander"] summary span div div p {
    opacity: 1 !important;
    visibility: visible !important;
    color: #18738c !important;
    -webkit-text-fill-color: #18738c !important;
    font-size: 1.1rem !important;
    font-weight: 600 !important;
    display: block !important;
    line-height: 1.4 !important;
}
</style>
""", unsafe_allow_html=True)

# ── JS FIXER: directly style expander titles via parent DOM ───────────────
# CSS specificity battles have been unreliable. This JavaScript runs inside
# a same-origin iframe and directly applies inline styles to expander title
# <p> elements in the parent document, bypassing all CSS cascade issues.
components.html("""
<script>
(function() {
    var TITLE_COLOR = '#18738c';
    var pd;
    try { pd = window.parent.document; } catch(e) { return; }

    function fixExpanders() {
        try {
            pd.querySelectorAll('[data-testid="stExpander"] summary').forEach(function(summary) {
                // Style every <p> inside the expander summary
                summary.querySelectorAll('p').forEach(function(p) {
                    p.style.setProperty('color', TITLE_COLOR, 'important');
                    p.style.setProperty('-webkit-text-fill-color', TITLE_COLOR, 'important');
                    p.style.setProperty('opacity', '1', 'important');
                    p.style.setProperty('visibility', 'visible', 'important');
                    p.style.setProperty('display', 'block', 'important');
                    p.style.setProperty('font-size', '1.05rem', 'important');
                    p.style.setProperty('font-weight', '600', 'important');
                    p.style.setProperty('line-height', '1.4', 'important');
                    p.style.setProperty('max-height', 'none', 'important');
                    p.style.setProperty('overflow', 'visible', 'important');
                });
                // Ensure the summary itself is not clipping content
                summary.style.setProperty('overflow', 'visible', 'important');
                summary.style.setProperty('min-height', '2rem', 'important');
            });
        } catch(e) {}
    }

    // Run immediately
    fixExpanders();
    // Watch for Streamlit re-renders
    try {
        new MutationObserver(fixExpanders).observe(pd.body, { childList: true, subtree: true });
    } catch(e) {}
    // Fallback timer
    setInterval(fixExpanders, 800);
})();
</script>
""", height=0, scrolling=False)

# Password hashing functions using bcrypt
def hash_password(password):
    """Hash password using bcrypt for secure storage."""
    salt = bcrypt.gensalt(rounds=12)
    return bcrypt.hashpw(password.encode('utf-8'), salt).decode('utf-8')

def verify_password(password, hashed):
    """Verify password against hash. Supports bcrypt and legacy SHA-256."""
    # Check if it's a bcrypt hash (starts with $2b$, $2a$, or $2y$)
    if hashed.startswith('$2'):
        try:
            return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))
        except Exception:
            return False
    else:
        # Legacy SHA-256 hash support for existing users
        legacy_hash = hashlib.sha256(password.encode()).hexdigest()
        return legacy_hash == hashed

def upgrade_password_hash(user_id, password):
    """Upgrade legacy SHA-256 hash to bcrypt."""
    new_hash = hash_password(password)
    conn = sqlite3.connect('lps_data.db')
    c = conn.cursor()
    c.execute("UPDATE users SET password_hash = ? WHERE id = ?", (new_hash, user_id))
    conn.commit()
    conn.close()
    return new_hash

# Database Backup System
BACKUP_DIR = "backups"
MAX_BACKUPS = 5  # Keep last 5 backups

def ensure_backup_dir():
    """Ensure backup directory exists."""
    if not os.path.exists(BACKUP_DIR):
        os.makedirs(BACKUP_DIR)

def create_database_backup():
    """Create a backup of the SQLite database."""
    ensure_backup_dir()
    
    db_path = 'lps_data.db'
    if not os.path.exists(db_path):
        return None
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    backup_filename = f"lps_data_backup_{timestamp}.db"
    backup_path = os.path.join(BACKUP_DIR, backup_filename)
    
    try:
        # Use shutil to copy the database file
        shutil.copy2(db_path, backup_path)
        
        # Cleanup old backups (keep only MAX_BACKUPS most recent)
        cleanup_old_backups()
        
        return backup_path
    except Exception as e:
        print(f"Backup error: {e}")
        return None

def cleanup_old_backups():
    """Remove old backups, keeping only the most recent ones."""
    ensure_backup_dir()
    
    backups = []
    for f in os.listdir(BACKUP_DIR):
        if f.startswith('lps_data_backup_') and f.endswith('.db'):
            backup_path = os.path.join(BACKUP_DIR, f)
            backups.append((backup_path, os.path.getmtime(backup_path)))
    
    # Sort by modification time (newest first)
    backups.sort(key=lambda x: x[1], reverse=True)
    
    # Remove old backups beyond MAX_BACKUPS
    for backup_path, _ in backups[MAX_BACKUPS:]:
        try:
            os.remove(backup_path)
        except Exception:
            pass

def restore_from_backup(backup_filename):
    """Restore database from a backup file."""
    backup_path = os.path.join(BACKUP_DIR, backup_filename)
    
    if not os.path.exists(backup_path):
        return False, "Arquivo de backup nao encontrado"
    
    try:
        # Create backup of current db before restoring
        current_backup = create_database_backup()
        
        # Restore
        shutil.copy2(backup_path, 'lps_data.db')
        return True, f"Banco restaurado com sucesso. Backup anterior: {current_backup}"
    except Exception as e:
        return False, f"Erro ao restaurar: {e}"

def get_available_backups():
    """Get list of available backups."""
    ensure_backup_dir()
    
    backups = []
    for f in os.listdir(BACKUP_DIR):
        if f.startswith('lps_data_backup_') and f.endswith('.db'):
            backup_path = os.path.join(BACKUP_DIR, f)
            size_mb = os.path.getsize(backup_path) / (1024 * 1024)
            mod_timestamp = os.path.getmtime(backup_path)
            mod_time = datetime.fromtimestamp(mod_timestamp)
            backups.append({
                'filename': f,
                'size_mb': round(size_mb, 2),
                'created': mod_time.strftime('%d/%m/%Y %H:%M:%S'),
                'timestamp': mod_timestamp  # Keep numeric timestamp for sorting
            })
    
    # Sort by timestamp (newest first) - using numeric timestamp for correct ordering
    backups.sort(key=lambda x: x['timestamp'], reverse=True)
    return backups

def auto_backup_on_startup():
    """Create an automatic backup when the application starts."""
    ensure_backup_dir()
    
    # Check if we already have a backup today
    today = datetime.now().strftime('%Y%m%d')
    for f in os.listdir(BACKUP_DIR):
        if f.startswith(f'lps_data_backup_{today}'):
            return None  # Already have today's backup
    
    # Create new backup
    return create_database_backup()

# Database Setup
def init_db():
    conn = sqlite3.connect('lps_data.db')
    c = conn.cursor()
    
    # Users table for authentication
    c.execute('''CREATE TABLE IF NOT EXISTS users (
        id TEXT PRIMARY KEY,
        email TEXT UNIQUE,
        password_hash TEXT,
        name TEXT,
        user_type TEXT DEFAULT 'manager',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    
    # Managers table
    c.execute('''CREATE TABLE IF NOT EXISTS managers (
        id TEXT PRIMARY KEY,
        user_id TEXT,
        session_id TEXT,
        name TEXT,
        email TEXT,
        profile_dominant TEXT,
        profile_secondary TEXT,
        profile_details TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users(id)
    )''')
    
    # Migration: Add user_id column to managers if it doesn't exist
    c.execute("PRAGMA table_info(managers)")
    columns = [col[1] for col in c.fetchall()]
    if 'user_id' not in columns:
        c.execute("ALTER TABLE managers ADD COLUMN user_id TEXT")
    
    # Employees table
    c.execute('''CREATE TABLE IF NOT EXISTS employees (
        id TEXT PRIMARY KEY,
        manager_id TEXT,
        link_token TEXT UNIQUE,
        slot_number INTEGER,
        name TEXT,
        email TEXT,
        profile_dominant TEXT,
        profile_secondary TEXT,
        profile_details TEXT,
        bion_role TEXT,
        completed INTEGER DEFAULT 0,
        consent_given INTEGER DEFAULT 0,
        consent_date TIMESTAMP,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (manager_id) REFERENCES managers(id)
    )''')
    
    # Migration: Add consent columns if they don't exist
    c.execute("PRAGMA table_info(employees)")
    columns = [col[1] for col in c.fetchall()]
    if 'consent_given' not in columns:
        c.execute("ALTER TABLE employees ADD COLUMN consent_given INTEGER DEFAULT 0")
    if 'consent_date' not in columns:
        c.execute("ALTER TABLE employees ADD COLUMN consent_date TIMESTAMP")
    if 'employee_password' not in columns:
        c.execute("ALTER TABLE employees ADD COLUMN employee_password TEXT")
    
    # Course progress table
    c.execute('''CREATE TABLE IF NOT EXISTS course_progress (
        id TEXT PRIMARY KEY,
        user_id TEXT,
        progress_data TEXT,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users(id)
    )''')
    
    # Payments table for subscription/access control
    c.execute('''CREATE TABLE IF NOT EXISTS payments (
        id TEXT PRIMARY KEY,
        user_id TEXT,
        plan_type TEXT DEFAULT 'basic',
        status TEXT DEFAULT 'pending',
        amount REAL,
        payment_date TIMESTAMP,
        expiry_date TIMESTAMP,
        payment_method TEXT,
        transaction_id TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users(id)
    )''')
    
    # AI Chat logs table
    c.execute('''CREATE TABLE IF NOT EXISTS ai_chat_logs (
        id TEXT PRIMARY KEY,
        user_id TEXT,
        manager_id TEXT,
        message_type TEXT,
        message_content TEXT,
        response_content TEXT,
        tokens_used INTEGER,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users(id),
        FOREIGN KEY (manager_id) REFERENCES managers(id)
    )''')
    
    # Mentorship scheduling table
    c.execute('''CREATE TABLE IF NOT EXISTS mentorship_sessions (
        id TEXT PRIMARY KEY,
        user_id TEXT,
        manager_id TEXT,
        session_date TIMESTAMP,
        session_type TEXT DEFAULT 'individual',
        status TEXT DEFAULT 'scheduled',
        notes TEXT,
        meeting_link TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users(id),
        FOREIGN KEY (manager_id) REFERENCES managers(id)
    )''')
    
    # Assessment responses table - stores each individual response (1-5) for future AI analysis
    c.execute('''CREATE TABLE IF NOT EXISTS assessment_responses (
        id TEXT PRIMARY KEY,
        respondent_id TEXT,
        respondent_type TEXT,
        block_name TEXT,
        question_index INTEGER,
        question_text TEXT,
        response_value INTEGER,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    
    # Authorized users table - pre-registered emails for invite-based access
    c.execute('''CREATE TABLE IF NOT EXISTS authorized_users (
        id TEXT PRIMARY KEY,
        email TEXT UNIQUE,
        name TEXT,
        invite_type TEXT DEFAULT 'equipe',
        status TEXT DEFAULT 'pendente',
        invited_by TEXT,
        completed_at TIMESTAMP,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (invited_by) REFERENCES users(id)
    )''')
    
    # Invite links table - unique URLs for leader/team assessments
    c.execute('''CREATE TABLE IF NOT EXISTS invite_links (
        id TEXT PRIMARY KEY,
        token TEXT UNIQUE,
        invite_type TEXT DEFAULT 'equipe',
        created_by TEXT,
        used_by_email TEXT,
        is_used INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        used_at TIMESTAMP,
        FOREIGN KEY (created_by) REFERENCES users(id)
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS laudos (
        id TEXT PRIMARY KEY,
        respondent_id TEXT NOT NULL,
        respondent_type TEXT NOT NULL DEFAULT 'gestor',
        respondent_name TEXT,
        profile_dominant TEXT,
        profile_secondary TEXT,
        bion_role TEXT,
        laudo_text TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    
    # Migration: Add is_admin column to users if it doesn't exist
    c.execute("PRAGMA table_info(users)")
    user_columns = [col[1] for col in c.fetchall()]
    if 'is_admin' not in user_columns:
        c.execute("ALTER TABLE users ADD COLUMN is_admin INTEGER DEFAULT 0")
    
    # Migration: Add invite_token column to authorized_users if it doesn't exist
    c.execute("PRAGMA table_info(authorized_users)")
    au_columns = [col[1] for col in c.fetchall()]
    if 'invite_token' not in au_columns:
        c.execute("ALTER TABLE authorized_users ADD COLUMN invite_token TEXT")
    
    conn.commit()
    conn.close()

init_db()

def _ensure_admin_user(email, password, name, is_viviane=False):
    """Create or update an admin user. Sets full course progress if is_viviane."""
    conn = sqlite3.connect('lps_data.db')
    c = conn.cursor()
    c.execute("SELECT id FROM users WHERE email = ?", (email,))
    row = c.fetchone()
    pw_hash = hash_password(password)
    if row:
        c.execute("UPDATE users SET password_hash = ?, is_admin = 1, name = ? WHERE email = ?",
                  (pw_hash, name, email))
        user_id = row[0]
    else:
        user_id = str(uuid.uuid4())
        c.execute("INSERT INTO users (id, email, password_hash, name, is_admin) VALUES (?, ?, ?, ?, 1)",
                  (user_id, email, pw_hash, name))
        c.execute("SELECT id FROM managers WHERE email = ?", (email,))
        if not c.fetchone():
            manager_id = str(uuid.uuid4())
            c.execute("INSERT INTO managers (id, user_id, email, name) VALUES (?, ?, ?, ?)",
                      (manager_id, user_id, email, name))
    if is_viviane:
        all_complete = {f"m{mid}_v{vi}": True for mid in range(0, 8) for vi in range(7)}
        c.execute("SELECT id FROM course_progress WHERE user_id = ?", (user_id,))
        if c.fetchone():
            c.execute("UPDATE course_progress SET progress_data = ? WHERE user_id = ?",
                      (json.dumps(all_complete), user_id))
        else:
            c.execute("INSERT INTO course_progress (id, user_id, progress_data) VALUES (?, ?, ?)",
                      (str(uuid.uuid4()), user_id, json.dumps(all_complete)))
    conn.commit()
    conn.close()

# ── OFFICIAL ACCOUNTS — the only two admin users on this platform ─────────
_ensure_admin_user("contato@lps.com.br", "lps2024",    "Viviane Nishiura", is_viviane=True)
_ensure_admin_user("suporte@lps.com.br", "lpsdev2026", "Suporte LPS",      is_viviane=False)

# Run automatic backup on startup (once per day)
auto_backup_on_startup()

# Authentication functions
def register_user(email, password, name):
    conn = get_db()
    c = conn.cursor()
    try:
        user_id = str(uuid.uuid4())
        password_hash = hash_password(password)
        c.execute("INSERT INTO users (id, email, password_hash, name) VALUES (?, ?, ?, ?)",
                  (user_id, email.lower(), password_hash, name))
        manager_id = str(uuid.uuid4())
        c.execute("INSERT INTO managers (id, user_id, email, name) VALUES (?, ?, ?, ?)",
                  (manager_id, user_id, email.lower(), name))
        conn.commit()
        conn.close()
        return user_id, None
    except sqlite3.IntegrityError:
        conn.close()
        return None, "E-mail já cadastrado."

def authenticate_user(email, password):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT id, password_hash, name FROM users WHERE email = ?", (email.lower(),))
    result = c.fetchone()
    conn.close()
    if result and verify_password(password, result[1]):
        # Upgrade legacy SHA-256 hash to bcrypt on successful login
        if not result[1].startswith('$2'):
            upgrade_password_hash(result[0], password)
        return {"id": result[0], "name": result[2], "email": email.lower()}
    return None

def try_employee_login(email, password):
    """Try to log in as an employee using email and password.
    Returns the employee's link_token if credentials match, None otherwise."""
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT link_token, employee_password FROM employees WHERE email = ? AND completed = 1", (email.lower(),))
    result = c.fetchone()
    conn.close()
    if result and result[1]:
        if verify_password(password, result[1]):
            return result[0]
    return None

def get_manager_by_user(user_id):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT id, profile_dominant, profile_secondary, profile_details FROM managers WHERE user_id = ?", (user_id,))
    result = c.fetchone()
    conn.close()
    if result:
        return {
            "id": result[0],
            "dominant": result[1],
            "secondary": result[2],
            "details": json.loads(result[3]) if result[3] else {}
        }
    return None

@st.cache_resource
def get_gemini_client():
    api_key = get_secret("GOOGLE_API_KEY", "")
    if api_key and len(api_key) >= 10:
        return genai.Client(api_key=api_key)
    return None

@st.cache_resource
def get_openai_client():
    api_key = get_secret("OPENAI_API_KEY", "")
    if api_key and len(api_key) >= 10:
        from openai import OpenAI
        return OpenAI(api_key=api_key)
    return None

def get_db():
    return sqlite3.connect('lps_data.db')

def get_manager_by_id(manager_id):
    """Get manager info by manager ID."""
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT id, user_id, name, email, profile_dominant, profile_secondary FROM managers WHERE id = ?", (manager_id,))
    result = c.fetchone()
    conn.close()
    if result:
        return {
            "id": result[0],
            "user_id": result[1],
            "name": result[2],
            "email": result[3],
            "dominant": result[4],
            "secondary": result[5]
        }
    return None

def validate_manager_ownership(user_id, manager_id):
    """Validate that the manager_id belongs to the user_id (multitenancy check)."""
    if not user_id or not manager_id:
        return False
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT id FROM managers WHERE id = ? AND user_id = ?", (manager_id, user_id))
    result = c.fetchone()
    conn.close()
    return result is not None

def get_secure_manager_employees(user_id, manager_id):
    """Get employees for a manager with ownership validation (multitenancy safe)."""
    if not validate_manager_ownership(user_id, manager_id):
        return []
    return get_manager_employees(manager_id)

def generate_employee_link(manager_id, slot_number):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT link_token FROM employees WHERE manager_id = ? AND slot_number = ?", (manager_id, slot_number))
    result = c.fetchone()
    if result:
        conn.close()
        return result[0]
    token = str(uuid.uuid4())[:8]
    employee_id = str(uuid.uuid4())
    c.execute("INSERT INTO employees (id, manager_id, link_token, slot_number) VALUES (?, ?, ?, ?)", 
              (employee_id, manager_id, token, slot_number))
    conn.commit()
    conn.close()
    return token

def save_manager_profile(user_id, dominant, secondary, details):
    conn = get_db()
    c = conn.cursor()
    c.execute("""UPDATE managers SET profile_dominant = ?, profile_secondary = ?, profile_details = ? 
                 WHERE user_id = ?""", (dominant, secondary, json.dumps(details), user_id))
    conn.commit()
    conn.close()

def get_manager_employees(manager_id):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM employees WHERE manager_id = ? ORDER BY slot_number", (manager_id,))
    employees = c.fetchall()
    conn.close()
    return employees

def save_employee_result(token, name, email, dominant, secondary, details, bion_role):
    conn = get_db()
    c = conn.cursor()
    
    # Get manager info for email notification
    c.execute("""SELECT m.name, m.email, m.user_id FROM employees e 
                 JOIN managers m ON e.manager_id = m.id 
                 WHERE e.link_token = ?""", (token,))
    manager_info = c.fetchone()
    manager_name = manager_info[0] if manager_info else "Seu Gestor"
    manager_email = manager_info[1] if manager_info else None
    
    # Update employee record
    c.execute("""UPDATE employees SET name = ?, email = ?, profile_dominant = ?, profile_secondary = ?, 
                 profile_details = ?, bion_role = ?, completed = 1 WHERE link_token = ?""",
              (name, email, dominant, secondary, json.dumps(details), bion_role, token))
    conn.commit()
    conn.close()
    
    # Update authorized_users status to completed if email exists
    if email:
        update_authorized_user_status(email, "concluido")
    
    # Send email to employee with their result
    send_employee_result_email(name, email, dominant, secondary, bion_role, manager_name)
    
    # Send notification to manager
    if manager_email:
        employee_profile = f"{dominant} + {secondary}"
        send_manager_notification_email(manager_email, manager_name, name, employee_profile, bion_role)

def get_employee_by_token(token):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM employees WHERE link_token = ?", (token,))
    employee = c.fetchone()
    conn.close()
    return employee

def save_employee_consent(token):
    """Save employee consent to database with timestamp."""
    conn = get_db()
    c = conn.cursor()
    c.execute("""UPDATE employees SET consent_given = 1, consent_date = ? 
                 WHERE link_token = ?""", (datetime.now(), token))
    conn.commit()
    conn.close()
    return True

def get_employee_consent(token):
    """Check if employee has given consent."""
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT consent_given FROM employees WHERE link_token = ?", (token,))
    result = c.fetchone()
    conn.close()
    if result:
        return result[0] == 1
    return False

def is_user_admin(user_id):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT is_admin FROM users WHERE id = ?", (user_id,))
    result = c.fetchone()
    conn.close()
    return result and result[0] == 1

def set_user_admin(user_id, is_admin=True):
    conn = get_db()
    c = conn.cursor()
    c.execute("UPDATE users SET is_admin = ? WHERE id = ?", (1 if is_admin else 0, user_id))
    conn.commit()
    conn.close()

def create_invite_link(invite_type, created_by):
    conn = get_db()
    c = conn.cursor()
    link_id = str(uuid.uuid4())
    token = str(uuid.uuid4())[:12]
    c.execute("INSERT INTO invite_links (id, token, invite_type, created_by) VALUES (?, ?, ?, ?)",
              (link_id, token, invite_type, created_by))
    conn.commit()
    conn.close()
    return token

def get_invite_links(created_by=None):
    conn = get_db()
    c = conn.cursor()
    if created_by:
        c.execute("SELECT * FROM invite_links WHERE created_by = ? ORDER BY created_at DESC", (created_by,))
    else:
        c.execute("SELECT * FROM invite_links ORDER BY created_at DESC")
    links = c.fetchall()
    conn.close()
    return links

def get_invite_by_token(token):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM invite_links WHERE token = ?", (token,))
    result = c.fetchone()
    conn.close()
    return result

def mark_invite_used(token, email):
    conn = get_db()
    c = conn.cursor()
    c.execute("UPDATE invite_links SET is_used = 1, used_by_email = ?, used_at = ? WHERE token = ?",
              (email, datetime.now(), token))
    conn.commit()
    conn.close()

def save_laudo(respondent_id, respondent_type, respondent_name, profile_dominant, profile_secondary, bion_role, laudo_text):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT id FROM laudos WHERE respondent_id = ? AND respondent_type = ?", (respondent_id, respondent_type))
    existing = c.fetchone()
    if existing:
        c.execute("""UPDATE laudos SET laudo_text = ?, respondent_name = ?, profile_dominant = ?, 
                     profile_secondary = ?, bion_role = ?, updated_at = CURRENT_TIMESTAMP 
                     WHERE respondent_id = ? AND respondent_type = ?""",
                  (laudo_text, respondent_name, profile_dominant, profile_secondary, bion_role, respondent_id, respondent_type))
    else:
        laudo_id = str(uuid.uuid4())
        c.execute("""INSERT INTO laudos (id, respondent_id, respondent_type, respondent_name, profile_dominant, 
                     profile_secondary, bion_role, laudo_text) VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                  (laudo_id, respondent_id, respondent_type, respondent_name, profile_dominant, profile_secondary, bion_role, laudo_text))
    conn.commit()
    conn.close()

def get_laudo(respondent_id, respondent_type):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM laudos WHERE respondent_id = ? AND respondent_type = ?", (respondent_id, respondent_type))
    result = c.fetchone()
    conn.close()
    return result

def get_all_laudos():
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM laudos ORDER BY created_at DESC")
    results = c.fetchall()
    conn.close()
    return results

def add_authorized_user(email, name, invite_type, invited_by):
    conn = get_db()
    c = conn.cursor()
    try:
        auth_id = str(uuid.uuid4())
        invite_token = str(uuid.uuid4())[:12]
        link_id = str(uuid.uuid4())
        c.execute("INSERT INTO authorized_users (id, email, name, invite_type, invited_by, invite_token) VALUES (?, ?, ?, ?, ?, ?)",
                  (auth_id, email.lower().strip(), name, invite_type, invited_by, invite_token))
        c.execute("INSERT INTO invite_links (id, token, invite_type, created_by) VALUES (?, ?, ?, ?)",
                  (link_id, invite_token, invite_type, invited_by))
        conn.commit()
        conn.close()
        return True, None
    except sqlite3.IntegrityError:
        conn.close()
        return False, "E-mail já cadastrado na lista de autorizados."

def get_authorized_users(invited_by=None):
    conn = get_db()
    c = conn.cursor()
    if invited_by:
        c.execute("SELECT * FROM authorized_users WHERE invited_by = ? ORDER BY created_at DESC", (invited_by,))
    else:
        c.execute("SELECT * FROM authorized_users ORDER BY created_at DESC")
    users = c.fetchall()
    conn.close()
    return users

def check_email_authorized(email):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM authorized_users WHERE email = ?", (email.lower().strip(),))
    result = c.fetchone()
    conn.close()
    return result

def update_authorized_user_status(email, status):
    conn = get_db()
    c = conn.cursor()
    if status == "concluido":
        c.execute("UPDATE authorized_users SET status = ?, completed_at = ? WHERE email = ?",
                  (status, datetime.now(), email.lower().strip()))
    else:
        c.execute("UPDATE authorized_users SET status = ? WHERE email = ?",
                  (status, email.lower().strip()))
    conn.commit()
    conn.close()

def delete_authorized_user(auth_id):
    conn = get_db()
    c = conn.cursor()
    c.execute("DELETE FROM authorized_users WHERE id = ?", (auth_id,))
    conn.commit()
    conn.close()

def get_admin_monitoring_data(admin_user_id):
    conn = get_db()
    c = conn.cursor()
    c.execute("""
        SELECT 
            au.name as nome,
            au.email,
            au.invite_type,
            au.status,
            au.completed_at,
            au.created_at,
            au.id as auth_id
        FROM authorized_users au
        WHERE au.invited_by = ?
        ORDER BY au.created_at DESC
    """, (admin_user_id,))
    results = c.fetchall()
    conn.close()
    return results

def get_all_leaders_results():
    conn = get_db()
    c = conn.cursor()
    c.execute("""
        SELECT 
            m.id as manager_id,
            m.user_id,
            u.name,
            u.email,
            m.profile_dominant,
            m.profile_secondary,
            m.profile_details,
            m.created_at
        FROM managers m
        JOIN users u ON m.user_id = u.id
        WHERE m.profile_dominant IS NOT NULL AND m.profile_dominant != ''
        ORDER BY m.created_at DESC
    """)
    results = c.fetchall()
    conn.close()
    return results

def get_all_employees_results():
    conn = get_db()
    c = conn.cursor()
    c.execute("""
        SELECT 
            e.id as employee_id,
            e.name,
            e.email,
            e.profile_dominant,
            e.profile_secondary,
            e.bion_role,
            e.profile_details,
            e.completed,
            e.manager_id,
            m.name as manager_name,
            e.created_at
        FROM employees e
        LEFT JOIN managers m ON e.manager_id = m.id
        WHERE e.completed = 1 AND e.profile_dominant IS NOT NULL AND e.profile_dominant != ''
        ORDER BY e.created_at DESC
    """)
    results = c.fetchall()
    conn.close()
    return results

def get_manager_profile_by_user(user_id):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT profile_dominant, profile_secondary, profile_details FROM managers WHERE user_id = ?", (user_id,))
    result = c.fetchone()
    conn.close()
    if result and result[0]:
        return {
            "dominant": result[0],
            "secondary": result[1],
            "details": json.loads(result[2]) if result[2] else {}
        }
    return None

def get_assessment_block_sums(respondent_id, respondent_type="manager"):
    """Calculate block sums from saved assessment responses for radar chart."""
    conn = get_db()
    c = conn.cursor()
    c.execute("""SELECT block_name, SUM(response_value) as block_sum
                 FROM assessment_responses 
                 WHERE respondent_id = ? AND respondent_type = ?
                 GROUP BY block_name""", (respondent_id, respondent_type))
    results = c.fetchall()
    conn.close()
    
    if not results:
        return None
    
    block_sums = {}
    for row in results:
        block_sums[row[0]] = row[1]
    return block_sums

# Course Progress Functions
def get_course_progress(user_id):
    """Get course progress for a user"""
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT progress_data FROM course_progress WHERE user_id = ?", (user_id,))
    result = c.fetchone()
    conn.close()
    if result and result[0]:
        return json.loads(result[0])
    return {}

def save_course_progress(user_id, progress_data):
    """Save course progress for a user"""
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT id FROM course_progress WHERE user_id = ?", (user_id,))
    existing = c.fetchone()
    if existing:
        c.execute("UPDATE course_progress SET progress_data = ?, updated_at = CURRENT_TIMESTAMP WHERE user_id = ?",
                  (json.dumps(progress_data), user_id))
    else:
        progress_id = str(uuid.uuid4())
        c.execute("INSERT INTO course_progress (id, user_id, progress_data) VALUES (?, ?, ?)",
                  (progress_id, user_id, json.dumps(progress_data)))
    conn.commit()
    conn.close()

def get_module_completion_status(user_id):
    """Get completion status for each module"""
    progress = get_course_progress(user_id)
    module_status = {}
    for mod in MODULES_DATA:
        module_lessons = [f"m{mod['id']}_v{v_idx}" for v_idx in range(len(mod['videos']))]
        completed = sum(1 for lesson in module_lessons if progress.get(lesson, False))
        total = len(module_lessons)
        module_status[mod['id']] = {
            'name': mod['name'],
            'completed': completed,
            'total': total,
            'percentage': (completed / total * 100) if total > 0 else 0
        }
    return module_status

def is_course_completed(user_id):
    """Check if user has completed all theoretical modules (first 5 modules)"""
    module_status = get_module_completion_status(user_id)
    theoretical_modules = [1, 2, 3, 4, 5]  # First 5 modules are theoretical
    for mod_id in theoretical_modules:
        if mod_id in module_status and module_status[mod_id]['percentage'] < 100:
            return False
    return True

def get_assessment_stats(manager_id):
    """Get assessment statistics for a manager"""
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM employees WHERE manager_id = ? AND completed = 1", (manager_id,))
    applied = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM employees WHERE manager_id = ?", (manager_id,))
    total_slots = c.fetchone()[0]
    conn.close()
    max_employees = 4  # Package limit
    remaining = max(0, max_employees - applied)  # Remaining based on completed assessments
    return {
        'applied': applied,
        'remaining': remaining,
        'total_slots': total_slots,
        'max_employees': max_employees
    }

def get_user_payment_status(user_id):
    """Check if user has an active payment/subscription"""
    conn = get_db()
    c = conn.cursor()
    c.execute("""SELECT id, plan_type, status, expiry_date 
                 FROM payments 
                 WHERE user_id = ? AND status = 'active'
                 ORDER BY created_at DESC LIMIT 1""", (user_id,))
    result = c.fetchone()
    conn.close()
    
    if result:
        # Check if payment is still valid (not expired)
        if result[3]:  # expiry_date
            try:
                # Use fromisoformat for robust parsing (handles microseconds)
                expiry = datetime.fromisoformat(result[3]) if isinstance(result[3], str) else result[3]
            except (ValueError, TypeError):
                # Fallback: try common format without microseconds
                try:
                    expiry = datetime.strptime(result[3], '%Y-%m-%d %H:%M:%S')
                except:
                    expiry = datetime.now()  # Safe fallback
            if expiry > datetime.now():
                return {'active': True, 'plan': result[1], 'expiry': result[3]}
        else:
            # No expiry date = lifetime access
            return {'active': True, 'plan': result[1], 'expiry': None}
    
    return {'active': False, 'plan': None, 'expiry': None}

def activate_user_payment(user_id, plan_type='premium', amount=997.0, days_valid=365):
    """Activate a payment for a user (called after WhatsApp confirmation)"""
    conn = get_db()
    c = conn.cursor()
    payment_id = str(uuid.uuid4())
    payment_date = datetime.now()
    expiry_date = payment_date + timedelta(days=days_valid)
    
    c.execute("""INSERT INTO payments (id, user_id, plan_type, status, amount, payment_date, expiry_date, payment_method)
                 VALUES (?, ?, ?, 'active', ?, ?, ?, 'whatsapp')""",
              (payment_id, user_id, plan_type, amount, payment_date, expiry_date))
    conn.commit()
    conn.close()
    return payment_id

def can_access_premium_features(user_id):
    """Check if user can access premium features (LPChat, Mentoria)
    Requires: Active payment AND completed theoretical course modules
    Admin users get full access automatically."""
    if is_user_admin(user_id):
        return {
            'can_access': True,
            'payment_active': True,
            'course_completed': True,
            'plan': 'admin',
            'message': None
        }
    payment_status = get_user_payment_status(user_id)
    course_completed = is_course_completed(user_id)
    
    return {
        'can_access': payment_status['active'] and course_completed,
        'payment_active': payment_status['active'],
        'course_completed': course_completed,
        'plan': payment_status['plan'],
        'message': get_access_message(payment_status['active'], course_completed)
    }

def get_access_message(payment_active, course_completed):
    """Generate appropriate message for access status"""
    if not payment_active:
        return "Acesso liberado após confirmação de pagamento via WhatsApp."
    if not course_completed:
        return "Complete os módulos teóricos do curso para liberar este recurso."
    return None

def render_premium_gate(feature_name="esta funcionalidade"):
    st.markdown("""
        <style>
        .premium-gate-card {
            max-width: 600px;
            margin: 40px auto;
            background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
            border: 2px solid #d19f09;
            border-radius: 16px;
            padding: 48px 40px;
            text-align: center;
            box-shadow: 0 8px 32px rgba(0,0,0,0.08);
            font-family: 'Open Sans', sans-serif;
        }
        .premium-gate-icon {
            font-size: 48px;
            margin-bottom: 16px;
        }
        .premium-gate-title {
            font-family: 'Open Sans', sans-serif;
            color: #18738c;
            font-size: 1.6rem;
            font-weight: 700;
            margin-bottom: 12px;
        }
        .premium-gate-text {
            font-family: 'Ubuntu', sans-serif;
            color: #444;
            font-size: 1.05rem;
            line-height: 1.7;
            margin-bottom: 24px;
        }
        .premium-gate-divider {
            width: 60px;
            height: 3px;
            background: #d19f09;
            margin: 20px auto;
            border-radius: 2px;
        }
        .premium-gate-cta {
            display: inline-block;
            background: linear-gradient(135deg, #d19f09 0%, #c5a059 100%);
            color: white !important;
            font-family: 'Ubuntu', sans-serif;
            font-weight: 700;
            font-size: 1.1rem;
            padding: 14px 36px;
            border-radius: 8px;
            text-decoration: none;
            transition: all 0.3s ease;
            box-shadow: 0 4px 12px rgba(209,159,9,0.3);
        }
        .premium-gate-cta:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(209,159,9,0.4);
            color: white !important;
        }
        .premium-gate-footer {
            color: #888;
            font-size: 0.85rem;
            margin-top: 20px;
        }
        </style>
    """, unsafe_allow_html=True)
    
    st.markdown(f"""
        <div class="premium-gate-card">
            <div class="premium-gate-icon">🔒</div>
            <div class="premium-gate-title">Desbloqueie seu Potencial de Liderança</div>
            <div class="premium-gate-divider"></div>
            <div class="premium-gate-text">
                O acesso completo a esta ferramenta e aos conteúdos exclusivos da Metodologia LPS
                está disponível para membros oficiais.
            </div>
            <a href="https://wa.me/5511999999999?text=Olá! Gostaria de liberar meu acesso à Plataforma LPS." 
               target="_blank" class="premium-gate-cta">
                Liberar Acesso via WhatsApp
            </a>
            <div class="premium-gate-footer">
                Após a confirmação do pagamento, seu acesso será ativado em até 24 horas.
            </div>
        </div>
    """, unsafe_allow_html=True)
    return True

def is_user_premium(user_id):
    if is_user_admin(user_id):
        return True
    payment_status = get_user_payment_status(user_id)
    return payment_status['active']

def log_ai_chat(user_id, manager_id, message_type, message_content, response_content, tokens_used=0):
    """Log AI chat interactions"""
    conn = get_db()
    c = conn.cursor()
    log_id = str(uuid.uuid4())
    c.execute("""INSERT INTO ai_chat_logs (id, user_id, manager_id, message_type, message_content, response_content, tokens_used)
                 VALUES (?, ?, ?, ?, ?, ?, ?)""",
              (log_id, user_id, manager_id, message_type, message_content, response_content, tokens_used))
    conn.commit()
    conn.close()
    return log_id

def get_ai_insights(manager_id, user_id):
    """Generate AI insights about the team with multitenancy validation"""
    # Use secure function to validate ownership
    employees = get_secure_manager_employees(user_id, manager_id)
    manager_profile = get_manager_profile_by_user(user_id)
    
    insights = []
    
    # Check if manager has taken the assessment
    if not manager_profile or not manager_profile.get('dominant'):
        insights.append({
            'type': 'warning',
            'message': 'Complete seu LPTest para receber insights personalizados sobre sua liderança.'
        })
    
    # Check team size
    completed_employees = [e for e in employees if e[10] == 1]  # completed column
    if len(completed_employees) == 0:
        insights.append({
            'type': 'info',
            'message': 'Envie os links de assessment para sua equipe para começar a receber insights.'
        })
    elif len(completed_employees) < 3:
        insights.append({
            'type': 'info',
            'message': f'Você tem {len(completed_employees)} funcionário(s) mapeado(s). Mapeie mais membros para análises mais completas.'
        })
    else:
        # Analyze Bion roles distribution
        bion_roles = [e[9] for e in completed_employees if e[9]]  # bion_role column
        if bion_roles:
            role_counts = {}
            for role in bion_roles:
                role_counts[role] = role_counts.get(role, 0) + 1
            
            most_common = max(role_counts, key=role_counts.get)
            if role_counts[most_common] > len(bion_roles) / 2:
                insights.append({
                    'type': 'alert',
                    'message': f'Concentração de papéis: {role_counts[most_common]} membros com papel "{most_common}". Considere diversificar.'
                })
            else:
                insights.append({
                    'type': 'success',
                    'message': 'Boa diversidade de papéis na equipe! Isso favorece a dinâmica grupal.'
                })
    
    return insights[:3]  # Return max 3 insights

def get_app_url():
    """Get the public-facing base URL of this app (used for invite links).
    Priority: deployed production URL > known production domain.
    Never uses the dev-workspace domain (REPLIT_DEV_DOMAIN) because
    that URL requires Replit login and employees cannot access it.
    """
    # 1. Deployment URL (set when the app is published/deployed)
    deploy_url = os.environ.get('REPLIT_DEPLOYMENT_URL', '')
    if deploy_url:
        return deploy_url.rstrip('/') if deploy_url.startswith('https://') else f"https://{deploy_url.rstrip('/')}"
    # 2. REPLIT_DOMAINS env var (available on deployed apps, comma-separated)
    domains_env = os.environ.get('REPLIT_DOMAINS', '')
    if domains_env:
        first_domain = domains_env.split(',')[0].strip()
        if first_domain:
            return f"https://{first_domain}"
    # 3. Hardcoded production URL — safe fallback for sharing
    return "https://lps-plataforma.replit.app"

# Estilização Customizada
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Open+Sans:wght@400;600;700;800&family=Ubuntu:wght@300;400;500;700&display=swap');
    
    :root {
        --primary-blue: #18738c;
        --accent-gold: #d19f09;
        --bg-gray: #F5F5F5;
        --light-gold: #FFF9E6;
    }
    
    /* ============================================================
       WHITE-LABEL: hide only the 3-dots menu and Streamlit footer
       ============================================================ */
    [data-testid="stMainMenu"] { display: none !important; }
    footer { display: none !important; }
    /* ============================================================
       HAMBURGER / SIDEBAR TOGGLE — always gold, always visible
       ============================================================ */
    /* The sidebar expand button shown when sidebar is collapsed */
    [data-testid="collapsedControl"],
    [data-testid="stExpandSidebarButton"],
    button[data-testid="collapsedControl"] {
        background-color: #d19f09 !important;
        border-radius: 8px !important;
        border: 2px solid #d19f09 !important;
        box-shadow: 0 2px 8px rgba(0,0,0,0.3) !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        min-width: 36px !important;
        min-height: 36px !important;
        opacity: 1 !important;
        visibility: visible !important;
    }
    [data-testid="collapsedControl"] svg,
    [data-testid="stExpandSidebarButton"] svg {
        fill: #FFFFFF !important;
        stroke: #FFFFFF !important;
        color: #FFFFFF !important;
        width: 20px !important;
        height: 20px !important;
    }
    [data-testid="collapsedControl"]:hover,
    [data-testid="stExpandSidebarButton"]:hover {
        background-color: #c5a059 !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.4) !important;
    }

    .main { background-color: var(--bg-gray); }
    h1, h2, h3 { 
        color: var(--primary-blue) !important;
        font-family: 'Open Sans', sans-serif !important;
    }
    body, p, div, label {
        font-family: 'Ubuntu', sans-serif !important;
    }
    /* Do NOT override font on icon spans — they need Material Symbols font */
    span:not([data-testid="stIconMaterial"]):not([translate="no"]) {
        font-family: 'Ubuntu', sans-serif !important;
    }
    /* Restore Material Symbols font for Streamlit icon spans */
    [data-testid="stIconMaterial"],
    span[translate="no"] {
        font-family: 'Material Symbols Rounded', 'Material Icons', sans-serif !important;
    }
    
    .stButton>button {
        background-color: var(--primary-blue);
        color: white;
        border-radius: 8px;
        border: 2px solid var(--accent-gold);
        width: 100%;
        font-weight: bold;
    }
    
    .stDownloadButton>button {
        background-color: var(--accent-gold) !important;
        color: black !important;
        font-weight: bold !important;
        border-radius: 8px !important;
        border: 1px solid var(--primary-blue) !important;
    }
    
    div[data-testid="stExpander"] {
        border: 1px solid var(--primary-blue);
        border-radius: 8px;
        background-color: var(--light-gold);
    }

    div[data-testid="stExpander"] summary p {
        color: var(--primary-blue) !important;
        font-weight: bold !important;
    }

    [data-testid="stSidebar"] {
        background-color: white;
        border-right: 1px solid #e0e0e0;
    }
    
    .result-card {
        background-color: var(--light-gold);
        padding: 2.5rem;
        border-radius: 15px;
        border: 5px solid var(--accent-gold);
        margin-top: 1.5rem;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    
    .profile-title {
        color: var(--primary-blue);
        font-size: 2.2rem;
        font-weight: 800;
        text-align: center;
        margin-bottom: 1rem;
    }

    .question-text {
        color: var(--primary-blue);
        font-weight: 600;
        margin-top: 15px;
        margin-bottom: 5px;
        font-size: 1.1rem;
    }

    .section-header {
        color: var(--primary-blue);
        border-bottom: 2px solid var(--accent-gold);
        padding-bottom: 5px;
        margin-top: 20px;
        font-weight: bold;
    }

    .employee-card {
        background-color: white;
        padding: 1.5rem;
        border-radius: 10px;
        border: 2px solid var(--accent-gold);
        margin: 10px 0;
    }

    .link-box {
        background-color: var(--light-gold);
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid var(--primary-blue);
        font-family: monospace;
        word-break: break-all;
    }

    /* Slider styling - blue bar, gold thumb */
    div[data-testid="stSlider"] > div > div > div {
        background-color: var(--primary-blue) !important;
    }
    div[data-testid="stSlider"] [role="slider"] {
        background-color: var(--accent-gold) !important;
        border: 2px solid var(--primary-blue) !important;
        box-shadow: 0 2px 6px rgba(0,0,0,0.2) !important;
    }
    div[data-testid="stSlider"] [data-testid="stTickBarMin"],
    div[data-testid="stSlider"] [data-testid="stTickBarMax"] {
        background-color: var(--primary-blue) !important;
    }

    /* Select slider styling */
    div[data-baseweb="slider"] > div > div:first-child {
        background-color: var(--primary-blue) !important;
    }
    div[data-baseweb="slider"] [role="slider"] {
        background-color: var(--accent-gold) !important;
        border-color: var(--primary-blue) !important;
    }
    div[data-baseweb="slider"] > div > div:nth-child(3) {
        background-color: #c0d8e0 !important;
    }

    .bion-badge {
        background-color: var(--primary-blue);
        color: white;
        padding: 5px 12px;
        border-radius: 15px;
        font-size: 0.9rem;
        display: inline-block;
        margin-top: 5px;
    }
    
    .module-card {
        background: white;
        border: 2px solid var(--accent-gold);
        border-radius: 15px;
        padding: 1.5rem;
        text-align: center;
        transition: transform 0.2s, box-shadow 0.2s;
        cursor: pointer;
        height: 100%;
    }
    
    .module-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 8px 25px rgba(13, 59, 102, 0.2);
    }
    
    .module-icon {
        font-size: 3rem;
        margin-bottom: 1rem;
    }
    
    .module-title {
        color: var(--primary-blue);
        font-size: 1.1rem;
        font-weight: bold;
        margin-bottom: 0.5rem;
    }
    
    .module-desc {
        color: #666;
        font-size: 0.9rem;
        line-height: 1.4;
    }
    
    
    .header-bar {
        background: linear-gradient(135deg, var(--primary-blue) 0%, #1a5490 100%);
        padding: 1rem 2rem;
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin: -1rem -1rem 1rem -1rem;
        border-radius: 0 0 15px 15px;
    }
    
    .header-logo {
        display: flex;
        align-items: center;
        gap: 15px;
    }
    
    .header-title {
        color: var(--accent-gold);
        font-size: 1.5rem;
        font-weight: bold;
        margin: 0;
    }
    
    .btn-area-aluno {
        background-color: var(--accent-gold);
        color: var(--primary-blue);
        border: none;
        padding: 10px 25px;
        border-radius: 25px;
        font-weight: bold;
        cursor: pointer;
        transition: transform 0.2s;
    }
    
    .btn-area-aluno:hover {
        transform: scale(1.05);
    }
    
    .hero-section {
        background: linear-gradient(135deg, var(--primary-blue) 0%, #1a5490 100%);
        padding: 3rem 2rem;
        border-radius: 20px;
        text-align: center;
        margin-bottom: 2rem;
    }
    
    .hero-title {
        color: var(--accent-gold);
        font-size: 2.5rem;
        font-weight: bold;
        margin-bottom: 1rem;
    }
    
    .hero-subtitle {
        color: white;
        font-size: 1.2rem;
        margin-bottom: 1.5rem;
    }
    
    .restricted-modal {
        background: linear-gradient(135deg, #18738c 0%, #1a5490 100%);
        padding: 2rem;
        border-radius: 15px;
        text-align: center;
        border: 3px solid var(--accent-gold);
    }
    
    .restricted-title {
        color: var(--accent-gold);
        font-size: 1.5rem;
        font-weight: bold;
        margin-bottom: 1rem;
    }
    
    .restricted-text {
        color: white;
        font-size: 1.1rem;
        margin-bottom: 1.5rem;
    }
    
    .nav-menu {
        display: flex;
        gap: 10px;
        flex-wrap: wrap;
        justify-content: center;
        margin-bottom: 1rem;
    }
    
    .nav-btn {
        background: transparent;
        border: none;
        color: var(--primary-blue);
        font-weight: 600;
        padding: 8px 15px;
        cursor: pointer;
        border-radius: 20px;
        transition: all 0.2s;
    }
    
    .nav-btn:hover {
        background-color: var(--light-gold);
    }
    
    .nav-btn.active {
        background-color: var(--primary-blue);
        color: white;
    }
    
    .btn-entrar {
        background-color: var(--accent-gold) !important;
        color: var(--primary-blue) !important;
        font-weight: bold !important;
        border-radius: 25px !important;
        padding: 10px 25px !important;
    }
    
    .section-title {
        color: var(--primary-blue);
        font-size: 2rem;
        font-weight: bold;
        text-align: center;
        margin: 2rem 0 1rem 0;
        border-bottom: 3px solid var(--accent-gold);
        padding-bottom: 0.5rem;
    }
    
    .about-card {
        background: white;
        padding: 2rem;
        border-radius: 15px;
        border-left: 5px solid var(--accent-gold);
        margin: 1rem 0;
    }
    
    .solution-card {
        background: linear-gradient(135deg, var(--primary-blue) 0%, #1a5490 100%);
        padding: 2rem;
        border-radius: 15px;
        text-align: center;
        color: white;
    }
    
    .solution-title {
        color: var(--accent-gold);
        font-size: 1.3rem;
        font-weight: bold;
        margin-bottom: 1rem;
    }
    
    .cta-button {
        display: inline-block;
        background-color: var(--accent-gold);
        color: var(--primary-blue);
        padding: 12px 30px;
        border-radius: 30px;
        font-weight: bold;
        text-decoration: none;
        margin-top: 1rem;
        transition: transform 0.2s;
    }
    
    .cta-button:hover {
        transform: scale(1.05);
        color: var(--primary-blue);
    }
    
    .paywall-box {
        background: linear-gradient(135deg, #18738c 0%, #1a5490 100%);
        padding: 2rem;
        border-radius: 15px;
        text-align: center;
        border: 3px solid var(--accent-gold);
        margin: 1rem 0;
    }
    
    .paywall-icon {
        font-size: 3rem;
        margin-bottom: 1rem;
    }
    
    .paywall-title {
        color: var(--accent-gold);
        font-size: 1.5rem;
        font-weight: bold;
        margin-bottom: 0.5rem;
    }
    
    .paywall-text {
        color: white;
        font-size: 1rem;
        margin-bottom: 1rem;
    }
    </style>
    """, unsafe_allow_html=True)

def vimeo_video(url):
    video_id = url.split('/')[-1].split('?')[0]
    embed_url = (
        f"https://player.vimeo.com/video/{video_id}"
        f"?dnt=1&badge=0&autopause=0&player_id=0&app_id=58479&byline=0&portrait=0&title=0"
    )
    iframe_html = f"""
    <div style="padding:56.25% 0 0 0;position:relative;margin-bottom:1rem;overflow:visible;">
        <iframe
            src="{embed_url}"
            frameborder="0"
            allow="autoplay; fullscreen; picture-in-picture; clipboard-write; encrypted-media"
            allowfullscreen
            referrerpolicy="no-referrer-when-downgrade"
            style="position:absolute;top:0;left:0;width:100%;height:100%;border-radius:8px;border:none;"
            title="LPS Video"
        ></iframe>
    </div>
    <script src="https://player.vimeo.com/api/player.js"></script>
    """
    st.markdown(iframe_html, unsafe_allow_html=True)

# Inicialização do Estado de Sessão
if 'page' not in st.session_state:
    st.session_state.page = "Home"
if 'section' not in st.session_state:
    st.session_state.section = "home"
if 'progress' not in st.session_state:
    st.session_state.progress = {}
if 'assessment_results' not in st.session_state:
    st.session_state.assessment_results = None
if 'ai_laudo' not in st.session_state:
    st.session_state.ai_laudo = None
if 'laudo_requested' not in st.session_state:
    st.session_state.laudo_requested = False
if 'employee_token' not in st.session_state:
    st.session_state.employee_token = None
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False
if 'user' not in st.session_state:
    st.session_state.user = None
if 'manager_data' not in st.session_state:
    st.session_state.manager_data = None
if 'login_mode' not in st.session_state:
    st.session_state.login_mode = "login"
if 'show_test_form' not in st.session_state:
    st.session_state.show_test_form = False
if 'show_login_modal' not in st.session_state:
    st.session_state.show_login_modal = False
if 'selected_module' not in st.session_state:
    st.session_state.selected_module = None

WHATSAPP_URL = "https://wa.me/5511971419453"
LOGO_PATH = "attached_assets/logotipo_2__1768529013163.jpeg"

# Definicoes das questoes do Assessment (8 blocos: 7x8 + 1x14 = 70 questoes totais)
ASSESSMENT_QUESTIONS = {
    "Bloco 1 – Autoridade Interna, Autoimagem e Superego": [
        "Quando recebo um feedback negativo sobre minha liderança, minha primeira reação interna costuma ser de justificativa ou defesa.",
        "Sinto uma pressão interna para corresponder a um ideal de líder perfeito, mesmo que isso me desgaste.",
        "Tenho dificuldade em admitir não sei ou pedir ajuda à minha equipe, pois temo parecer fraco(a).",
        "A opinião dos meus superiores sobre mim tem um peso significativo na minha autoconfiança como líder.",
        "Hesito em tomar decisões que possam desagradar a equipe, mesmo que sejam necessárias para o negócio.",
        "Quando um projeto sob minha liderança falha, sinto uma culpa ou responsabilidade desproporcional.",
        "Prefiro manter uma postura mais distante e formal para garantir o respeito da equipe.",
        "Fico ruminando críticas ou comentários negativos sobre meu trabalho por mais tempo do que gostaria."
    ],
    "Bloco 2 – Contenção Emocional, Empatia e Círculo de Segurança": [
        "Em momentos de alta tensão na equipe, consigo manter minha calma interna e pensar com clareza.",
        "Percebo quando minha própria ansiedade ou frustração está contaminando o clima da equipe.",
        "Sinto-me confortável em abordar conversas difíceis sobre emoções ou conflitos interpessoais no time.",
        "Consigo ouvir reclamações ou desabafos da equipe sem me sentir pessoalmente atacado(a) ou sobrecarregado(a).",
        "Priorizo ativamente a criação de um ambiente onde as pessoas se sintam seguras para expressar opiniões divergentes ou admitir erros.",
        "Tenho facilidade em ler o estado emocional não verbal da equipe (linguagem corporal, tom de voz).",
        "Após reuniões emocionalmente carregadas, preciso de um tempo considerável para me recuperar energeticamente.",
        "Hesito em dar más notícias ou feedbacks corretivos por medo da reação emocional da equipe."
    ],
    "Bloco 3 – Narcisismo, Reconhecimento e Motivação": [
        "Sinto uma satisfação particular quando minhas ideias ou projetos recebem destaque e admiração.",
        "Fico genuinamente incomodado(a) quando o crédito por um trabalho meu é dado a outra pessoa.",
        "A ausência de feedback positivo sobre meu desempenho me desmotiva significativamente.",
        "Tenho uma tendência a comparar meus resultados e reconhecimento com os de outros líderes que admiro.",
        "Em discussões, foco mais em defender meu ponto de vista do que em entender a perspectiva do outro.",
        "Busco ativamente projetos ou tarefas que me coloquem em evidência ou aumentem minha visibilidade.",
        "Tenho dificuldade em celebrar o sucesso de outros líderes, especialmente se me sinto em competição.",
        "Uso histórias sobre meus próprios sucessos para inspirar ou motivar a equipe."
    ],
    "Bloco 4 – Estrutura, Controle e Tolerância à Ambiguidade": [
        "Sinto-me desconfortável ou ansioso(a) quando os planos mudam inesperadamente ou há falta de clareza nas diretrizes.",
        "Prefiro ter controle sobre os detalhes das tarefas da equipe para garantir que tudo saia conforme o esperado.",
        "Tenho dificuldade em delegar tarefas críticas, preferindo fazê-las eu mesmo(a) para assegurar a qualidade.",
        "Acredito que regras e processos bem definidos são a melhor forma de evitar o caos e garantir a produtividade.",
        "Intervenho rapidamente quando percebo que um membro da equipe está se desviando do processo estabelecido.",
        "Valorizo mais a previsibilidade e a estabilidade do que a flexibilidade e a experimentação.",
        "Sinto uma necessidade de preencher o vazio com tarefas, mesmo quando não há urgência real.",
        "Foco mais em identificar e prevenir riscos do que em explorar novas oportunidades que envolvam incerteza."
    ],
    "Bloco 5 – Dinâmicas Relacionais, Transferência e Contratransferência": [
        "Percebo que certos tipos de pessoas na equipe despertam em mim reações emocionais (positivas ou negativas) mais intensas que outras.",
        "Já me peguei tratando um membro da equipe de forma semelhante a como lidei com figuras importantes do meu passado (pais, irmãos, professores).",
        "Sinto que alguns membros da equipe me colocam em um papel específico (o salvador, o crítico, o idealizado) que não corresponde totalmente à realidade.",
        "Tenho dificuldade em manter a neutralidade emocional com membros da equipe que me lembram pessoas com quem tive conflitos no passado.",
        "Noto padrões repetitivos na forma como diferentes membros da equipe se relacionam comigo (ex: sempre buscando aprovação, sempre desafiando).",
        "Minhas reações a um determinado funcionário às vezes me surpreendem pela sua intensidade ou inadequação.",
        "Sinto necessidade de manter uma distância segura de certos membros da equipe para não me envolver emocionalmente.",
        "Reflito sobre como minhas próprias experiências e fantasmas do passado podem estar influenciando minhas interações atuais com a equipe."
    ],
    "Bloco 6 – Autoconsciência, Mentalização e Defesas do Ego": [
        "Consigo identificar quando estou usando um mecanismo de defesa (ex: racionalização, negação, projeção) para lidar com uma situação desconfortável.",
        "Sou capaz de analisar minhas próprias motivações (conscientes e talvez inconscientes) por trás de uma decisão importante.",
        "Busco ativamente feedback sobre meus pontos cegos, mesmo que seja desconfortável ouvi-los.",
        "Consigo diferenciar uma crítica ao meu comportamento/decisão de um ataque pessoal à minha identidade.",
        "Reflito sobre o impacto do meu estado emocional (cansaço, estresse, empolgação) na minha forma de liderar.",
        "Estou aberto(a) a revisar minhas crenças e meu estilo de liderança com base em novas experiências ou aprendizados.",
        "Consigo mentalizar, ou seja, imaginar o que se passa na mente do outro (seus pensamentos, sentimentos, intenções) sem projetar excessivamente minhas próprias ideias.",
        "Reconheço que minha percepção da realidade é subjetiva e pode ser influenciada por meus vieses e história pessoal."
    ],
    "Bloco 7 – Ação, Urgência e Foco na Entrega": [
        "Quando há indecisão na equipe, sinto um forte impulso para intervir e dar a direção final imediatamente.",
        "Minha principal métrica de sucesso na liderança é a capacidade de fazer as coisas acontecerem de forma rápida.",
        "Prefiro resolver um problema com ação direta do que gastar tempo excessivo em análise ou discussões emocionais.",
        "Tomo decisões críticas rapidamente, mesmo que isso implique risco ou desagrade alguns membros da equipe.",
        "Sinto uma frustração intensa quando a equipe demonstra lentidão, passividade ou falta de senso de urgência.",
        "Sou mais eficaz em situações de crise ou pressão, onde minha assertividade e foco são necessários.",
        "Valorizo mais a autonomia e a liberdade de ação do que a necessidade de seguir processos rígidos.",
        "Tenho dificuldade em delegar tarefas de alta prioridade, preferindo a garantia da minha própria execução."
    ],
    "Bloco 8 – Temperamento e Caráter (Cloninger)": [
        "Sinto vontade de experimentar coisas novas regularmente. (Busca de Novidade - BN)",
        "Gosto de desafios e situações não previsíveis. (Busca de Novidade - BN)",
        "Sinto ansiedade em situações de risco ou incerteza. (Evitação de Dano - ED)",
        "Evito decisões que possam gerar críticas ou falhas. (Evitação de Dano - ED)",
        "Busco apoio e aprovação de outros antes de decidir. (Dependência de Recompensa - DR)",
        "Sinto motivação extra quando sou reconhecido publicamente. (Dependência de Recompensa - DR)",
        "Continuo esforços mesmo diante de frustrações. (Persistência - P)",
        "Finalizo tarefas mesmo quando enfrentam resistência. (Persistência - P)",
        "Tenho clareza de objetivos e valores que guiam decisões. (Autodireção - AD)",
        "Assumo responsabilidade pelos resultados, bons ou ruins. (Autodireção - AD)",
        "Tento entender perspectiva de todos antes de agir. (Cooperatividade - C)",
        "Priorizo harmonia e colaboração sobre competição. (Cooperatividade - C)",
        "Busco propósito maior no que faço. (Autotranscendência - AT)",
        "Sinto-me conectado(a) a algo maior que minhas tarefas diárias. (Autotranscendência - AT)"
    ]
}

EMPLOYEE_ASSESSMENT_QUESTIONS = {
    "Bloco 1 – Idealista Exigente (Funcionário)": [
        "Busco constantemente a excelência em tudo que faço, mesmo que isso exija mais esforço da equipe.",
        "Tenho dificuldade em aceitar trabalhos ou resultados que não atendam aos meus altos padrões de qualidade.",
        "Quando percebo um erro ou falha no grupo, sinto a necessidade de apontá-lo para que seja corrigido.",
        "Frequentemente me vejo tentando inspirar a equipe a alcançar um objetivo maior ou uma visão mais ambiciosa.",
        "Sinto-me frustrado(a) quando o grupo parece satisfeito com resultados medianos ou 'bons o suficiente'.",
        "Acredito que a disciplina e o rigor são fundamentais para o sucesso de qualquer projeto em equipe.",
        "Se sinto que minhas ideias ou padrões não são valorizados, posso me tornar crítico(a) em relação ao progresso do grupo.",
        "Prefiro trabalhar com pessoas que compartilham meu nível de exigência e comprometimento.",
        "Em discussões de grupo, foco mais nos padrões e na qualidade do que nas relações interpessoais.",
        "Às vezes, sinto que minhas críticas, mesmo bem-intencionadas, me isolam ou me tornam alvo de ressentimento no grupo.",
        "Tenho uma visão clara de como as coisas deveriam ser feitas e tento guiar o grupo nessa direção.",
        "A mediocridade ou a falta de ambição no grupo me incomodam profundamente."
    ],
    "Bloco 2 – Contenedor Empático (Funcionário)": [
        "Sou frequentemente procurado(a) por colegas para desabafar ou buscar apoio emocional.",
        "Percebo facilmente o clima emocional do grupo e tento agir para torná-lo mais harmonioso.",
        "Priorizo o bem-estar e as boas relações na equipe, mesmo que isso signifique ceder em alguns pontos.",
        "Sinto-me responsável por ajudar colegas que estão passando por dificuldades ou estresse.",
        "Tenho facilidade em ouvir diferentes pontos de vista e mediar conflitos na equipe.",
        "Em momentos de tensão no grupo, minha primeira reação é tentar acalmar os ânimos.",
        "Às vezes, sinto que absorvo demais os problemas ou a negatividade dos outros no ambiente de trabalho.",
        "Acredito que um ambiente de segurança psicológica é essencial para a produtividade da equipe.",
        "Consigo manter a calma e oferecer suporte mesmo quando o grupo está sob pressão.",
        "Preocupo-me mais em garantir que todos se sintam incluídos e ouvidos do que em impor minha própria opinião.",
        "Evito confrontos diretos e busco soluções que agradem a maioria das pessoas no grupo.",
        "Sinto-me energizado(a) ao ajudar a equipe a superar desafios emocionais ou relacionais."
    ],
    "Bloco 3 – Buscador de Reconhecimento (Funcionário)": [
        "Gosto quando minhas contribuições são publicamente reconhecidas pelo líder ou pela equipe.",
        "Busco oportunidades para apresentar ideias ou liderar iniciativas que me coloquem em evidência.",
        "Sinto-me motivado(a) por elogios e pela admiração dos meus colegas e superiores.",
        "Comparo meu desempenho com o dos outros e me esforço para ser um dos melhores da equipe.",
        "Fico desapontado(a) ou ressentido(a) quando sinto que meu trabalho não recebe a atenção que merece.",
        "Tenho facilidade em 'vender' minhas ideias e influenciar a opinião do grupo a meu favor.",
        "Às vezes, posso exagerar um pouco minhas conquistas para causar uma boa impressão.",
        "A ideia de passar despercebido(a) ou ser apenas 'mais um' no grupo me incomoda.",
        "Se sinto que outro colega está recebendo mais destaque, posso sentir inveja ou tentar minimizar suas conquistas.",
        "Acredito que minha presença e minhas ideias agregam um valor especial à equipe.",
        "Em reuniões, faço questão de participar ativamente para que minhas opiniões sejam ouvidas.",
        "Já fui criticado(a) por parecer buscar demais os holofotes ou por ser competitivo(a) demais."
    ],
    "Bloco 4 – Estruturador Cauteloso (Funcionário)": [
        "Sinto-me mais confortável quando há regras claras, processos definidos e um plano bem estruturado.",
        "Preocupo-me com os riscos e prefiro analisar cuidadosamente antes de tomar decisões ou apoiar novas ideias.",
        "Tenho o hábito de verificar se os procedimentos estão sendo seguidos corretamente pela equipe.",
        "Resisto a mudanças repentinas ou a abordagens que me pareçam desorganizadas ou impulsivas.",
        "Prefiro métodos comprovados e seguros a experimentar soluções inovadoras, mas incertas.",
        "Em situações de caos ou incerteza no grupo, tento trazer ordem e estrutura para a discussão.",
        "A falta de planejamento ou a ambiguidade nas tarefas me gera ansiedade.",
        "Valorizo a estabilidade e a previsibilidade no ambiente de trabalho.",
        "Quando surgem conflitos, foco em encontrar soluções práticas e racionais, evitando discussões emocionais.",
        "Sou detalhista e atento(a) a possíveis problemas ou inconsistências nos planos da equipe.",
        "Posso ser visto(a) como alguém resistente a novas ideias ou que 'coloca dificuldade' em propostas mais ousadas.",
        "Acredito que seguir as regras é a melhor forma de evitar problemas e garantir a eficiência do grupo."
    ],
    "Bloco 5 – Relacional Reativo (Funcionário)": [
        "Sou muito sensível ao clima interpessoal e percebo facilmente quando há tensões não ditas no grupo.",
        "Preocupo-me bastante com o que os outros pensam de mim e evito fazer coisas que possam desagradar.",
        "Tenho dificuldade em dizer 'não' ou em expressar opiniões que possam gerar conflito.",
        "Adapto meu comportamento para me encaixar melhor no grupo ou para evitar ser excluído(a).",
        "Em situações de conflito, tendo a ficar quieto(a) ou a concordar com a maioria para evitar problemas.",
        "Sinto-me desconfortável ou ansioso(a) quando percebo críticas ou desaprovação, mesmo que indiretas.",
        "Já me senti injustiçado(a) ou culpabilizado(a) por problemas do grupo, mesmo sem ter responsabilidade direta.",
        "Prefiro seguir o fluxo do grupo a tentar impor minhas próprias ideias ou vontades.",
        "Se me sinto ameaçado(a) ou desvalorizado(a), posso me retrair ou expressar meu descontentamento de forma indireta (ex: fofoca, atrasos).",
        "Busco constantemente a aprovação e a aceitação dos meus colegas e do meu líder.",
        "Tenho medo de ser mal interpretado(a) ou de causar uma má impressão no grupo.",
        "Às vezes, sinto que minhas próprias necessidades ficam em segundo plano para manter a harmonia do grupo."
    ],
    "Bloco 6 – Observador Consciente (Funcionário)": [
        "Prefiro analisar as situações de forma lógica e objetiva antes de me envolver emocionalmente.",
        "Mantenho um certo distanciamento emocional para conseguir avaliar melhor as dinâmicas do grupo.",
        "Frequentemente ofereço insights ou análises que ajudam o grupo a entender melhor um problema.",
        "Tenho dificuldade em lidar com explosões emocionais ou dramas interpessoais no ambiente de trabalho.",
        "Baseio minhas opiniões e decisões em fatos e dados, mais do que em intuição ou sentimentos.",
        "Posso ser percebido(a) como alguém frio(a) ou indiferente às questões emocionais do grupo.",
        "Em discussões acaloradas, costumo observar mais do que participar ativamente.",
        "Acredito que a lógica e a razão são as melhores ferramentas para resolver os problemas da equipe.",
        "Avalio a coerência e a lógica dos argumentos dos outros, apontando falhas quando as percebo.",
        "Não me deixo levar facilmente pelo 'entusiasmo' ou pela 'pressão' do grupo.",
        "Meu distanciamento às vezes é interpretado como falta de comprometimento ou crítica velada.",
        "Busco entender as causas raízes dos problemas do grupo, em vez de focar apenas nos sintomas."
    ],
    "Bloco 7 – Executor Decidido (Funcionário)": [
        "Em vez de ficar analisando, prefiro partir logo para a ação e corrigir o curso se for preciso.",
        "Quando há indecisão na equipe, eu tomo a frente e dou a direção.",
        "Sinto-me mais motivado(a) quando tenho autonomia total para decidir o como da execução.",
        "Tenho facilidade em tomar decisões rápidas, mesmo com informações incompletas.",
        "Em uma crise, meu foco principal é resolver o problema e garantir o resultado, deixando as emoções para depois.",
        "Sou visto(a) como alguém assertivo(a) e direto(a) na comunicação, sem rodeios.",
        "A lentidão ou a passividade de um colega me irritam profundamente e me fazem querer agir por ele.",
        "Priorizo o resultado final e o cumprimento da meta acima de quaisquer burocracias ou detalhes.",
        "Não tenho medo de confrontar um problema ou uma pessoa se isso for necessário para alcançar o objetivo.",
        "Sinto-me energizado(a) pela pressão e pelos desafios que exigem uma reação imediata.",
        "Minha principal contribuição para a equipe é a capacidade de fazer as coisas acontecerem.",
        "Acredito que o tempo é um recurso valioso e me esforço para manter um alto senso de urgência."
    ],
    "Bloco 8 – Temperamento e Caráter (Cloninger Funcionário)": [
        "Busco frequentemente experiências novas e estímulos variados (BN).",
        "Sinto ansiedade em situações de risco ou desconhecidas (ED).",
        "Procuro recompensas externas e aceitação social em minhas decisões (DR).",
        "Persevero em tarefas mesmo quando difíceis ou frustrantes (P).",
        "Estabeleço metas claras e sigo princípios próprios de forma consistente (AD).",
        "Colaboro com os outros e sou compreensivo(a) mesmo em conflitos (C).",
        "Sinto conexão com algo maior ou propósito transcendente (AT).",
        "Evito riscos desnecessários mesmo que perca oportunidades (ED).",
        "Procuro desafios e novas experiências mesmo que envolvam riscos (BN).",
        "Tenho dificuldade em cooperar quando meus objetivos individuais estão em jogo (C).",
        "Ajusto meu comportamento para agradar os outros ou obter aprovação (DR).",
        "Sigo minhas convicções pessoais mesmo quando não reconhecido pelos outros (AD).",
        "Sou flexível e me adapto a mudanças e incertezas (P).",
        "Sinto que minha vida tem um sentido maior ou espiritual (AT)."
    ]
}

EMPLOYEE_BLOCK_TO_PROFILE = {
    "Bloco 1 – Idealista Exigente (Funcionário)": "O Idealista Exigente",
    "Bloco 2 – Contenedor Empático (Funcionário)": "O Contenedor Empático",
    "Bloco 3 – Buscador de Reconhecimento (Funcionário)": "O Buscador de Reconhecimento",
    "Bloco 4 – Estruturador Cauteloso (Funcionário)": "O Estruturador Cauteloso",
    "Bloco 5 – Relacional Reativo (Funcionário)": "O Relacional Reativo",
    "Bloco 6 – Observador Consciente (Funcionário)": "O Observador Consciente",
    "Bloco 7 – Executor Decidido (Funcionário)": "O Executor Decidido",
    "Bloco 1 – Idealista Exigente (Funcionario)": "O Idealista Exigente",
    "Bloco 2 – Contenedor Empatico (Funcionario)": "O Contenedor Empático",
    "Bloco 3 – Buscador de Reconhecimento (Funcionario)": "O Buscador de Reconhecimento",
    "Bloco 4 – Estruturador Cauteloso (Funcionario)": "O Estruturador Cauteloso",
    "Bloco 5 – Relacional Reativo (Funcionario)": "O Relacional Reativo",
    "Bloco 6 – Observador Consciente (Funcionario)": "O Observador Consciente",
    "Bloco 7 – Executor Decidido (Funcionario)": "O Executor Decidido"
}

EMPLOYEE_BLOCK_SHORT_NAMES = {
    "Bloco 1 – Idealista Exigente (Funcionário)": "Autoridade",
    "Bloco 2 – Contenedor Empático (Funcionário)": "Contenção",
    "Bloco 3 – Buscador de Reconhecimento (Funcionário)": "Narcisismo",
    "Bloco 4 – Estruturador Cauteloso (Funcionário)": "Estrutura",
    "Bloco 5 – Relacional Reativo (Funcionário)": "Relacional",
    "Bloco 6 – Observador Consciente (Funcionário)": "Observação",
    "Bloco 7 – Executor Decidido (Funcionário)": "Execução",
    "Bloco 1 – Idealista Exigente (Funcionario)": "Autoridade",
    "Bloco 2 – Contenedor Empatico (Funcionario)": "Contenção",
    "Bloco 3 – Buscador de Reconhecimento (Funcionario)": "Narcisismo",
    "Bloco 4 – Estruturador Cauteloso (Funcionario)": "Estrutura",
    "Bloco 5 – Relacional Reativo (Funcionario)": "Relacional",
    "Bloco 6 – Observador Consciente (Funcionario)": "Observação",
    "Bloco 7 – Executor Decidido (Funcionario)": "Execução"
}

EMPLOYEE_CLONINGER_SUBDIMENSIONS = {
    "BN": {"name": "Busca de Novidade", "questions": [0, 8]},
    "ED": {"name": "Evitação de Dano", "questions": [1, 7]},
    "DR": {"name": "Dependência de Recompensa", "questions": [2, 10]},
    "P": {"name": "Persistência", "questions": [3, 12]},
    "AD": {"name": "Autodireção", "questions": [4, 11]},
    "C": {"name": "Cooperatividade", "questions": [5, 9]},
    "AT": {"name": "Autotranscendência", "questions": [6, 13]}
}

ARCHETYPE_TO_SIGLA = {
    "O Idealista Exigente": "IE",
    "O Contenedor Empático": "CE",
    "O Contenedor Empatico": "CE",
    "O Buscador de Reconhecimento": "BR",
    "O Estruturador Cauteloso": "EC",
    "O Relacional Reativo": "RR",
    "O Observador Consciente": "OC",
    "O Executor Decidido": "ED"
}

CLONINGER_SUBDIMENSIONS = {
    "BN": {"name": "Busca de Novidade", "questions": [0, 1]},
    "ED": {"name": "Evitação de Dano", "questions": [2, 3]},
    "DR": {"name": "Dependência de Recompensa", "questions": [4, 5]},
    "P": {"name": "Persistência", "questions": [6, 7]},
    "AD": {"name": "Autodireção", "questions": [8, 9]},
    "C": {"name": "Cooperatividade", "questions": [10, 11]},
    "AT": {"name": "Autotranscendência", "questions": [12, 13]}
}

BLOCK_SHORT_NAMES = {
    "Bloco 1 – Autoridade Interna, Autoimagem e Superego": "Autoridade",
    "Bloco 2 – Contenção Emocional, Empatia e Círculo de Segurança": "Contenção",
    "Bloco 3 – Narcisismo, Reconhecimento e Motivação": "Narcisismo",
    "Bloco 4 – Estrutura, Controle e Tolerância à Ambiguidade": "Estrutura",
    "Bloco 5 – Dinâmicas Relacionais, Transferência e Contratransferência": "Relacional",
    "Bloco 6 – Autoconsciência, Mentalização e Defesas do Ego": "Observação",
    "Bloco 7 – Ação, Urgência e Foco na Entrega": "Execução",
    "Bloco 2 – Contencao Emocional, Empatia e Circulo de Seguranca": "Contenção",
    "Bloco 3 – Narcisismo, Reconhecimento e Motivacao": "Narcisismo",
    "Bloco 4 – Estrutura, Controle e Tolerancia a Ambiguidade": "Estrutura",
    "Bloco 5 – Dinamicas Relacionais, Transferencia e Contratransferencia": "Relacional",
    "Bloco 6 – Autoconsciencia, Mentalizacao e Defesas do Ego": "Observação",
    "Bloco 7 – Acao, Urgencia e Foco na Entrega": "Execução"
}

BLOCK_TO_PROFILE = {
    "Bloco 1 – Autoridade Interna, Autoimagem e Superego": "O Idealista Exigente",
    "Bloco 2 – Contenção Emocional, Empatia e Círculo de Segurança": "O Contenedor Empático",
    "Bloco 3 – Narcisismo, Reconhecimento e Motivação": "O Buscador de Reconhecimento",
    "Bloco 4 – Estrutura, Controle e Tolerância à Ambiguidade": "O Estruturador Cauteloso",
    "Bloco 5 – Dinâmicas Relacionais, Transferência e Contratransferência": "O Relacional Reativo",
    "Bloco 6 – Autoconsciência, Mentalização e Defesas do Ego": "O Observador Consciente",
    "Bloco 7 – Ação, Urgência e Foco na Entrega": "O Executor Decidido",
    "Bloco 2 – Contencao Emocional, Empatia e Circulo de Seguranca": "O Contenedor Empático",
    "Bloco 3 – Narcisismo, Reconhecimento e Motivacao": "O Buscador de Reconhecimento",
    "Bloco 4 – Estrutura, Controle e Tolerancia a Ambiguidade": "O Estruturador Cauteloso",
    "Bloco 5 – Dinamicas Relacionais, Transferencia e Contratransferencia": "O Relacional Reativo",
    "Bloco 6 – Autoconsciencia, Mentalizacao e Defesas do Ego": "O Observador Consciente",
    "Bloco 7 – Acao, Urgencia e Foco na Entrega": "O Executor Decidido"
}

def classify_bion_role(block_sums):
    """
    Classifica o papel grupal segundo Bion (max 40 pontos por bloco, 8 questoes x 5):
    - Porta-voz: Alta Autoridade + Alta Observacao (expressa o que o grupo sente)
    - Bode Expiatorio: Alto Relacional + Baixa Contencao (absorve projecoes negativas)
    - Dependente: Alta Contencao + Baixa Autoridade (busca proteção no lider)
    - Lider de Luta-Fuga: Alto Narcisismo + Alta Execucao (reativo a ameacas)
    - Sabotador Silencioso: Alta Estrutura + Baixa Observacao (resiste passivamente)
    - Harmonizador Ansioso: Alto Relacional + Alta Contencao (evita conflitos a todo custo)
    """
    autoridade = block_sums.get("Autoridade", 20)
    contencao = block_sums.get("Contencao", 20)
    narcisismo = block_sums.get("Narcisismo", 20)
    estrutura = block_sums.get("Estrutura", 20)
    relacional = block_sums.get("Relacional", 20)
    observacao = block_sums.get("Observacao", 20)
    execucao = block_sums.get("Execucao", 20)
    
    high = 28
    low = 20
    
    if relacional >= high and contencao >= high:
        return "Harmonizador Ansioso"
    elif autoridade >= high and observacao >= high:
        return "Porta-voz"
    elif relacional >= high and contencao <= low:
        return "Bode Expiatorio"
    elif contencao >= high and autoridade <= low:
        return "Dependente"
    elif narcisismo >= high and execucao >= high:
        return "Lider de Luta-Fuga"
    elif estrutura >= high and observacao <= low:
        return "Sabotador Silencioso"
    else:
        return "Neutro/Adaptavel"

BION_DESCRIPTIONS = {
    "Porta-voz": "Expressa verbalmente o que o grupo sente mas nao consegue dizer. Canaliza tensoes coletivas.",
    "Bode Expiatorio": "Absorve projecoes negativas do grupo. Frequentemente culpado por falhas sistemicas.",
    "Dependente": "Busca proteção e direcao no lider. Evita autonomia e delega responsabilidade emocional.",
    "Lider de Luta-Fuga": "Reativo a ameacas reais ou imaginarias. Mobiliza o grupo para atacar ou fugir.",
    "Sabotador Silencioso": "Resiste passivamente as mudancas. Cumpre tarefas sem engajamento emocional.",
    "Harmonizador Ansioso": "Evita conflitos a todo custo. Busca manter harmonia mesmo sacrificando a autenticidade.",
    "Neutro/Adaptavel": "Perfil equilibrado. Adapta-se as necessidades do grupo sem assumir papel fixo."
}

def classify_employee_bion_role(block_sums):
    """
    Classifica o papel grupal do funcionario segundo Bion (max 60 pontos por bloco, 12 questoes x 5):
    5 papeis grupais para funcionarios:
    - Bode Expiatorio: Alto Relacional + Baixa Contencao (absorve culpa do grupo)
    - Sabotador: Alta Estrutura + Baixa Observacao (resiste passivamente)
    - Lider Informal: Alto Narcisismo + Alta Execucao (assume lideranca sem cargo)
    - Patrulheiro: Alta Autoridade + Alta Observacao (vigia normas e comportamentos)
    - Apaziguador: Alta Contencao + Alto Relacional (evita conflitos a todo custo)
    """
    autoridade = block_sums.get("Autoridade", 30)
    contencao = block_sums.get("Contencao", 30)
    narcisismo = block_sums.get("Narcisismo", 30)
    estrutura = block_sums.get("Estrutura", 30)
    relacional = block_sums.get("Relacional", 30)
    observacao = block_sums.get("Observacao", 30)
    execucao = block_sums.get("Execucao", 30)
    
    high = 42
    low = 30
    
    if contencao >= high and relacional >= high:
        return "Apaziguador"
    elif autoridade >= high and observacao >= high:
        return "Patrulheiro"
    elif relacional >= high and contencao <= low:
        return "Bode Expiatorio"
    elif narcisismo >= high and execucao >= high:
        return "Lider Informal"
    elif estrutura >= high and observacao <= low:
        return "Sabotador"
    else:
        return "Neutro/Adaptavel"

EMPLOYEE_BION_DESCRIPTIONS = {
    "Bode Expiatorio": "Absorve as projecoes negativas do grupo. Frequentemente culpado por falhas que sao sistemicas.",
    "Sabotador": "Resiste passivamente as mudancas e decisoes. Cumpre tarefas sem engajamento real.",
    "Lider Informal": "Assume papel de lideranca sem ter o cargo formal. Mobiliza o grupo por carisma ou competencia.",
    "Patrulheiro": "Vigia as normas e comportamentos do grupo. Aponta desvios e cobra conformidade.",
    "Apaziguador": "Evita conflitos a todo custo. Busca manter a harmonia mesmo sacrificando a autenticidade.",
    "Neutro/Adaptavel": "Perfil equilibrado. Adapta-se as necessidades do grupo sem assumir papel fixo."
}

PROFILES_DB = {
    "O Idealista Exigente": {
        "O Contenedor Empatico": {
            "forcas": "Forca: Combina a busca por excelencia com sensibilidade emocional, criando padroes elevados sem perder a humanidade. Forca: Consegue inspirar a equipe a dar o melhor de si enquanto mantem um ambiente acolhedor. Forca: Equilibra a exigencia tecnica com o cuidado genuino pelas pessoas.",
            "riscos": "Risco: Pode se esgotar emocionalmente ao tentar manter padroes altissimos enquanto absorve as emocoes do grupo. Risco: Tende a se culpar duplamente quando algo falha, tanto pela qualidade quanto pelo bem-estar da equipe. Risco: Pode adiar feedbacks criticos por medo de ferir os outros, comprometendo a excelencia que busca.",
            "recomendacoes": "Acao: Estabeleca limites claros entre cuidar da equipe e manter seus padroes, sem tentar ser perfeito em ambos. Acao: Pratique o feedback direto e compassivo, entendendo que honestidade e uma forma de cuidado. Acao: Reserve momentos de autocuidado para evitar o esgotamento que vem da dupla demanda interna."
        },
        "O Buscador de Reconhecimento": {
            "forcas": "Forca: Une a busca por excelencia com o carisma e a capacidade de engajar os outros em torno de metas ambiciosas. Forca: A necessidade de reconhecimento impulsiona a entrega de resultados de alta qualidade. Forca: Consegue transformar padroes elevados em vitrines de competencia que inspiram a equipe.",
            "riscos": "Risco: Pode se tornar perfeccionista movido pela vaidade, buscando excelencia apenas para ser admirado. Risco: Tende a competir com subordinados ou pares quando sente que seu merito nao esta sendo reconhecido. Risco: A frustacao com a falta de reconhecimento pode amplificar a autocritica ja intensa do superego.",
            "recomendacoes": "Acao: Diferencie entre buscar excelencia por valores internos e buscar aprovacao externa, priorizando o primeiro. Acao: Pratique reconhecer publicamente as contribuicoes dos outros, fortalecendo a equipe e sua propria lideranca. Acao: Desenvolva fontes internas de validacao para reduzir a dependencia de aplausos externos."
        },
        "O Estruturador Cauteloso": {
            "forcas": "Forca: Combina altos padroes de qualidade com planejamento meticuloso e organizacao impecavel. Forca: Cria sistemas e processos que sustentam a excelencia de forma consistente e previsivel. Forca: A atenção ao detalhe do estruturador potencializa a busca por perfeicao do idealista.",
            "riscos": "Risco: Pode criar ambientes rigidos demais onde a busca por perfeicao se transforma em burocracia paralisante. Risco: Tende a microgerenciar obsessivamente cada etapa, gerando lentidao e desmotivacao na equipe. Risco: A combinacao de perfeccionismo com necessidade de controle pode impedir a inovacao e a adaptacao.",
            "recomendacoes": "Acao: Defina quais processos realmente precisam de controle rigoroso e quais podem ser mais flexiveis. Acao: Aceite que nem tudo precisa ser perfeito e estruturado ao mesmo tempo, priorizando o que gera mais impacto. Acao: Delegue etapas do processo confiando na competencia da equipe, mesmo que o resultado nao seja identico ao seu padrao."
        },
        "O Relacional Reativo": {
            "forcas": "Forca: A sensibilidade relacional permite perceber como as exigencias elevadas afetam cada membro da equipe. Forca: Consegue ajustar o nivel de cobranca de acordo com o momento emocional das relacoes no grupo. Forca: A consciencia das dinamicas interpessoais humaniza a busca por excelencia.",
            "riscos": "Risco: Pode ser ativado emocionalmente quando alguem nao atende seus padroes, reagindo de forma desproporcional. Risco: Tende a projetar nas relacoes de trabalho padroes de exigencia vindos de figuras de autoridade do passado. Risco: A reatividade emocional pode transformar feedbacks construtivos em confrontos pessoais.",
            "recomendacoes": "Acao: Antes de reagir a uma entrega abaixo do esperado, faca uma pausa e avalie se a intensidade da sua reacao e proporcional. Acao: Identifique quais relacoes profissionais ativam seus padroes de exigencia mais primitivos e trabalhe essas transferencias. Acao: Desenvolva protocolos de feedback estruturado para evitar que emocoes relacionais contaminem avaliacoes tecnicas."
        },
        "O Observador Consciente": {
            "forcas": "Forca: A autoconsciencia permite identificar quando a busca por perfeicao se torna autodestrutiva ou improdutiva. Forca: Consegue refletir sobre seus proprios padroes de exigencia e ajusta-los com maturidade. Forca: A capacidade de mentalizacao transforma a autocritica em crescimento pessoal e profissional.",
            "riscos": "Risco: Pode cair em ciclos de auto-analise paralisante, pensando demais sobre seus padroes em vez de agir. Risco: Tende a intelectualizar a busca por perfeicao sem conseguir modifica-la na pratica. Risco: A hiperreflexao sobre falhas pode amplificar a pressao interna do superego em vez de alivia-la.",
            "recomendacoes": "Acao: Use a autoconsciencia como ferramenta de acao, nao apenas de reflexao, definindo mudancas concretas a cada insight. Acao: Estabeleca um limite de tempo para a auto-analise e depois mude para o modo de execucao. Acao: Compartilhe suas reflexoes com um mentor ou coach para evitar o isolamento do pensamento circular."
        },
        "O Executor Decidido": {
            "forcas": "Forca: A combinacao de altos padroes com velocidade de execucao gera resultados excepcionais em prazos curtos. Forca: A decisividade do executor complementa a profundidade analitica do idealista, equilibrando pensamento e acao. Forca: Consegue transformar padroes elevados em entregas concretas sem ficar preso na paralisia da perfeicao.",
            "riscos": "Risco: Pode gerar entregas rapidas mas que nao atendem seus proprios padroes, causando frustracao e retrabalho. Risco: A tensao entre querer perfeicao e querer velocidade pode criar ansiedade cronica e estresse elevado. Risco: Tende a pressionar a equipe com prazos apertados e padroes altissimos simultaneamente, gerando burnout.",
            "recomendacoes": "Acao: Defina com clareza quando a velocidade e mais importante que a perfeicao e vice-versa, caso a caso. Acao: Crie categorias de entrega: rapida e funcional versus refinada e excelente, comunicando isso a equipe. Acao: Monitore os sinais de estresse em si mesmo e na equipe quando estiver operando em modo de alta exigencia e alta velocidade."
        }
    },
    "O Contenedor Empatico": {
        "O Idealista Exigente": {
            "forcas": "Forca: Cria ambientes seguros onde as pessoas se sentem apoiadas para buscar altos padroes de desempenho. Forca: A empatia permite comunicar expectativas elevadas de forma que motiva em vez de intimidar. Forca: Consegue sustentar emocionalmente a equipe durante periodos de pressao por resultados.",
            "riscos": "Risco: Pode absorver a ansiedade da equipe gerada por metas exigentes, esgotando-se emocionalmente. Risco: Tende a flexibilizar padroes para proteger os sentimentos dos outros, comprometendo a qualidade. Risco: A pressao interna por excelencia pode entrar em conflito com o desejo de acolher, gerando ambivalencia.",
            "recomendacoes": "Acao: Reconheca que manter padroes elevados e uma forma de respeitar o potencial da equipe, nao uma agressao. Acao: Desenvolva a capacidade de ser firme e acolhedor ao mesmo tempo, praticando a assertividade empatica. Acao: Crie rituais de descompressao pessoal para nao acumular a carga emocional do grupo."
        },
        "O Buscador de Reconhecimento": {
            "forcas": "Forca: A empatia combinada com o carisma cria uma lideranca magnetica que atrai e mantem talentos. Forca: Consegue fazer cada pessoa se sentir vista e valorizada, fortalecendo o engajamento da equipe. Forca: O desejo de reconhecimento motiva a criar experiencias positivas que sao lembradas e celebradas.",
            "riscos": "Risco: Pode usar a empatia como estrategia para ganhar admiracao, perdendo a autenticidade do cuidado. Risco: Tende a priorizar ser amado em vez de ser respeitado, evitando decisoes impopulares mas necessarias. Risco: A necessidade de validacao pode leva-lo a se sobrecarregar com as demandas emocionais dos outros para ser visto como indispensavel.",
            "recomendacoes": "Acao: Verifique regularmente se seu cuidado com os outros e genuino ou se esta servindo a uma necessidade de ser admirado. Acao: Pratique tomar decisoes dificeis mesmo quando isso significa desagradar temporariamente. Acao: Busque reconhecimento por impacto real e nao apenas por simpatia, valorizando resultados alem de relacoes."
        },
        "O Estruturador Cauteloso": {
            "forcas": "Forca: Combina seguranca emocional com seguranca estrutural, criando ambientes altamente estaveis para a equipe. Forca: Os processos organizados permitem que o cuidado emocional seja sistematico e nao apenas reativo. Forca: A previsibilidade estrutural reduz a ansiedade do grupo, potencializando o efeito contenedor.",
            "riscos": "Risco: Pode criar ambientes tao protegidos e controlados que impedem o crescimento e a autonomia da equipe. Risco: Tende a confundir seguranca emocional com ausencia de desafios, superprotegendo os liderados. Risco: A rigidez dos processos pode limitar a espontaneidade emocional necessaria para conexoes genuinas.",
            "recomendacoes": "Acao: Permita que a equipe enfrente desafios graduais, oferecendo suporte sem eliminar todas as dificuldades. Acao: Use a estrutura como base de seguranca, nao como gaiola, mantendo flexibilidade para adaptacao. Acao: Reserve espacos nao-estruturados para conversas espontaneas que fortalecem vinculos genuinos."
        },
        "O Relacional Reativo": {
            "forcas": "Forca: A sensibilidade as dinamicas relacionais permite antecipar conflitos e intervir antes que se agravem. Forca: Consegue perceber sutilezas emocionais nas interacoes que outros lideres ignorariam. Forca: A combinacao de empatia com consciencia relacional cria mediacao natural de conflitos.",
            "riscos": "Risco: Pode se envolver emocionalmente demais nos conflitos da equipe, perdendo a neutralidade necessaria para mediar. Risco: Tende a ser reativado por padroes relacionais do passado quando as dinamicas do grupo se intensificam. Risco: A sobrecarga emocional de absorver as tensoes e ao mesmo tempo ser ativado por elas pode causar colapso emocional.",
            "recomendacoes": "Acao: Desenvolva a capacidade de observar as dinamicas relacionais sem se fundir emocionalmente com elas. Acao: Identifique quais tipos de conflito ativam suas proprias transferencias e busque supervisao para esses casos. Acao: Estabeleca limites claros sobre ate onde vai seu papel de contenedor e quando e necessario buscar apoio externo."
        },
        "O Observador Consciente": {
            "forcas": "Forca: A autoconsciencia permite monitorar a propria capacidade emocional e evitar o esgotamento por empatia. Forca: Consegue refletir sobre as dinamicas emocionais do grupo com profundidade e clareza. Forca: A mentalizacao transforma a empatia intuitiva em compreensao estruturada das necessidades da equipe.",
            "riscos": "Risco: Pode intelectualizar as emocoes em vez de simplesmente acolhe-las, perdendo a espontaneidade do cuidado. Risco: Tende a analisar demais as interacoes emocionais, criando distancia em vez de conexao. Risco: A auto-observacao constante pode gerar duvidas sobre a autenticidade do proprio cuidado.",
            "recomendacoes": "Acao: Permita-se sentir e acolher antes de analisar, usando a reflexao como complemento e nao substituto da empatia. Acao: Use a autoconsciencia para identificar quando esta se distanciando emocionalmente e reconecte-se. Acao: Confie na sua intuicao emocional tanto quanto na sua capacidade analitica."
        },
        "O Executor Decidido": {
            "forcas": "Forca: A decisividade permite agir rapidamente quando a equipe precisa de intervenção emocional urgente. Forca: Combina cuidado com pragmatismo, oferecendo suporte que resulta em acoes concretas. Forca: A orientacao para acao impede que o acolhimento emocional se torne passividade ou estagnacao.",
            "riscos": "Risco: Pode atropelar processos emocionais do grupo por impaciencia, querendo resolver sentimentos como se fossem problemas tecnicos. Risco: Tende a oferecer solucoes rapidas quando as pessoas precisam apenas ser ouvidas. Risco: A urgencia de agir pode criar a impressao de que emocoes sao obstaculos a serem eliminados.",
            "recomendacoes": "Acao: Pratique a escuta ativa sem necessidade de resolver imediatamente, reconhecendo que processar emocoes leva tempo. Acao: Diferencie entre situacoes que exigem acao rapida e aquelas que precisam de espaco emocional. Acao: Use sua energia de execucao para criar estruturas de apoio duradouras, nao apenas intervencoes pontuais."
        }
    },
    "O Buscador de Reconhecimento": {
        "O Idealista Exigente": {
            "forcas": "Forca: O carisma natural combinado com altos padroes cria uma lideranca visionaria e inspiradora. Forca: A necessidade de reconhecimento impulsiona entregas de excelencia que destacam o lider e a equipe. Forca: Consegue mobilizar pessoas em torno de metas ambiciosas com entusiasmo contagiante.",
            "riscos": "Risco: Pode usar padroes elevados como instrumento de competicao, buscando superioridade em vez de excelencia. Risco: Tende a se frustrar intensamente quando resultados excelentes nao geram o reconhecimento esperado. Risco: A combinacao de vaidade com perfeccionismo pode tornar qualquer critica uma ferida narcisica.",
            "recomendacoes": "Acao: Separe a busca por excelencia da necessidade de aplausos, encontrando satisfacao no processo e nao apenas no resultado visivel. Acao: Desenvolva tolerancia a critica construtiva, entendendo que feedback e investimento no seu crescimento. Acao: Celebre as conquistas coletivas com a mesma intensidade que celebra as individuais."
        },
        "O Contenedor Empatico": {
            "forcas": "Forca: O carisma aliado a empatia cria uma presenca magnetica que faz as pessoas se sentirem valorizadas. Forca: Consegue usar a sensibilidade emocional para conectar-se profundamente com cada membro da equipe. Forca: A necessidade de ser admirado motiva a criacao de ambientes positivos onde todos prosperam.",
            "riscos": "Risco: Pode usar a empatia de forma manipulativa, acolhendo estrategicamente para ganhar lealdade e admiracao. Risco: Tende a absorver emocoes alheias nao por cuidado genuino, mas para se posicionar como heroi ou salvador. Risco: Pode negligenciar suas proprias necessidades emocionais enquanto busca validacao atraves do cuidado com os outros.",
            "recomendacoes": "Acao: Examine honestamente suas motivacoes ao oferecer apoio emocional, distinguindo cuidado genuino de busca por admiracao. Acao: Permita que outros tambem cuidem de voce, reduzindo a assimetria relacional que alimenta o narcisismo. Acao: Desenvolva fontes de validacao interna que nao dependam de ser o salvador emocional do grupo."
        },
        "O Estruturador Cauteloso": {
            "forcas": "Forca: A capacidade de estruturar processos da visibilidade e credibilidade a lideranca carismatica. Forca: O planejamento cuidadoso garante que as iniciativas do lider nao sejam apenas brilhantes mas tambem sustentaveis. Forca: A combinacao de carisma com organizacao cria uma lideranca que inspira e entrega resultados consistentes.",
            "riscos": "Risco: Pode usar estruturas e processos para controlar a narrativa e garantir que os creditos sempre cheguem a si. Risco: Tende a criar burocracias que centralizam as decisoes, mantendo-se como figura indispensavel. Risco: A cautela excessiva pode frustrar o desejo de brilhar rapidamente, gerando impaciencia com o proprio sistema.",
            "recomendacoes": "Acao: Crie estruturas que distribuam responsabilidades e visibilidade, em vez de centraliza-las em voce. Acao: Reconheca que lideranca sustentavel vem da construcao de sistemas que funcionam mesmo sem a sua presenca. Acao: Use sua capacidade de planejamento para criar oportunidades de destaque para toda a equipe."
        },
        "O Relacional Reativo": {
            "forcas": "Forca: A sensibilidade relacional permite perceber como esta sendo percebido e ajustar seu comportamento para manter influencia positiva. Forca: Consegue navegar dinamicas politicas complexas com carisma e inteligencia emocional. Forca: A consciencia das transferencias permite construir aliancas estrategicas genuinas.",
            "riscos": "Risco: Pode reagir de forma desproporcional quando sente que nao esta recebendo o reconhecimento merecido nas relacoes. Risco: Tende a interpretar dinamicas relacionais normais como ameacas ao seu status ou posicao. Risco: A reatividade emocional combinada com necessidade de validacao pode gerar ciclos de idealizacao e desvalorizacao dos outros.",
            "recomendacoes": "Acao: Quando se sentir ameacado em uma relacao profissional, pause e avalie se a ameaca e real ou uma projecao da sua necessidade de reconhecimento. Acao: Desenvolva relacionamentos profissionais baseados em reciprocidade, nao em admiracao unilateral. Acao: Trabalhe a capacidade de manter sua autoestima estavel independentemente das flutuacoes relacionais."
        },
        "O Observador Consciente": {
            "forcas": "Forca: A autoconsciencia permite reconhecer quando a busca por reconhecimento esta distorcendo decisoes ou relacoes. Forca: A capacidade reflexiva transforma o narcisismo bruto em autenticidade e autoconhecimento produtivo. Forca: Consegue usar o carisma de forma consciente e etica, potencializando sua influencia positiva.",
            "riscos": "Risco: Pode usar a auto-analise como mais uma forma de se admirar, transformando a reflexao em narcisismo intelectual. Risco: Tende a se perder em reflexoes sobre suas motivacoes sem mudar efetivamente o comportamento. Risco: A hiperconciencia pode gerar uma versao calculada de si mesmo que perde espontaneidade e autenticidade.",
            "recomendacoes": "Acao: Use a autoconsciencia para acao transformadora, nao como espelho sofisticado para auto-admiracao. Acao: Busque feedback externo honesto para contrastar com suas auto-percepcoes, aceitando discrepancias. Acao: Pratique acoes anonimas de contribuicao para exercitar a generosidade desvinculada do reconhecimento."
        },
        "O Executor Decidido": {
            "forcas": "Forca: A combinacao de carisma com capacidade de execucao rapida cria uma lideranca altamente impactante e visivel. Forca: A decisividade permite transformar ideias brilhantes em resultados concretos que geram reconhecimento legitimado. Forca: A energia de acao combinada com presenca carismatica mobiliza equipes para entregas extraordinarias.",
            "riscos": "Risco: Pode agir precipitadamente para gerar resultados visiveis, sacrificando qualidade ou sustentabilidade. Risco: Tende a monopolizar decisoes importantes para garantir que os creditos sejam seus. Risco: A impaciencia com processos lentos combinada com desejo de destaque pode gerar atropelos e ressentimento na equipe.",
            "recomendacoes": "Acao: Canalize sua energia de execucao para resultados coletivos, celebrando publicamente a contribuicao de cada um. Acao: Pratique a paciencia estrategica, entendendo que nem todos os resultados significativos sao imediatos ou visiveis. Acao: Delegue execucao e compartilhe creditos, fortalecendo sua reputacao como lider que desenvolve outros."
        }
    },
    "O Estruturador Cauteloso": {
        "O Idealista Exigente": {
            "forcas": "Forca: A organizacao metodica sustenta a busca por excelencia com processos claros e mensuracoes precisas. Forca: Altos padroes combinados com estrutura criam um sistema de qualidade robusto e confiavel. Forca: A previsibilidade dos processos reduz erros e aproxima os resultados dos padroes idealizados.",
            "riscos": "Risco: Pode criar uma maquina burocratica perfeccionista que sufoca a criatividade e a agilidade da equipe. Risco: Tende a confundir seguir processos a risca com atingir excelencia, priorizando forma sobre substancia. Risco: A dupla rigidez de controle e perfeccionismo pode tornar o ambiente de trabalho insuportavelmente tenso.",
            "recomendacoes": "Acao: Diferencie entre processos essenciais que garantem qualidade e processos excessivos que apenas alimentam a necessidade de controle. Acao: Inclua revisoes periodicas de processos para eliminar etapas que nao agregam valor real. Acao: Permita que a equipe sugira simplificacoes, valorizando eficiencia tanto quanto perfeicao."
        },
        "O Contenedor Empatico": {
            "forcas": "Forca: A estrutura clara combinada com sensibilidade emocional cria ambientes previsivos e acolhedores ao mesmo tempo. Forca: Os processos organizados permitem atender as necessidades emocionais da equipe de forma consistente. Forca: A previsibilidade estrutural reduz a ansiedade do grupo, complementando o acolhimento emocional.",
            "riscos": "Risco: Pode usar a estrutura como barreira emocional, mantendo as relacoes dentro de parametros controlados. Risco: Tende a processar sentimentos como itens de uma lista, perdendo a profundidade da conexao humana. Risco: A necessidade de ordem pode entrar em conflito com a natureza caotica das emocoes humanas.",
            "recomendacoes": "Acao: Reconheca que emocoes nao seguem cronogramas e permita flexibilidade nos processos quando o momento emocional exigir. Acao: Use a estrutura para criar espacos seguros de dialogo, nao para controlar o que pode ser dito. Acao: Pratique estar presente emocionalmente mesmo quando a situacao foge dos seus planos."
        },
        "O Buscador de Reconhecimento": {
            "forcas": "Forca: O carisma do buscador de reconhecimento torna os processos estruturados mais atrativos e faceis de adotar pela equipe. Forca: A necessidade de visibilidade motiva a criacao de sistemas que geram resultados mensuraveisos e celebraveis. Forca: A combinacao de organizacao com comunicação carismatica fortalece a adesao da equipe a processos.",
            "riscos": "Risco: Pode criar processos elaborados mais para impressionar do que para resolver problemas reais. Risco: Tende a resistir a mudancas nos processos que nao foram propostas por si mesmo, protegendo a autoria. Risco: A necessidade de reconhecimento pode fazer com que os processos sirvam ao ego do lider em vez de servir a equipe.",
            "recomendacoes": "Acao: Avalie seus processos pela eficacia real e nao pela impressao que causam em superiores ou pares. Acao: Encoraje a equipe a melhorar processos e reconheca publicamente essas contribuicoes. Acao: Foque em criar sistemas que funcionem independentemente de quem receba o credito."
        },
        "O Relacional Reativo": {
            "forcas": "Forca: A sensibilidade relacional permite perceber quando processos estao gerando resistencia ou desconforto na equipe. Forca: Consegue ajustar estruturas de acordo com as dinamicas interpessoais do grupo. Forca: A consciencia das transferencias ajuda a entender por que certas pessoas resistem a determinadas estruturas.",
            "riscos": "Risco: Pode usar processos e estruturas como defesa contra a ansiedade gerada pelas dinamicas relacionais. Risco: Tende a reagir emocionalmente quando alguem questiona ou desrespeita os processos estabelecidos. Risco: A reatividade relacional pode criar excecoes inconsistentes nos processos baseadas em preferencias pessoais.",
            "recomendacoes": "Acao: Observe quando a rigidez nos processos aumenta em resposta a conflitos relacionais e busque abordar a causa real. Acao: Aplique as mesmas regras para todos, independentemente da sua relacao pessoal com cada membro da equipe. Acao: Use os processos como ferramentas de gestao, nao como escudo emocional contra a complexidade das relacoes."
        },
        "O Observador Consciente": {
            "forcas": "Forca: A autoconsciencia permite identificar quando a necessidade de controle esta se tornando disfuncional. Forca: A reflexao sobre os proprios padroes de organizacao gera melhorias continuas nos processos. Forca: A mentalizacao ajuda a entender o impacto dos processos na experiencia subjetiva da equipe.",
            "riscos": "Risco: Pode analisar excessivamente cada processo sem nunca implementa-lo de fato, buscando a estrutura perfeita. Risco: Tende a usar a reflexao como mais uma forma de controle, planejando compulsivamente em vez de executar. Risco: A hiperanalise pode gerar paralisia decisoria sobre qual a melhor forma de organizar cada atividade.",
            "recomendacoes": "Acao: Estabeleca prazos para a fase de planejamento e force-se a implementar, mesmo que o processo nao esteja perfeito. Acao: Adote uma mentalidade iterativa onde processos sao testados e melhorados continuamente, nao planejados indefinidamente. Acao: Use a autoconsciencia para soltar o controle gradualmente, confiando na capacidade adaptativa da equipe."
        },
        "O Executor Decidido": {
            "forcas": "Forca: A decisividade do executor complementa o planejamento do estruturador, criando um ciclo eficiente de planejar e agir. Forca: A urgencia de execucao evita que os processos se tornem fins em si mesmos. Forca: A combinacao de estrutura com acao rapida gera resultados previsiveis em prazos competitivos.",
            "riscos": "Risco: A tensao entre querer planejar tudo e querer agir rapido pode gerar ansiedade e inconsistencia. Risco: Tende a atropelar seus proprios processos quando a urgencia fala mais alto, minando a credibilidade da estrutura. Risco: Pode oscilar entre periodos de planejamento excessivo e execucao precipitada, confundindo a equipe.",
            "recomendacoes": "Acao: Crie processos ageis que permitam velocidade sem sacrificar a previsibilidade que voce valoriza. Acao: Defina claramente em quais situacoes e aceitavel simplificar processos para ganhar velocidade. Acao: Pratique a consistencia entre o que planeja e o que executa, fortalecendo a confianca da equipe nos seus sistemas."
        }
    },
    "O Relacional Reativo": {
        "O Idealista Exigente": {
            "forcas": "Forca: A sensibilidade relacional permite calibrar exigencias de acordo com a capacidade e o momento de cada pessoa. Forca: A consciencia das transferencias ajuda a entender de onde vem a propria pressao por perfeicao. Forca: Consegue criar relacoes de confianca que sustentam feedbacks exigentes sem romper vinculos.",
            "riscos": "Risco: Pode alternar entre exigencia excessiva e permissividade dependendo do estado emocional da relacao com cada liderado. Risco: Tende a projetar seus padroes de perfeicao nas relacoes, esperando que os outros correspondam a idealizacoes. Risco: A reatividade emocional pode transformar avaliacoes de desempenho em julgamentos pessoais carregados de afeto.",
            "recomendacoes": "Acao: Separe a avaliacao de desempenho da qualidade do vinculo relacional, mantendo criterios objetivos. Acao: Identifique quais liderados ativam suas transferencias mais intensas e busque equilibrio nessas relacoes. Acao: Desenvolva consciencia de quando esta sendo exigente por padrao de qualidade e quando por necessidade relacional."
        },
        "O Contenedor Empatico": {
            "forcas": "Forca: A combinacao de reatividade relacional com empatia cria uma lideranca extremamente sintonizada com o clima emocional do grupo. Forca: Consegue perceber e acolher tensoes interpessoais antes que se transformem em conflitos abertos. Forca: A sensibilidade dupla permite mediar conflitos com profundidade e genuinidade.",
            "riscos": "Risco: Pode se fundir emocionalmente com os conflitos do grupo, perdendo a capacidade de funcionar como lider. Risco: Tende a absorver as projecoes e transferencias dos outros sem filtro, esgotando-se rapidamente. Risco: A combinacao de reatividade com empatia pode criar dependencia emocional mutua entre lider e equipe.",
            "recomendacoes": "Acao: Desenvolva a capacidade de estar emocionalmente presente sem se perder nas emocoes alheias, mantendo seu centro. Acao: Busque supervisao ou terapia regular para processar as cargas emocionais que absorve do grupo. Acao: Estabeleca rituais de autocuidado que renovem sua energia emocional apos interacoes intensas."
        },
        "O Buscador de Reconhecimento": {
            "forcas": "Forca: O carisma e a consciencia relacional permitem construir redes de influencia poderosas e autenticas. Forca: A sensibilidade as dinamicas interpessoais potencializa a capacidade de engajar e motivar pessoas. Forca: Consegue perceber e responder as necessidades nao-verbalizadas dos outros, gerando lealdade genuina.",
            "riscos": "Risco: Pode manipular dinamicas relacionais para garantir que o reconhecimento flua em sua direcao. Risco: Tende a interpretar a falta de reconhecimento como rejeicao pessoal, ativando padroes relacionais defensivos. Risco: A combinacao de necessidade de validacao com reatividade pode gerar dramas relacionais recorrentes.",
            "recomendacoes": "Acao: Observe quando esta usando suas habilidades relacionais para servir a equipe versus para servir a sua necessidade de validacao. Acao: Pratique receber falta de reconhecimento sem interpretar como ataque pessoal ou relacional. Acao: Invista em relacoes profissionais baseadas em contribuicao mutua, nao em admiracao."
        },
        "O Estruturador Cauteloso": {
            "forcas": "Forca: A estrutura e os processos claros criam um contencao que reduz a ansiedade gerada pela reatividade relacional. Forca: As regras e protocolos ajudam a manter relacoes profissionais dentro de limites saudaveis. Forca: A previsibilidade dos processos oferece seguranca quando as dinamicas interpessoais se tornam intensas.",
            "riscos": "Risco: Pode usar estruturas rigidas como defesa contra a vulnerabilidade das relacoes interpessoais. Risco: Tende a burocratizar as interacoes humanas para evitar o desconforto da espontaneidade relacional. Risco: A rigidez estrutural pode impedir a flexibilidade necessaria para lidar com situacoes emocionais imprevistas.",
            "recomendacoes": "Acao: Use processos como suporte para as relacoes, nao como substituto delas, mantendo espaco para o humano. Acao: Permita-se ser surpreendido nas interacoes, aceitando que nem tudo precisa ser planejado. Acao: Avalie periodicamente se seus processos estao facilitando ou dificultando as conexoes genuinas com a equipe."
        },
        "O Observador Consciente": {
            "forcas": "Forca: A autoconsciencia permite identificar padroes relacionais repetitivos e trabalhar para transforma-los. Forca: A mentalizacao ajuda a entender tanto suas proprias reacoes quanto as dos outros nas interacoes. Forca: A reflexao consciente transforma experiencias relacionais dificeis em aprendizado e crescimento.",
            "riscos": "Risco: Pode intelectualizar as reacoes emocionais em vez de vive-las e resolve-las genuinamente. Risco: Tende a usar a auto-analise como forma de evitar o confronto direto com dinamicas relacionais dificeis. Risco: A hiperreflexao sobre padroes relacionais pode gerar paralisia na hora de agir nas relacoes.",
            "recomendacoes": "Acao: Equilibre reflexao com acao relacional, usando os insights para mudar comportamentos concretos. Acao: Pratique conversas dificeis em vez de apenas analisa-las mentalmente, enfrentando o desconforto. Acao: Compartilhe suas reflexoes com pessoas de confianca para validar percepcoes e evitar vieses."
        },
        "O Executor Decidido": {
            "forcas": "Forca: A decisividade permite resolver conflitos interpessoais rapidamente em vez de deixa-los escalar. Forca: A orientacao para acao impede que dinamicas relacionais toxicas se perpetuem por inacao. Forca: Consegue tomar decisoes dificeis sobre pessoas quando as dinamicas do grupo exigem mudanca.",
            "riscos": "Risco: Pode agir impulsivamente em situacoes relacionais complexas que exigiriam mais reflexao e sensibilidade. Risco: Tende a tratar conflitos interpessoais como problemas a serem resolvidos com rapidez, ignorando a profundidade emocional. Risco: A impaciencia com processos relacionais lentos pode gerar rupturas desnecessarias nos vinculos.",
            "recomendacoes": "Acao: Antes de agir em situacoes interpessoais intensas, faca uma pausa para avaliar se a urgencia e real ou e sua reatividade. Acao: Reconheca que relacoes humanas precisam de tempo para se reparar e desenvolver, mesmo quando voce quer resolucao imediata. Acao: Use sua capacidade de decisao para criar espacos de dialogo, nao para encerrar conversas prematuramente."
        }
    },
    "O Observador Consciente": {
        "O Idealista Exigente": {
            "forcas": "Forca: A reflexao consciente direciona a busca por excelencia para areas de real impacto e significado. Forca: A autoconsciencia permite distinguir entre padroes saudaveis de qualidade e perfeccionismo autodestrutivo. Forca: Consegue articular com clareza por que determinados padroes sao importantes, inspirando compreensao na equipe.",
            "riscos": "Risco: Pode analisar excessivamente seus proprios padroes de excelencia sem conseguir relaxa-los na pratica. Risco: Tende a se tornar um critico interno hiperativo que pensa muito e age pouco sobre suas exigencias. Risco: A combinacao de hiperreflexao com perfeccionismo pode gerar uma forma sofisticada de procrastinacao.",
            "recomendacoes": "Acao: Transforme suas reflexoes sobre padroes em acoes concretas com prazos definidos. Acao: Pratique o conceito de bom o suficiente em areas de menor impacto, reservando a excelencia para o que realmente importa. Acao: Busque um accountability partner que o ajude a sair do ciclo de analise e entrar no ciclo de execucao."
        },
        "O Contenedor Empatico": {
            "forcas": "Forca: A autoconsciencia enriquece a empatia, criando um acolhimento profundo e informado por compreensao psicologica. Forca: Consegue oferecer suporte emocional sem se perder nas emocoes alheias, mantendo clareza interna. Forca: A reflexao consciente permite perceber necessidades emocionais sutis que outros lideres nao captariam.",
            "riscos": "Risco: Pode transformar o cuidado em projeto intelectual, analisando as emocoes dos outros em vez de simplesmente acolhe-las. Risco: Tende a manter distancia emocional sob pretexto de observacao consciente, evitando vulnerabilidade. Risco: A sobre-analise das dinamicas emocionais pode tornar as interacoes artificiais e calculadas.",
            "recomendacoes": "Acao: Pratique momentos de presenca emocional pura, sem a necessidade de compreender ou analisar tudo. Acao: Permita-se ser tocado emocionalmente pelas situacoes da equipe, usando a reflexao depois e nao durante. Acao: Busque feedback sobre como sua equipe percebe seu acolhimento para verificar se esta chegando de forma genuina."
        },
        "O Buscador de Reconhecimento": {
            "forcas": "Forca: A autoconsciencia sobre a necessidade de reconhecimento permite usa-la de forma construtiva e etica. Forca: O carisma informado por reflexao cria uma lideranca autentica e nao manipulativa. Forca: Consegue canalizar o desejo de destaque para contribuicoes que beneficiam genuinamente o grupo.",
            "riscos": "Risco: Pode usar a autoconsciencia como mais uma ferramenta de construcao de imagem sofisticada. Risco: Tende a ficar preso entre o desejo de reconhecimento e a autocritica por desejar reconhecimento. Risco: A reflexao excessiva sobre motivacoes pode levar a inacao por medo de ser percebido como narcisista.",
            "recomendacoes": "Acao: Aceite que o desejo de reconhecimento e humano e natural, focando em canaliza-lo de forma construtiva. Acao: Use sua capacidade reflexiva para garantir que suas acoes estejam alinhadas com seus valores, nao apenas com seu ego. Acao: Pratique a transparencia sobre suas motivacoes com pessoas de confianca, reduzindo a pressao interna."
        },
        "O Estruturador Cauteloso": {
            "forcas": "Forca: A reflexao consciente sobre processos gera sistemas mais inteligentes e adaptaveis. Forca: A autoconsciencia permite perceber quando a necessidade de controle esta limitando a eficacia da equipe. Forca: A mentalizacao ajuda a criar estruturas que consideram o fator humano alem do fator tecnico.",
            "riscos": "Risco: Pode analisar excessivamente cada processo, nunca chegando a uma versao final implementavel. Risco: Tende a criar sistemas teoricamente perfeitos mas praticamente impraticaveis por excesso de complexidade. Risco: A combinacao de cautela com hiperreflexao pode gerar paralisia organizacional sofisticada.",
            "recomendacoes": "Acao: Adote a regra dos 80 porcento: implemente quando o processo estiver bom o suficiente e melhore continuamente. Acao: Teste suas estruturas rapidamente no campo em vez de refina-las indefinidamente no papel. Acao: Peca feedback pratico da equipe sobre seus processos, priorizando usabilidade sobre perfeicao teorica."
        },
        "O Relacional Reativo": {
            "forcas": "Forca: A autoconsciencia permite identificar e nomear padroes transferenciais no momento em que acontecem. Forca: A mentalizacao transforma reacoes emocionais automaticas em respostas conscientes e escolhidas. Forca: Consegue observar as dinamicas relacionais do grupo com clareza suficiente para intervir de forma precisa.",
            "riscos": "Risco: Pode observar e analisar as dinamicas sem nunca intervir, mantendo-se na seguranca da posicao de espectador. Risco: Tende a intelectualizar conflitos relacionais, nomeando padroes sem resolve-los efetivamente. Risco: A sobre-analise das proprias reacoes pode gerar desconfianca de si mesmo nas interacoes espontaneas.",
            "recomendacoes": "Acao: Use seus insights sobre dinamicas relacionais para intervencoes praticas e oportunas, nao apenas para compreensao teorica. Acao: Confie nas suas reacoes emocionais como dados validos, nao apenas como objetos de analise. Acao: Pratique intervencoes relacionais imperfeitas, aceitando que a acao conta mais que a analise perfeita."
        },
        "O Executor Decidido": {
            "forcas": "Forca: A decisividade complementa a reflexao, garantindo que insights profundos se transformem em acoes concretas. Forca: A orientacao para acao evita que a auto-observacao se torne fim em si mesma. Forca: A combinacao de pensamento profundo com execucao rapida cria decisoes bem fundamentadas e implementadas com agilidade.",
            "riscos": "Risco: A tensao entre querer refletir mais e querer agir logo pode gerar estresse e decisoes ora precipitadas ora atrasadas. Risco: Tende a alternar entre periodos de analise excessiva e acao impulsiva, sem encontrar o equilibrio. Risco: A impaciencia do executor pode sabotar a profundidade da reflexao que o observador precisa.",
            "recomendacoes": "Acao: Defina momentos claros para reflexao e momentos para acao, respeitando ambos sem mistura-los. Acao: Use a reflexao para definir direcao e a execucao para gerar aprendizado pratico, criando um ciclo virtuoso. Acao: Aceite que nem toda decisao precisa de reflexao profunda e nem toda acao precisa ser imediata."
        }
    },
    "O Executor Decidido": {
        "O Idealista Exigente": {
            "forcas": "Forca: A velocidade de execucao aliada a altos padroes gera entregas rapidas e de qualidade quando bem equilibradas. Forca: A decisividade evita a paralisia da perfeicao, transformando ideais em resultados concretos. Forca: A combinacao de urgencia com excelencia cria uma lideranca de alto impacto e alta performance.",
            "riscos": "Risco: Pode entregar rapido mas se frustrar com a qualidade, gerando ciclos de retrabalho e insatisfacao. Risco: Tende a pressionar a equipe com prazos agressivos e padroes irrealistas simultaneamente. Risco: A tensao cronica entre velocidade e perfeicao pode causar burnout pessoal e coletivo.",
            "recomendacoes": "Acao: Defina explicitamente para cada projeto se a prioridade e velocidade ou qualidade, comunicando isso a equipe. Acao: Crie mecanismos de revisao rapida que garantam qualidade minima sem sacrificar a agilidade. Acao: Monitore seu nivel de estresse e o da equipe, ajustando expectativas quando necessario."
        },
        "O Contenedor Empatico": {
            "forcas": "Forca: A empatia suaviza a urgencia de acao, criando um ambiente onde a velocidade nao atropela as pessoas. Forca: Consegue mover a equipe rapidamente enquanto mantem a conexao emocional e o bem-estar do grupo. Forca: A sensibilidade emocional informa decisoes rapidas, tornando-as mais humanas e sustentaveis.",
            "riscos": "Risco: Pode ignorar sinais emocionais da equipe por estar focado na urgencia da execucao. Risco: Tende a alternar entre momentos de pressao intensa e tentativas de reparacao emocional, confundindo a equipe. Risco: A impaciencia natural pode ser percebida como falta de cuidado, mesmo quando a intencao e boa.",
            "recomendacoes": "Acao: Inclua check-ins emocionais rapidos na sua rotina de execucao, mesmo quando o ritmo e intenso. Acao: Comunique a equipe que sua urgencia nao significa falta de cuidado, explicitando suas intencoes. Acao: Aprenda a desacelerar quando perceber que a equipe precisa de processamento emocional antes de seguir adiante."
        },
        "O Buscador de Reconhecimento": {
            "forcas": "Forca: O carisma combinado com capacidade de entrega cria uma lideranca que inspira e produz resultados visiveis. Forca: A necessidade de reconhecimento canalizada em execucao gera realizacoes concretas e celebraveis. Forca: A visibilidade natural do executor potencializa o impacto positivo do seu carisma.",
            "riscos": "Risco: Pode priorizar acoes que geram visibilidade rapida em detrimento de trabalho importante mas invisivel. Risco: Tende a competir por projetos de alta visibilidade, negligenciando tarefas fundamentais mas pouco glamorosas. Risco: A combinacao de impaciencia com vaidade pode gerar decisoes precipitadas motivadas por ego.",
            "recomendacoes": "Acao: Valorize igualmente entregas visiveis e invisiveis, reconhecendo que sustentabilidade exige trabalho nos bastidores. Acao: Desenvolva satisfacao com o impacto real do trabalho, nao apenas com a percepcao externa. Acao: Pratique delegar projetos de alta visibilidade para membros da equipe, demonstrando lideranca madura."
        },
        "O Estruturador Cauteloso": {
            "forcas": "Forca: A estrutura e o planejamento direcionam a energia de execucao para acoes mais estrategicas e eficazes. Forca: Os processos organizados garantem que a velocidade de execucao nao gere caos ou inconsistencia. Forca: A cautela do estruturador equilibra a impulsividade do executor, criando decisoes rapidas mas fundamentadas.",
            "riscos": "Risco: Pode se sentir preso pelos proprios processos quando a situacao exige acao imediata e sem precedentes. Risco: Tende a oscilar entre seguir o plano rigidamente e abandona-lo completamente quando fica impaciente. Risco: A frustracao com a lentidao dos processos pode leva-lo a desrespeitar suas proprias estruturas.",
            "recomendacoes": "Acao: Crie processos que tenham caminhos acelerados para situacoes de urgencia, mantendo a estrutura sem perder agilidade. Acao: Respeite seus proprios processos para manter a credibilidade, pedindo contribuicoes da equipe para torna-los mais ageis. Acao: Reconheca que a estrutura existe para servir a execucao, ajustando-a quando necessario sem abandona-la."
        },
        "O Relacional Reativo": {
            "forcas": "Forca: A sensibilidade relacional permite perceber rapidamente quando a equipe precisa de abordagens diferentes para manter o ritmo. Forca: Consegue adaptar seu estilo de lideranca de acordo com as necessidades interpessoais de cada momento. Forca: A consciencia das dinamicas de grupo informa decisoes rapidas sobre alocacao e gestao de pessoas.",
            "riscos": "Risco: Pode tomar decisoes impulsivas sobre pessoas quando ativado emocionalmente por dinamicas interpessoais. Risco: Tende a reagir rapidamente demais em conflitos, antes de entender completamente o que esta acontecendo. Risco: A impaciencia com processos relacionais pode gerar rupturas desnecessarias que prejudicam o desempenho do grupo.",
            "recomendacoes": "Acao: Estabeleca a regra de esperar 24 horas antes de tomar decisoes sobre pessoas ou relacoes em momentos de ativacao emocional. Acao: Use sua capacidade de acao rapida para criar espacos de dialogo, nao para impor solucoes unilaterais. Acao: Desenvolva tolerancia ao desconforto relacional, entendendo que nem todo conflito exige acao imediata."
        },
        "O Observador Consciente": {
            "forcas": "Forca: A autoconsciencia permite identificar quando a urgencia de acao esta sendo produtiva ou destrutiva. Forca: A reflexao consciente informa decisoes rapidas com maior profundidade e menor risco de erro. Forca: Consegue equilibrar acao decisiva com momentos de pausa estrategica quando necessario.",
            "riscos": "Risco: A tensao entre agir e refletir pode gerar indecisao ou acoes seguidas de arrependimento. Risco: Tende a alternar entre impulso e sobre-analise, sem encontrar um ritmo consistente. Risco: Pode usar a reflexao como desculpa para procrastinar em situacoes que genuinamente exigem acao imediata.",
            "recomendacoes": "Acao: Desenvolva um sistema rapido de checagem interna antes de decisoes importantes: tres perguntas-chave em dois minutos. Acao: Aceite que algumas decisoes precisam ser tomadas com informacao incompleta e que a acao gera aprendizado. Acao: Use a reflexao pos-acao como ferramenta de melhoria continua em vez de tentar prever tudo antes de agir."
        }
    }
}


MODULES_DATA = [
    {
        "id": 0, 
        "name": "Introducao: A Jornada LPS", 
        "title": "A Jornada LPS", 
        "description": "Apresentação da metodologia, boas-vindas da Viviane Nishiura e o mapa da jornada entre Neurociência e Psicanálise. Conheça o caminho que vai transformar sua liderança.", 
        "icon": "🚀", 
        "file": "attached_assets/introdução_1768431876966.pdf", 
        "videos": [
            "https://vimeo.com/1154502544",
            "https://vimeo.com/1154502598",
            "https://vimeo.com/1154502492"
        ]
    },
    {
        "id": 1, 
        "name": "Modulo 1: Neurociência da Liderança", 
        "title": "Neurociência da Liderança", 
        "description": "Foco em químicos cerebrais (dopamina, ocitocina) e o Círculo de Segurança. Entenda como o cérebro processa decisões e aprenda a usar a neurociência para liderar com mais eficácia.", 
        "icon": "🧠", 
        "file": "attached_assets/Módulo_1_1768431876967.pdf", 
        "videos": [
            "https://vimeo.com/1154503073",
            "https://vimeo.com/1154503122",
            "https://vimeo.com/1154503201",
            "https://vimeo.com/1154503286",
            "https://vimeo.com/1154503332",
            "https://vimeo.com/1154502907",
            "https://vimeo.com/1154502997"
        ]
    },
    {
        "id": 2, 
        "name": "Modulo 2: Mergulho no Inconsciente", 
        "title": "Mergulho no Inconsciente", 
        "description": "Estudo do Id, Ego, Superego e mecanismos de defesa. Explore as camadas profundas da mente e descubra como padrões inconscientes influenciam sua liderança.", 
        "icon": "🌊", 
        "file": "attached_assets/Módulo_2_1768431876968.pdf", 
        "videos": [
            "https://vimeo.com/1154504282",
            "https://vimeo.com/1154503918",
            "https://vimeo.com/1154503996",
            "https://vimeo.com/1154504129",
            "https://vimeo.com/1154504216",
            "https://vimeo.com/1154504054"
        ]
    },
    {
        "id": 3, 
        "name": "Modulo 3: Relacoes e Transferencia", 
        "title": "Relacoes e Transferencia", 
        "description": "Dinâmicas líder-liderado e manejo de contratransferência. Compreenda as dinâmicas de transferência nas relações profissionais e como usá-las a seu favor.", 
        "icon": "🔄", 
        "file": "attached_assets/Módulo_3_1768431876969.pdf", 
        "videos": [
            "https://vimeo.com/1154508629",
            "https://vimeo.com/1154508577",
            "https://vimeo.com/1154508688",
            "https://vimeo.com/1154508745",
            "https://vimeo.com/1154508530"
        ]
    },
    {
        "id": 4, 
        "name": "Modulo 4: Autoconsciencia", 
        "title": "Autoconsciencia", 
        "description": "Identificação do seu arquétipo de liderança psicanalítica. Desenvolva autoconhecimento profundo e identifique seus gatilhos emocionais como líder.", 
        "icon": "🪞", 
        "file": "attached_assets/Módulo_4_1768431876970.pdf", 
        "videos": [
            "https://vimeo.com/1154509566",
            "https://vimeo.com/1154509679",
            "https://vimeo.com/1154509769",
            "https://vimeo.com/1154509511"
        ]
    },
    {
        "id": 5, 
        "name": "Modulo 5: Entendendo a Equipe", 
        "title": "Entendendo a Equipe", 
        "description": "Mapeamento com assessments e os papéis grupais de Bion. Aprenda a mapear perfis e dinâmicas grupais usando conceitos psicanalíticos.", 
        "icon": "👥", 
        "file": "attached_assets/Módulo_5_1768431876971.pdf", 
        "videos": [
            "https://vimeo.com/1154510241",
            "https://vimeo.com/1154510404",
            "https://vimeo.com/1154510309"
        ]
    },
    {
        "id": 6, 
        "name": "Modulo 6: Aplicacao Pratica", 
        "title": "Aplicacao Pratica", 
        "description": "Análise de casos reais e construção do plano de ação personalizado. Coloque em prática as ferramentas psicanalíticas no dia a dia da liderança.", 
        "icon": "🛠️", 
        "file": "attached_assets/Módulo_6_1768431876972.pdf", 
        "videos": [
            "https://vimeo.com/1154510682",
            "https://vimeo.com/1154510710",
            "https://vimeo.com/1154510729",
            "https://vimeo.com/1154510816"
        ]
    },
    {
        "id": 7, 
        "name": "Modulo 7: Liderança de Alta Performance", 
        "title": "Liderança de Alta Performance e Sustentabilidade", 
        "description": "Como manter os resultados a longo prazo, gestão da cultura organizacional e o fechamento do ciclo de desenvolvimento. Consolide sua transformação como líder.", 
        "icon": "🏆", 
        "file": "attached_assets/Módulo_7_1768431876973.pdf", 
        "videos": [
            "https://vimeo.com/1154511020",
            "https://vimeo.com/1154511064"
        ]
    }
]

# Check for employee token in URL (takes priority over auth)
query_params = st.query_params
is_employee_access = False

# Invite link detection: ?tipo=equipe&ref=token123
if 'ref' in query_params and 'tipo' in query_params:
    invite_ref = query_params['ref']
    invite_tipo = query_params['tipo']
    st.session_state.invite_ref = invite_ref
    st.session_state.invite_tipo = invite_tipo
    st.session_state.page = "InviteWelcome"

# Token in URL - set session state (legacy employee access)
if 'token' in query_params:
    st.session_state.employee_token = query_params['token']
    st.session_state.page = "EmployeeAssessment"
    is_employee_access = True

# Invite session state persistence
if st.session_state.get('invite_ref') and not st.session_state.get('authenticated') and not st.session_state.get('invite_email_verified'):
    if st.session_state.get('page') != "InviteWelcome":
        st.session_state.page = "InviteWelcome"

# Token already in session state - maintain employee access lock
if st.session_state.get('employee_token') and not st.session_state.get('authenticated'):
    st.session_state.page = "EmployeeAssessment"
    is_employee_access = True

# Default: if not authenticated and no token/invite, show Home (vitrine)
if not st.session_state.get('authenticated') and not is_employee_access and not st.session_state.get('invite_ref'):
    if st.session_state.page not in ("Home", "Login", "InviteWelcome", "EmployeeAssessment"):
        st.session_state.page = "Home"

# Navigation Menu Sections with Icons
MENU_SECTIONS = [
    {"key": "home", "label": "Home"},
    {"key": "sobre", "label": "Sobre"},
    {"key": "curso", "label": "Curso"},
    {"key": "lpstest", "label": "LPTest"},
    {"key": "lpschat", "label": "LPChat"},
    {"key": "mentoria", "label": "Mentoria"},
    {"key": "soluções", "label": "Trilha LPS"},
    {"key": "insights", "label": "Insights"},
    {"key": "contato", "label": "Contato"}
]

# Sidebar Navigation - Uses unique keys per page context
def render_sidebar_navigation():
    # Use page context for unique widget keys
    page_ctx = st.session_state.get('page', 'Home')
    key_prefix = f"sb_{page_ctx}_"
    
    with st.sidebar:
        # Sidebar CSS styling - Premium Design with Native Toggle Styled
        st.markdown("""
            <style>
            /* Global page background */
            .stApp {
                background-color: #F5F5F5 !important;
            }
            .stApp > header {
                background-color: #F5F5F5 !important;
            }
            /* Main content text color */
            .stApp .stMarkdown p, .stApp .stMarkdown li, .stApp .stMarkdown span {
                color: #000000;
            }
            /* Style the native Streamlit sidebar toggle button */
            [data-testid="collapsedControl"],
            [data-testid="stBaseButton-headerNoPadding"],
            button[data-testid="collapsedControl"] {
                background-color: #d19f09 !important;
                border-radius: 8px !important;
                border: 2px solid #d19f09 !important;
                box-shadow: 0 2px 8px rgba(0,0,0,0.25) !important;
                display: flex !important;
                align-items: center !important;
                justify-content: center !important;
                min-width: 36px !important;
                min-height: 36px !important;
            }
            [data-testid="collapsedControl"] svg,
            [data-testid="stBaseButton-headerNoPadding"] svg {
                fill: #FFFFFF !important;
                stroke: #FFFFFF !important;
                color: #FFFFFF !important;
            }
            [data-testid="collapsedControl"]:hover,
            [data-testid="stBaseButton-headerNoPadding"]:hover {
                background-color: #c5a059 !important;
                box-shadow: 0 4px 12px rgba(0,0,0,0.3) !important;
            }
            /* Sidebar nav icon spans — hide only within sidebar nav links */
            [data-testid="stSidebarNavLink"] [data-testid="stIconMaterial"],
            [data-testid="stSidebarNav"] [data-testid="stIconMaterial"] {
                font-size: 0px !important;
                color: transparent !important;
            }
            /* Sidebar toggle / hamburger always overrides any generic header button rule */
            [data-testid="collapsedControl"],
            [data-testid="stExpandSidebarButton"] {
                background-color: #d19f09 !important;
                border: 2px solid #d19f09 !important;
                border-radius: 8px !important;
                color: white !important;
            }
            /* Sidebar styling */
            [data-testid="stSidebar"] {
                background-color: #18738c !important;
                padding-top: 0;
            }
            [data-testid="stSidebar"] > div:first-child {
                background-color: #18738c !important;
            }
            /* Sidebar close button (X) - Yellow styling */
            [data-testid="stSidebar"] [data-testid="stSidebarCollapseButton"] button,
            [data-testid="stSidebar"] button[aria-label="Close sidebar"] {
                background-color: transparent !important;
                border: 2px solid #d19f09 !important;
                color: #d19f09 !important;
                border-radius: 8px !important;
            }
            [data-testid="stSidebar"] [data-testid="stSidebarCollapseButton"] button:hover,
            [data-testid="stSidebar"] button[aria-label="Close sidebar"]:hover {
                background-color: #d19f09 !important;
                color: #18738c !important;
            }
            [data-testid="stSidebar"] [data-testid="stSidebarCollapseButton"] svg,
            [data-testid="stSidebar"] button[aria-label="Close sidebar"] svg {
                fill: #d19f09 !important;
                stroke: #d19f09 !important;
            }
            [data-testid="stSidebar"] [data-testid="stSidebarCollapseButton"]:hover svg,
            [data-testid="stSidebar"] button[aria-label="Close sidebar"]:hover svg {
                fill: #18738c !important;
                stroke: #18738c !important;
            }
            /* Sidebar buttons */
            [data-testid="stSidebar"] .stButton > button {
                background-color: transparent;
                color: #FFFFFF !important;
                border: none;
                text-align: left;
                padding: 1rem 1.25rem;
                font-size: 1.05rem;
                font-weight: 500;
                width: 100%;
                border-radius: 8px;
                margin-bottom: 0.5rem;
                transition: all 0.2s ease;
                white-space: nowrap;
            }
            [data-testid="stSidebar"] .stButton > button:hover {
                background-color: rgba(209, 159, 9, 0.25) !important;
                color: #d19f09 !important;
            }
            [data-testid="stSidebar"] .stButton > button[kind="primary"] {
                background-color: #d19f09 !important;
                color: #18738c !important;
                font-weight: bold;
            }
            [data-testid="stSidebar"] .stButton > button[kind="primary"]:hover {
                background-color: #e6c654 !important;
            }
            /* Logo container - Full width, no borders */
            .sidebar-logo-container {
                text-align: center;
                padding: 0;
                margin: 0;
                background-color: transparent;
            }
            .sidebar-logo-container img {
                width: 100% !important;
                max-width: 100% !important;
            }
            /* Remove top padding from sidebar */
            [data-testid="stSidebar"] > div:first-child > div:first-child {
                padding-top: 0 !important;
                margin-top: 0 !important;
            }
            /* Sidebar divider */
            .sidebar-divider {
                border-top: 1px solid rgba(255,255,255,0.15);
                margin: 1rem 0;
            }
            /* User badge */
            .user-badge {
                background-color: rgba(209, 159, 9, 0.15);
                border: 2px solid #d19f09;
                color: #d19f09;
                padding: 0.75rem 1rem;
                border-radius: 25px;
                text-align: center;
                font-weight: bold;
                margin-bottom: 1rem;
                font-size: 0.95rem;
            }
            /* Navigation label */
            .nav-section-label {
                color: #d19f09;
                font-size: 0.75rem;
                font-weight: 600;
                text-transform: uppercase;
                letter-spacing: 1px;
                padding-left: 1.25rem;
                margin-bottom: 0.75rem;
            }
            /* Ensure banner doesn't overlap toggle button */
            .main .block-container {
                padding-top: 3rem !important;
            }
            </style>
        """, unsafe_allow_html=True)
        
        # Logo Only (Full Width, No Text Title)
        st.markdown('<div class="sidebar-logo-container">', unsafe_allow_html=True)
        if os.path.exists(LOGO_PATH):
            st.image(LOGO_PATH, use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Login/User Section at Top
        if st.session_state.authenticated:
            st.markdown(f"""
                <div class="user-badge">{st.session_state.user['name']}</div>
            """, unsafe_allow_html=True)
            if st.button("Sair", key=f"{key_prefix}logout", use_container_width=True):
                st.session_state.authenticated = False
                st.session_state.user = None
                st.session_state.page = "Home"
                st.rerun()
        else:
            if st.button("Entrar", key=f"{key_prefix}login", use_container_width=True, type="primary"):
                st.session_state.page = "Login"
                st.rerun()
        
        st.markdown('<div class="sidebar-divider"></div>', unsafe_allow_html=True)
        
        if st.session_state.authenticated:
            st.markdown('<div class="nav-section-label">Minha Área</div>', unsafe_allow_html=True)
            auth_menu = [
                {"key": "Curso", "label": "Curso LPS", "page": "LPS Curso"},
                {"key": "Dashboard", "label": "Dashboard", "page": "Dashboard"},
                {"key": "LPTest", "label": "LPTest", "page": "LPTest"},
                {"key": "LPChat", "label": "LPChat", "page": "LPChat"},
                {"key": "Equipe", "label": "Gestão de Equipe", "page": "TeamManagement"},
            ]
            for item in auth_menu:
                is_active = (st.session_state.page == item["page"])
                if st.button(item['label'], key=f"{key_prefix}auth_{item['key']}", use_container_width=True, type="primary" if is_active else "secondary"):
                    st.session_state.page = item["page"]
                    st.rerun()
            
            st.markdown('<div class="sidebar-divider"></div>', unsafe_allow_html=True)
            st.markdown("""<style>
            [data-testid="stSidebar"] button[kind="secondary"]:last-of-type {
                background-color: #c5a059 !important;
                color: white !important;
                border: none !important;
                font-weight: bold !important;
            }
            </style>""", unsafe_allow_html=True)
            if st.button("Voltar ao Dashboard", key=f"{key_prefix}back_home", use_container_width=True, type="secondary"):
                st.session_state.page = "Dashboard"
                st.rerun()
        else:
            current = st.session_state.section
            for item in MENU_SECTIONS:
                is_active = (current == item["key"])
                if st.button(item['label'], key=f"{key_prefix}nav_{item['key']}", use_container_width=True, type="primary" if is_active else "secondary"):
                    st.session_state.section = item["key"]
                    st.session_state.show_login_modal = False
                    st.rerun()
        
        # Footer in sidebar
        st.markdown('<div class="sidebar-divider"></div>', unsafe_allow_html=True)
        st.markdown("""
            <p style="color: rgba(255,255,255,0.6); font-size: 0.7rem; text-align: center; padding: 0.5rem;">
                2026 Viviane Nishiura<br>
                Todos os direitos reservados
            </p>
        """, unsafe_allow_html=True)

# Public Header (simplified - navigation moved to sidebar)
def render_public_header():
    # Empty header for Home - the main banner is rendered in the section
    # This function is kept for other pages that may need a header
    pass

# Login Page Function
def render_login_page():
    st.markdown("""
        <style>
        .login-card {
            background: white;
            border-radius: 18px;
            box-shadow: 0 12px 40px rgba(24, 115, 140, 0.15);
            max-width: 440px;
            width: 100%;
            overflow: hidden;
        }
        .login-header {
            background: linear-gradient(135deg, #18738c 0%, #0e5a6e 100%);
            padding: 2.5rem 2rem 2rem 2rem;
            text-align: center;
        }
        .login-header img {
            width: 90px;
            height: 90px;
            border-radius: 50%;
            border: 3px solid #c5a059;
            margin-bottom: 1rem;
            object-fit: cover;
        }
        .login-header h2 {
            color: white;
            font-size: 1.5rem;
            font-weight: 700;
            margin: 0 0 0.3rem 0;
            letter-spacing: 0.5px;
        }
        .login-header p {
            color: #c5a059;
            font-size: 0.9rem;
            margin: 0;
            font-style: italic;
        }
        .login-divider {
            height: 3px;
            background: linear-gradient(90deg, #18738c, #c5a059, #18738c);
        }
        </style>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2.2, 1])
    with col2:
        logo_html = ""
        if os.path.exists(LOGO_PATH):
            import base64 as b64mod
            with open(LOGO_PATH, "rb") as lf:
                logo_b64 = b64mod.b64encode(lf.read()).decode()
            ext = LOGO_PATH.rsplit(".", 1)[-1]
            logo_html = f'<img src="data:image/{ext};base64,{logo_b64}" alt="LPS Logo">'
        
        st.markdown(f"""
            <div class="login-card" style="margin: 0 auto;">
                <div class="login-header">
                    {logo_html}
                    <h2>Liderança Psicanalítica</h2>
                    <p>Transforme sua lideranca com a ciencia do inconsciente</p>
                </div>
                <div class="login-divider"></div>
            </div>
        """, unsafe_allow_html=True)
        
        tab1, tab2 = st.tabs(["Entrar", "Cadastrar"])
        
        with tab1:
            with st.form("login_form"):
                email = st.text_input("E-mail", key="login_email", placeholder="seu@email.com")
                password = st.text_input("Senha", type="password", key="login_password")
                submit = st.form_submit_button("Entrar", use_container_width=True)
                
                if submit:
                    if email and password:
                        user = authenticate_user(email, password)
                        if user:
                            st.session_state.authenticated = True
                            st.session_state.user = user
                            st.session_state.manager_data = get_manager_by_user(user['id'])
                            st.session_state.show_welcome = True
                            st.session_state.page = "Home"
                            st.rerun()
                        else:
                            emp_login = try_employee_login(email, password)
                            if emp_login:
                                st.session_state.employee_token = emp_login
                                st.session_state.page = "EmployeeAssessment"
                                st.rerun()
                            else:
                                st.error("E-mail ou senha incorretos.")
                    else:
                        st.error("Preencha todos os campos.")
        
        with tab2:
            with st.form("register_form"):
                name = st.text_input("Nome Completo", key="reg_name")
                email = st.text_input("E-mail", key="reg_email", placeholder="seu@email.com")
                password = st.text_input("Senha", type="password", key="reg_password")
                password2 = st.text_input("Confirmar Senha", type="password", key="reg_password2")
                submit = st.form_submit_button("Criar Conta", use_container_width=True)
                
                if submit:
                    if not all([name, email, password, password2]):
                        st.error("Preencha todos os campos.")
                    elif password != password2:
                        st.error("As senhas não coincidem.")
                    elif len(password) < 6:
                        st.error("Senha deve ter no mínimo 6 caracteres.")
                    else:
                        user_id, error = register_user(email, password, name)
                        if user_id:
                            st.success("Conta criada! Faça login para continuar.")
                        else:
                            st.error(error)
        
        st.write("")
        if st.button("Voltar para Home", key="back_to_vitrine", use_container_width=True):
            st.session_state.section = "home"
            st.session_state.page = "Home"
            st.rerun()

page = st.session_state.page

# Assessment Form Component
def render_assessment_form(form_key, is_employee=False):
    responses = {}
    if is_employee:
        questions_dict = EMPLOYEE_ASSESSMENT_QUESTIONS
        cloninger_block = "Bloco 8 – Temperamento e Caráter (Cloninger Funcionário)"
        cloninger_dims = EMPLOYEE_CLONINGER_SUBDIMENSIONS
    else:
        questions_dict = ASSESSMENT_QUESTIONS
        cloninger_block = "Bloco 8 – Temperamento e Caráter (Cloninger)"
        cloninger_dims = CLONINGER_SUBDIMENSIONS

    total_blocks = len(questions_dict)
    block_names_list = list(questions_dict.keys())

    total_questions = sum(len(qs) for qs in questions_dict.values())
    questions_so_far = 0

    for block_idx, (block_name, questions) in enumerate(questions_dict.items()):
        questions_so_far += len(questions)
        progress_pct = (questions_so_far / total_questions) * 100 if total_questions > 0 else 0
        block_num = block_idx + 1
        st.markdown(f"""
            <div style="margin-bottom: 0.3rem;">
                <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 0.3rem;">
                    <span style="font-size: 0.8rem; color: #18738c; font-weight: 600;">Bloco {block_num} de {total_blocks}</span>
                    <span style="font-size: 0.8rem; color: #666;">{int(progress_pct)}% concluido</span>
                </div>
                <div style="width: 100%; background: #e8e8e8; border-radius: 10px; height: 10px; overflow: hidden;">
                    <div style="width: {max(progress_pct, 2)}%; background: linear-gradient(90deg, #18738c, #c5a059); height: 100%; border-radius: 10px; transition: width 0.3s;"></div>
                </div>
            </div>
        """, unsafe_allow_html=True)
        st.markdown(f"### {block_name}")
        if block_name == cloninger_block:
            st.markdown("""<p style='color: #666; font-size: 0.9rem; margin-bottom: 1rem;'>
                Este bloco avalia 7 subdimensoes de temperamento e carater segundo o modelo de Cloninger. 
                Cada subdimensao possui 2 questoes (pontuacao maxima: 10 por subdimensao).</p>""", unsafe_allow_html=True)
            question_to_dim = {}
            for code, info in cloninger_dims.items():
                for q_idx in info["questions"]:
                    question_to_dim[q_idx] = f"{info['name']} ({code})"
            shown_dims = set()
            for i, q in enumerate(questions):
                dim_label = question_to_dim.get(i, "")
                if dim_label and dim_label not in shown_dims:
                    shown_dims.add(dim_label)
                    st.markdown(f"""<div style='background: linear-gradient(135deg, #18738c15, #d19f0915); 
                        padding: 0.5rem 1rem; border-radius: 8px; margin: 0.8rem 0 0.4rem 0; 
                        border-left: 3px solid #18738c;'>
                        <strong style='color: #18738c;'>{dim_label}</strong></div>""", unsafe_allow_html=True)
                q_display = q.split("(")[0].strip() if "(" in q else q
                st.markdown(f'<div class="question-text">{i+1}. {q_display}</div>', unsafe_allow_html=True)
                responses[f"{block_name}_{i}"] = st.select_slider(
                    "Nota:",
                    options=[1, 2, 3, 4, 5],
                    value=3,
                    key=f"{form_key}_{block_name}_{i}"
                )
        else:
            for i, q in enumerate(questions):
                st.markdown(f'<div class="question-text">{i+1}. {q}</div>', unsafe_allow_html=True)
                responses[f"{block_name}_{i}"] = st.select_slider(
                    "Nota:",
                    options=[1, 2, 3, 4, 5],
                    value=3,
                    key=f"{form_key}_{block_name}_{i}"
                )
        st.write("---")
    st.markdown(f"""
        <div style="margin-bottom: 1rem;">
            <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 0.3rem;">
                <span style="font-size: 0.85rem; color: #18738c; font-weight: 700;">Questionario completo!</span>
                <span style="font-size: 0.85rem; color: #c5a059; font-weight: 600;">100%</span>
            </div>
            <div style="width: 100%; background: #e8e8e8; border-radius: 10px; height: 10px; overflow: hidden;">
                <div style="width: 100%; background: linear-gradient(90deg, #18738c, #c5a059); height: 100%; border-radius: 10px;"></div>
            </div>
        </div>
    """, unsafe_allow_html=True)
    return responses

def get_profile_tendency(profile):
    """Return the leadership tendency for a given profile."""
    tendencies = {
        "O Idealista Exigente": "excelencia e padroes elevados de lideranca",
        "O Contenedor Empatico": "estabilidade emocional e seguranca do grupo",
        "O Buscador de Reconhecimento": "inspiracao e motivacao da equipe",
        "O Estruturador Cauteloso": "organizacao, controle e previsibilidade",
        "O Relacional Reativo": "sensibilidade as dinamicas interpessoais",
        "O Observador Consciente": "autoconsciencia e analise profunda",
        "O Executor Decidido": "acao decisiva e foco em resultados"
    }
    return tendencies.get(profile, "desenvolvimento da equipe")

# ==========================================
# PDF EXPORT FUNCTIONS
# ==========================================

def create_pdf_styles():
    """Create custom styles for PDF reports with LPS branding."""
    styles = getSampleStyleSheet()
    
    # Title style
    styles.add(ParagraphStyle(
        name='LPSTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=HexColor('#18738c'),
        alignment=TA_CENTER,
        spaceAfter=20
    ))
    
    # Subtitle style
    styles.add(ParagraphStyle(
        name='LPSSubtitle',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=HexColor('#d19f09'),
        alignment=TA_CENTER,
        spaceAfter=15
    ))
    
    # Section header
    styles.add(ParagraphStyle(
        name='LPSSection',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=HexColor('#18738c'),
        spaceAfter=10,
        spaceBefore=15
    ))
    
    # Body text
    styles.add(ParagraphStyle(
        name='LPSBody',
        parent=styles['Normal'],
        fontSize=11,
        textColor=HexColor('#333333'),
        alignment=TA_JUSTIFY,
        spaceAfter=8
    ))
    
    # Info text
    styles.add(ParagraphStyle(
        name='LPSInfo',
        parent=styles['Normal'],
        fontSize=10,
        textColor=HexColor('#666666'),
        alignment=TA_LEFT
    ))
    
    return styles

def create_pdf_header_table():
    """Create a styled header table with LPS logo and branding."""
    logo_path = "attached_assets/logotipo_1768443722848.jpeg"
    
    # Check if logo exists
    if os.path.exists(logo_path):
        logo = Image(logo_path, width=1.2*inch, height=1.2*inch)
        header_data = [[
            logo,
            Paragraph("<font color='white' size='14'><b>Liderança Psicanalítica</b></font><br/><font color='#d19f09' size='10'>Viviane Nishiura & Equipe LPS</font>", getSampleStyleSheet()['Normal'])
        ]]
    else:
        # Fallback to text-based header if logo not found
        header_data = [[
            Paragraph("<font color='#d19f09' size='28'><b>LPS</b></font>", getSampleStyleSheet()['Normal']),
            Paragraph("<font color='white' size='12'>Liderança Psicanalítica<br/><font size='9'>Viviane Nishiura & Equipe</font></font>", getSampleStyleSheet()['Normal'])
        ]]
    
    header_table = Table(header_data, colWidths=[1.5*inch, 4*inch])
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), HexColor('#18738c')),
        ('TEXTCOLOR', (0, 0), (0, 0), HexColor('#d19f09')),
        ('TEXTCOLOR', (1, 0), (1, 0), HexColor('#FFFFFF')),
        ('ALIGN', (0, 0), (0, 0), 'CENTER'),
        ('ALIGN', (1, 0), (1, 0), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('LEFTPADDING', (0, 0), (-1, -1), 12),
        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ('BOX', (0, 0), (-1, -1), 1, HexColor('#18738c')),
    ]))
    return header_table

def generate_team_pdf_report(manager_name, employees_data, include_date=True):
    """Generate a PDF report with team assessment results."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = create_pdf_styles()
    elements = []
    
    # Styled header with LPS branding
    elements.append(create_pdf_header_table())
    elements.append(Spacer(1, 20))
    
    # Report title
    elements.append(Paragraph("Relatorio da Equipe", styles['LPSSection']))
    elements.append(Spacer(1, 20))
    
    # Manager info
    date_str = datetime.now().strftime("%d/%m/%Y") if include_date else ""
    elements.append(Paragraph(f"<b>Gestor:</b> {manager_name}", styles['LPSInfo']))
    elements.append(Paragraph(f"<b>Data:</b> {date_str}", styles['LPSInfo']))
    elements.append(Spacer(1, 20))
    
    # Table header
    table_data = [["Nome", "E-mail", "Perfil Dominante", "Perfil Secundario", "Papel de Bion"]]
    
    # Add employee data
    for emp in employees_data:
        if emp[10] == 1:  # completed
            emp_name = emp[4] or f'Funcionario {emp[3]}'
            table_data.append([
                emp_name,
                emp[5] or "N/A",
                emp[6] or "N/A",
                emp[7] or "N/A",
                emp[9] or "N/A"
            ])
    
    if len(table_data) > 1:
        table = Table(table_data, colWidths=[1.3*inch, 1.6*inch, 1.2*inch, 1.2*inch, 1.2*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#18738c')),
            ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#FFFFFF')),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), HexColor('#F5F5F5')),
            ('GRID', (0, 0), (-1, -1), 1, HexColor('#CCCCCC')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#FFFFFF'), HexColor('#F5F5F5')])
        ]))
        elements.append(table)
    else:
        elements.append(Paragraph("Nenhum funcionario completou o assessment ainda.", styles['LPSBody']))
    
    # Footer
    elements.append(Spacer(1, 30))
    elements.append(Paragraph("_" * 60, styles['LPSInfo']))
    elements.append(Paragraph("Gerado pela Plataforma LPS - Liderança Psicanalítica", styles['LPSInfo']))
    elements.append(Paragraph("Viviane Nishiura & Equipe LPS", styles['LPSInfo']))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()

def generate_individual_pdf_report(employee_data, manager_name):
    """Generate a PDF report for an individual employee."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = create_pdf_styles()
    elements = []
    
    emp_name = employee_data[4] or f'Funcionario {employee_data[3]}'
    
    # Styled header with LPS branding
    elements.append(create_pdf_header_table())
    elements.append(Spacer(1, 20))
    
    # Report title
    elements.append(Paragraph("Relatorio Individual de Assessment", styles['LPSSection']))
    elements.append(Spacer(1, 15))
    
    # Employee info
    elements.append(Paragraph(f"<b>Funcionario:</b> {emp_name}", styles['LPSInfo']))
    elements.append(Paragraph(f"<b>E-mail:</b> {employee_data[5] or 'N/A'}", styles['LPSInfo']))
    elements.append(Paragraph(f"<b>Gestor:</b> {manager_name}", styles['LPSInfo']))
    elements.append(Paragraph(f"<b>Data:</b> {datetime.now().strftime('%d/%m/%Y')}", styles['LPSInfo']))
    elements.append(Spacer(1, 25))
    
    # Profile section
    elements.append(Paragraph("Perfil de Liderança", styles['LPSSection']))
    elements.append(Paragraph(f"<b>Perfil Dominante:</b> {employee_data[6] or 'N/A'}", styles['LPSBody']))
    elements.append(Paragraph(f"<b>Perfil Secundario:</b> {employee_data[7] or 'N/A'}", styles['LPSBody']))
    elements.append(Spacer(1, 15))
    
    # Bion role section
    elements.append(Paragraph("Dinamica Grupal (Bion)", styles['LPSSection']))
    bion_role = employee_data[9] or 'N/A'
    elements.append(Paragraph(f"<b>Papel de Bion:</b> {bion_role}", styles['LPSBody']))
    
    # Bion role description
    bion_descriptions = {
        "Porta-voz": "Expressa o que o grupo sente mas nao consegue dizer. Captam tensoes inconscientes do grupo.",
        "Bode Expiatorio": "Absorve projecoes negativas do grupo. Frequentemente culpado por problemas sistemicos.",
        "Dependente": "Busca proteção constante no lider, evita autonomia. Requer contencao e orientacao.",
        "Lider de Luta-Fuga": "Reativo a ameacas reais ou imaginarias, mobiliza o grupo para ataque ou fuga.",
        "Sabotador Silencioso": "Resiste passivamente as mudancas. Concordancia superficial, boicote sutil."
    }
    if bion_role in bion_descriptions:
        elements.append(Paragraph(f"<i>{bion_descriptions[bion_role]}</i>", styles['LPSBody']))
    
    # Footer
    elements.append(Spacer(1, 40))
    elements.append(Paragraph("_" * 60, styles['LPSInfo']))
    elements.append(Paragraph("Gerado pela Plataforma LPS - Liderança Psicanalítica", styles['LPSInfo']))
    elements.append(Paragraph("Viviane Nishiura & Equipe LPS", styles['LPSInfo']))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()

def generate_ai_analysis_pdf(manager_name, analysis_text, employees_data):
    """Generate a PDF report from AI analysis with LPS branding."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = create_pdf_styles()
    elements = []
    
    # Styled header with LPS branding
    elements.append(create_pdf_header_table())
    elements.append(Spacer(1, 20))
    
    # Report title
    elements.append(Paragraph("Analise de IA - Dinamicas de Equipe", styles['LPSSection']))
    elements.append(Spacer(1, 15))
    
    # Manager info
    elements.append(Paragraph(f"<b>Gestor:</b> {manager_name}", styles['LPSInfo']))
    elements.append(Paragraph(f"<b>Data da Analise:</b> {datetime.now().strftime('%d/%m/%Y as %H:%M')}", styles['LPSInfo']))
    elements.append(Spacer(1, 20))
    
    # Team summary
    elements.append(Paragraph("Composição da Equipe Analisada", styles['LPSSection']))
    completed_count = sum(1 for emp in employees_data if emp[10] == 1)
    elements.append(Paragraph(f"<b>Funcionarios mapeados:</b> {completed_count}", styles['LPSBody']))
    
    # Bion distribution
    bion_roles = {}
    for emp in employees_data:
        if emp[10] == 1 and emp[9]:
            bion_roles[emp[9]] = bion_roles.get(emp[9], 0) + 1
    if bion_roles:
        elements.append(Paragraph("<b>Distribuicao de Papeis de Bion:</b>", styles['LPSBody']))
        for role, count in bion_roles.items():
            elements.append(Paragraph(f"  - {role}: {count} funcionario(s)", styles['LPSBody']))
    
    elements.append(Spacer(1, 20))
    
    # AI Analysis
    elements.append(Paragraph("Analise da Consultora de IA", styles['LPSSection']))
    
    # Split analysis into paragraphs
    paragraphs = analysis_text.split('\n\n')
    for para in paragraphs:
        if para.strip():
            import re as re_mod
            clean_para = para.replace('\n', ' ').strip()
            clean_para = re_mod.sub(r'\*\*\*(.+?)\*\*\*', r'<b><i>\1</i></b>', clean_para)
            clean_para = re_mod.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', clean_para)
            clean_para = re_mod.sub(r'\*(.+?)\*', r'<i>\1</i>', clean_para)
            try:
                elements.append(Paragraph(clean_para, styles['LPSBody']))
            except Exception:
                safe_para = para.replace('\n', ' ').strip().replace("<", "&lt;").replace(">", "&gt;")
                elements.append(Paragraph(safe_para, styles['LPSBody']))
            elements.append(Spacer(1, 5))
    
    # Footer
    elements.append(Spacer(1, 30))
    elements.append(Paragraph("_" * 60, styles['LPSInfo']))
    elements.append(Paragraph("Analise gerada pela LPChat - Consultora de IA em Psicanálise e Neurociência", styles['LPSInfo']))
    elements.append(Paragraph("Plataforma LPS - Viviane Nishiura & Equipe LPS", styles['LPSInfo']))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()

def _pdf_page_decorator(canvas_obj, doc_obj):
    """Draw LPS logo at top-right and page number at bottom of every page."""
    canvas_obj.saveState()
    logo_path = "attached_assets/logotipo_1768443722848.jpeg"
    page_width = A4[0]
    if os.path.exists(logo_path):
        try:
            canvas_obj.drawImage(logo_path,
                page_width - doc_obj.rightMargin - 0.7*inch,
                A4[1] - 0.62*inch,
                width=0.65*inch, height=0.65*inch,
                preserveAspectRatio=True, mask='auto')
        except Exception:
            pass
    canvas_obj.setStrokeColor(HexColor('#18738c'))
    canvas_obj.setLineWidth(0.5)
    canvas_obj.line(doc_obj.leftMargin, A4[1] - 0.72*inch,
                    page_width - doc_obj.rightMargin, A4[1] - 0.72*inch)
    canvas_obj.setStrokeColor(HexColor('#c5a059'))
    canvas_obj.setLineWidth(0.4)
    canvas_obj.line(doc_obj.leftMargin, 0.55*inch,
                    page_width - doc_obj.rightMargin, 0.55*inch)
    canvas_obj.setFont('Helvetica', 8)
    canvas_obj.setFillColor(HexColor('#888888'))
    canvas_obj.drawRightString(page_width - doc_obj.rightMargin, 0.38*inch,
        f"Pág. {canvas_obj.getPageNumber()}")
    canvas_obj.drawCentredString(page_width / 2, 0.38*inch,
        "Plataforma LPS | Viviane Nishiura | Confidencial")
    canvas_obj.restoreState()

def generate_laudo_pdf(laudo_text, respondent_name, dominant, secondary, bion_role, respondent_type="gestor"):
    """Generate PDF that is a direct mirror of the .docx content.
    Reads the .docx file directly paragraph by paragraph.
    Every paragraph goes into the PDF story as-is. No filtering, no summarizing.
    SimpleDocTemplate handles pagination automatically - unlimited pages."""
    from docx import Document as DocxDocument

    docx_path = find_docx_file(dominant, secondary, respondent_type)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=0.85*inch, bottomMargin=0.7*inch,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('LaudoTitle', parent=styles['Normal'],
        fontSize=20, textColor=HexColor('#18738c'), fontName='Helvetica-Bold',
        alignment=TA_CENTER, spaceBefore=4, spaceAfter=4, leading=24)
    subtitle_style = ParagraphStyle('LaudoSub', parent=styles['Normal'],
        fontSize=12, textColor=HexColor('#555555'), fontName='Helvetica',
        alignment=TA_CENTER, spaceAfter=12)
    tipo_style = ParagraphStyle('LaudoTipo', parent=styles['Normal'],
        fontSize=11, textColor=HexColor('#c5a059'), fontName='Helvetica-Bold',
        alignment=TA_CENTER, spaceAfter=14)
    meta_style = ParagraphStyle('LaudoMeta', parent=styles['Normal'],
        fontSize=11, textColor=HexColor('#333333'), fontName='Helvetica',
        spaceAfter=2, leading=14)
    section_style = ParagraphStyle('LaudoSection', parent=styles['Normal'],
        fontSize=13, textColor=HexColor('#18738c'), fontName='Helvetica-Bold',
        spaceBefore=16, spaceAfter=4, leading=16)
    body_style = ParagraphStyle('LaudoBody', parent=styles['Normal'],
        fontSize=11, textColor=HexColor('#2c2c2c'), fontName='Helvetica',
        alignment=TA_JUSTIFY, spaceAfter=6, leading=15)

    story = []

    story.append(Spacer(1, 6))
    story.append(Paragraph("Relatório de Perfil Psicanalítico", title_style))
    story.append(Paragraph("Viviane Nishiura", subtitle_style))
    tipo_label = "Laudo Psicanalítico de Liderança" if respondent_type == "gestor" else "Laudo Psicanalítico — Perfil de Equipe"
    story.append(Paragraph(tipo_label, tipo_style))

    meta_rows = [
        ["<b>Avaliado(a):</b>", respondent_name or "---"],
        ["<b>Perfil:</b>", f"{dominant} + {secondary}"],
        ["<b>Papel de Bion:</b>", bion_role or "---"],
        ["<b>Data:</b>", datetime.now().strftime('%d/%m/%Y')],
    ]
    meta_tbl = Table(
        [[Paragraph(r[0], meta_style), Paragraph(r[1], meta_style)] for r in meta_rows],
        colWidths=[1.6*inch, 4.4*inch])
    meta_tbl.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('TOPPADDING', (0,0), (-1,-1), 2),
        ('BOTTOMPADDING', (0,0), (-1,-1), 2),
        ('LEFTPADDING', (0,0), (0,-1), 0),
    ]))
    story.append(meta_tbl)
    story.append(Spacer(1, 6))

    sep = Table([[""]], colWidths=[doc.width])
    sep.setStyle(TableStyle([
        ('LINEABOVE', (0,0), (-1,0), 1.5, HexColor('#18738c')),
        ('TOPPADDING', (0,0), (-1,0), 0),
        ('BOTTOMPADDING', (0,0), (-1,0), 8),
    ]))
    story.append(sep)

    if docx_path and os.path.exists(docx_path):
        docx_doc = DocxDocument(docx_path)
        for para in docx_doc.paragraphs:
            raw = para.text.strip()
            if not raw:
                story.append(Spacer(1, 4))
                continue

            is_bold = bool(para.runs) and all(r.bold for r in para.runs if r.text.strip())
            looks_like_header = False
            if raw and raw[0].isdigit():
                dot_pos = raw.find('.')
                if 0 < dot_pos <= 3:
                    looks_like_header = True

            safe = raw.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            safe = safe.replace("\u201c", "&ldquo;").replace("\u201d", "&rdquo;")
            safe = safe.replace("\u2013", "&ndash;").replace("\u2014", "&mdash;")

            if looks_like_header:
                story.append(Paragraph(safe, section_style))
                gold = Table([[""]], colWidths=[doc.width])
                gold.setStyle(TableStyle([
                    ('LINEABOVE', (0,0), (-1,0), 0.6, HexColor('#c5a059')),
                    ('TOPPADDING', (0,0), (-1,0), 0),
                    ('BOTTOMPADDING', (0,0), (-1,0), 5),
                ]))
                story.append(gold)
            elif is_bold:
                story.append(Paragraph(f"<b>{safe}</b>", body_style))
            else:
                story.append(Paragraph(safe, body_style))
    else:
        story.append(Paragraph("Documento .docx nao encontrado para este perfil.", body_style))

    story.append(Spacer(1, 20))
    end_sep = Table([[""]], colWidths=[doc.width])
    end_sep.setStyle(TableStyle([
        ('LINEABOVE', (0,0), (-1,0), 1, HexColor('#18738c')),
        ('TOPPADDING', (0,0), (-1,0), 8),
    ]))
    story.append(end_sep)
    footer_style = ParagraphStyle('LaudoFooter', parent=styles['Normal'],
        fontSize=9, textColor=HexColor('#777777'), fontName='Helvetica',
        alignment=TA_CENTER, spaceAfter=2)
    story.append(Paragraph("Plataforma LPS — Liderança Psicanalítica | Viviane Nishiura &amp; Equipe LPS", footer_style))
    story.append(Paragraph(f"Documento gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}", footer_style))
    story.append(Paragraph("Este documento é confidencial.", footer_style))

    doc.build(story, onFirstPage=_pdf_page_decorator, onLaterPages=_pdf_page_decorator)
    buffer.seek(0)
    return buffer.getvalue()

def generate_manager_guide_pdf():
    """Generate the Manager's Guide PDF with LPS methodology."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = create_pdf_styles()
    elements = []
    
    # Styled header with LPS branding
    elements.append(create_pdf_header_table())
    elements.append(Spacer(1, 20))
    
    # Title
    elements.append(Paragraph("Manual do Gestor LPS", styles['LPSTitle']))
    elements.append(Paragraph("Guia Pratico para Liderança Psicanalítica", styles['LPSInfo']))
    elements.append(Spacer(1, 30))
    
    # Section 1: Como interpretar seu Perfil
    elements.append(Paragraph("1. Como Interpretar seu Perfil", styles['LPSSection']))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(
        "Seu perfil de lideranca revela seus arquetipos inconscientes dominantes - padroes de comportamento "
        "que operam abaixo da consciencia e influenciam como voce lidera sua equipe.",
        styles['LPSBody']))
    elements.append(Spacer(1, 8))
    
    # Profile descriptions
    profiles = [
        ("O Idealista Exigente", "Pressao interna constante por excelencia. Seu ponto forte e manter padroes altos de qualidade. "
         "Atenção: autocritica excessiva pode gerar esgotamento e paralisia."),
        ("O Contenedor Empatico", "Absorve tensoes do grupo e cria seguranca emocional. Essencial em crises. "
         "Atenção: risco de sobrecarga emocional e esgotamento."),
        ("O Buscador de Reconhecimento", "Inspira e motiva atraves de carisma e visao. Mobiliza a equipe para grandes objetivos. "
         "Atenção: pode depender demais de validacao externa."),
        ("O Estruturador Cauteloso", "Organiza processos e garante previsibilidade. Equipes bem estruturadas. "
         "Atenção: rigidez excessiva pode inibir criatividade e inovacao."),
        ("O Relacional Reativo", "Sensivel as dinamicas interpessoais e transferencias. Percebe padroes relacionais. "
         "Atenção: pode ser ativado emocionalmente por conflitos, reagindo de forma desproporcional."),
        ("O Observador Consciente", "Analisa profundamente e reflete antes de agir. Decisoes ponderadas e conscientes. "
         "Atenção: pode parecer distante ou demorar demais para decidir."),
        ("O Executor Decidido", "Orientado a acao e resultados rapidos. Faz as coisas acontecerem. "
         "Atenção: pode atropelar processos e pessoas por impaciencia.")
    ]
    
    for name, desc in profiles:
        elements.append(Paragraph(f"<b>{name}:</b> {desc}", styles['LPSBody']))
        elements.append(Spacer(1, 5))
    
    elements.append(Spacer(1, 20))
    
    # Section 2: Mapeamento de Equipe
    elements.append(Paragraph("2. Mapeamento de Equipe", styles['LPSSection']))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(
        "O LPTest mapeia os perfis inconscientes de sua equipe, permitindo entender dinamicas grupais "
        "que impactam produtividade, conflitos e turnover.",
        styles['LPSBody']))
    elements.append(Spacer(1, 8))
    
    elements.append(Paragraph("<b>Como usar para reduzir turnover:</b>", styles['LPSBody']))
    elements.append(Paragraph(
        "- Identifique funcionarios cujo perfil nao se adequa ao cargo atual", styles['LPSBody']))
    elements.append(Paragraph(
        "- Realoque pessoas com base em suas tendencias naturais", styles['LPSBody']))
    elements.append(Paragraph(
        "- Crie pares complementares (ex: Estruturador + Criativo)", styles['LPSBody']))
    elements.append(Spacer(1, 8))
    
    elements.append(Paragraph("<b>Como usar para reduzir conflitos:</b>", styles['LPSBody']))
    elements.append(Paragraph(
        "- Identifique Bodes Expiatorios e proteja-os de projecoes negativas", styles['LPSBody']))
    elements.append(Paragraph(
        "- Reconheca Porta-Vozes como sensores do clima organizacional", styles['LPSBody']))
    elements.append(Paragraph(
        "- Transforme Lideres de Luta-Fuga em agentes de mudanca construtiva", styles['LPSBody']))
    
    elements.append(Spacer(1, 20))
    
    # Section 3: Uso Estratégico do LPChat
    elements.append(Paragraph("3. Uso Estrategico do LPChat", styles['LPSSection']))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(
        "O LPChat é sua consultora de IA especializada em psicanálise e neurociência aplicada à liderança. "
        "Para obter os melhores insights, faça perguntas específicas sobre sua equipe.",
        styles['LPSBody']))
    elements.append(Spacer(1, 8))
    
    elements.append(Paragraph("<b>Perguntas recomendadas:</b>", styles['LPSBody']))
    questions = [
        "Quais conflitos inconscientes podem surgir entre [Funcionario A] e [Funcionario B]?",
        "Qual funcionario seria ideal para liderar o projeto X, considerando os perfis mapeados?",
        "Como posso dar feedback construtivo para um Dependente sem gerar mais dependencia?",
        "Quais dinamicas de transferencia podem estar afetando minha relacao com a equipe?",
        "Como aplicar conceitos de neurociencia para reduzir o estresse no time?"
    ]
    for q in questions:
        elements.append(Paragraph(f"- {q}", styles['LPSBody']))
    
    elements.append(Spacer(1, 20))
    
    # Section 4: Passo a Passo da Mentoria
    elements.append(Paragraph("4. Passo a Passo da Mentoria", styles['LPSSection']))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(
        "A Mentoria Executiva LPS e o momento de aprofundar sua jornada de liderança consciente. "
        "Prepare-se adequadamente para maximizar os resultados.",
        styles['LPSBody']))
    elements.append(Spacer(1, 8))
    
    elements.append(Paragraph("<b>Antes da sessão:</b>", styles['LPSBody']))
    elements.append(Paragraph("1. Revise seu perfil de lideranca no Dashboard", styles['LPSBody']))
    elements.append(Paragraph("2. Analise os resultados do mapeamento de equipe", styles['LPSBody']))
    elements.append(Paragraph("3. Identifique 2-3 desafios especificos que deseja abordar", styles['LPSBody']))
    elements.append(Paragraph("4. Anote situacoes concretas para discutir", styles['LPSBody']))
    elements.append(Spacer(1, 8))
    
    elements.append(Paragraph("<b>Durante a sessão:</b>", styles['LPSBody']))
    elements.append(Paragraph("- Compartilhe abertamente seus desafios", styles['LPSBody']))
    elements.append(Paragraph("- Pergunte sobre padroes inconscientes que nao consegue ver", styles['LPSBody']))
    elements.append(Paragraph("- Solicite exercicios praticos para aplicar no dia-a-dia", styles['LPSBody']))
    elements.append(Spacer(1, 8))
    
    elements.append(Paragraph("<b>Como agendar:</b>", styles['LPSBody']))
    elements.append(Paragraph(
        "Acesse o menu 'Mentoria' no Dashboard ou envie mensagem via WhatsApp para agendar sua sessao.",
        styles['LPSBody']))
    
    # Footer
    elements.append(Spacer(1, 30))
    elements.append(Paragraph("_" * 60, styles['LPSInfo']))
    elements.append(Paragraph("Plataforma LPS - Liderança Psicanalítica", styles['LPSInfo']))
    elements.append(Paragraph("Viviane Nishiura & Equipe LPS", styles['LPSInfo']))
    elements.append(Paragraph(f"Versao: {datetime.now().strftime('%m/%Y')}", styles['LPSInfo']))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()

def generate_team_chart(employees_data, manager_name):
    """Generate a team profile distribution chart as PNG."""
    # Count profiles
    profile_counts = {}
    bion_counts = {}
    
    for emp in employees_data:
        if emp[10] == 1:  # completed
            # Count dominant profiles
            if emp[6]:
                profile_counts[emp[6]] = profile_counts.get(emp[6], 0) + 1
            # Count Bion roles
            if emp[9]:
                bion_counts[emp[9]] = bion_counts.get(emp[9], 0) + 1
    
    if not profile_counts and not bion_counts:
        return None
    
    # Create figure with two subplots
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))
    fig.suptitle(f'Perfil da Equipe - {manager_name}', fontsize=16, color='#18738c', fontweight='bold')
    
    # Colors matching LPS brand
    colors_profile = ['#18738c', '#1a4f7a', '#2d6a9f', '#4080b5', '#5596c9', '#6aabdc']
    colors_bion = ['#d19f09', '#e6c54e', '#d9b83e', '#ccab2e', '#bf9e1e', '#b2910e']
    
    # Profile distribution pie chart
    if profile_counts:
        labels1 = list(profile_counts.keys())
        sizes1 = list(profile_counts.values())
        ax1.pie(sizes1, labels=labels1, colors=colors_profile[:len(labels1)], autopct='%1.0f%%', startangle=90)
        ax1.set_title('Perfis de Liderança', fontsize=12, color='#18738c')
    else:
        ax1.text(0.5, 0.5, 'Sem dados', ha='center', va='center')
        ax1.set_title('Perfis de Liderança', fontsize=12, color='#18738c')
    
    # Bion roles bar chart
    if bion_counts:
        labels2 = list(bion_counts.keys())
        sizes2 = list(bion_counts.values())
        bars = ax2.barh(labels2, sizes2, color=colors_bion[:len(labels2)])
        ax2.set_xlabel('Quantidade')
        ax2.set_title('Papeis de Bion', fontsize=12, color='#18738c')
        ax2.set_xlim(0, max(sizes2) + 1)
        for bar, size in zip(bars, sizes2):
            ax2.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height()/2, str(size), va='center')
    else:
        ax2.text(0.5, 0.5, 'Sem dados', ha='center', va='center')
        ax2.set_title('Papeis de Bion', fontsize=12, color='#18738c')
    
    plt.tight_layout()
    
    # Save to buffer
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    buffer.seek(0)
    return buffer.getvalue()

def generate_radar_chart(block_sums, profile_name=""):
    """Generate a radar chart showing the 7 axes of the assessment."""
    import numpy as np
    
    categories = [
        'Autoridade\nSuperego',
        'Contencao\nEmpatia',
        'Narcisismo\nMotivacao',
        'Estrutura\nControle',
        'Relacional\nTransferencia',
        'Observacao\nMentalizacao',
        'Execucao\nUrgencia'
    ]
    
    block_keys = [
        "Bloco 1 – Autoridade Interna, Autoimagem e Superego",
        "Bloco 2 – Contenção Emocional, Empatia e Círculo de Segurança",
        "Bloco 3 – Narcisismo, Reconhecimento e Motivação",
        "Bloco 4 – Estrutura, Controle e Tolerância à Ambiguidade",
        "Bloco 5 – Dinâmicas Relacionais, Transferência e Contratransferência",
        "Bloco 6 – Autoconsciência, Mentalização e Defesas do Ego",
        "Bloco 7 – Ação, Urgência e Foco na Entrega"
    ]
    
    values = []
    for key in block_keys:
        val = block_sums.get(key, 20)
        normalized = (val / 40) * 100
        values.append(normalized)
    
    # Number of variables
    N = len(categories)
    
    # Compute angle for each category
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    values_plot = values + [values[0]]  # Complete the loop
    angles += angles[:1]  # Complete the loop
    
    # Create figure
    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))
    
    # Plot data
    ax.fill(angles, values_plot, color='#18738c', alpha=0.25)
    ax.plot(angles, values_plot, color='#18738c', linewidth=2)
    
    # Add markers
    ax.scatter(angles[:-1], values, color='#c5a059', s=100, zorder=5)
    
    # Set category labels
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, size=10, color='#18738c')
    
    # Set y-axis limits
    ax.set_ylim(0, 100)
    ax.set_yticks([20, 40, 60, 80, 100])
    ax.set_yticklabels(['20%', '40%', '60%', '80%', '100%'], size=8, color='gray')
    
    # Title
    title = f'Perfil de Liderança: {profile_name}' if profile_name else 'Perfil de Liderança'
    ax.set_title(title, size=14, color='#18738c', fontweight='bold', pad=20)
    
    # Style
    ax.grid(True, color='gray', linestyle='-', linewidth=0.5, alpha=0.5)
    ax.spines['polar'].set_color('#18738c')
    
    # Save to buffer
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    buffer.seek(0)
    return buffer.getvalue()

def find_docx_file(dominant, secondary, respondent_type="gestor"):
    """Find the matching .docx file for a profile combination.
    Returns the file path or None if not found."""
    import glob as glob_module
    
    dom_sigla = ARCHETYPE_TO_SIGLA.get(dominant, "")
    sec_sigla = ARCHETYPE_TO_SIGLA.get(secondary, "")
    
    if not dom_sigla or not sec_sigla:
        return None
    
    combo = f"{dom_sigla}{sec_sigla}"
    pattern = f"attached_assets/{combo}_{respondent_type}_*.docx"
    matches = glob_module.glob(pattern)
    
    if not matches:
        pattern_alt = f"attached_assets/{sec_sigla}{dom_sigla}_{respondent_type}_*.docx"
        matches = glob_module.glob(pattern_alt)
    
    if matches:
        return sorted(matches)[-1]
    return None

def extract_docx_profile_text(dominant, secondary, respondent_type="gestor"):
    """Extract the COMPLETE, UNMODIFIED text from matching .docx file.
    Preserves all original paragraph breaks and bold/italic formatting.
    Bold text is wrapped in **text** and italic in *text* to preserve
    the original Word formatting for display and PDF rendering.
    No text is summarized, altered, or omitted."""
    try:
        from docx import Document as DocxDocument
        
        file_path = find_docx_file(dominant, secondary, respondent_type)
        if not file_path:
            return None
        
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            if not para.runs:
                full_text.append(para.text)
                continue
            para_parts = []
            for run in para.runs:
                text = run.text
                if not text:
                    continue
                if run.bold and run.italic:
                    text = f"***{text}***"
                elif run.bold:
                    text = f"**{text}**"
                elif run.italic:
                    text = f"*{text}*"
                para_parts.append(text)
            full_text.append("".join(para_parts))
        
        result = "\n".join(full_text)
        while result.endswith("\n"):
            result = result[:-1]
        return result
    except Exception as e:
        print(f"[DOCX] Error extracting profile: {e}")
        return None

GESTOR_HEADER_PATTERNS = [
    ("1. Visao Geral", [
        "visão geral",
        "visao geral",
    ]),
    ("2. Essencia Psicanalítica", [
        "essência psicanalítica",
        "essencia psicanalítica",
    ]),
    ("3. Motivações Inconscientes", [
        "motivações inconscientes",
        "motivacoes inconscientes",
    ]),
    ("4. Forcas", [
        "forças (manifestações positivas)",
        "forcas (manifestacoes positivas)",
        "forças",
        "forcas",
    ]),
    ("5. Sombra", [
        "sombra (riscos e manifestações negativas)",
        "sombra (riscos e manifestacoes negativas)",
        "sombra (riscos)",
        "sombra",
    ]),
    ("6. Estilo de Liderança e Impacto", [
        "estilo de liderança e impacto",
        "estilo de lideranca e impacto",
        "estilo de liderança",
        "estilo de lideranca",
        "impacto no círculo de segurança",
        "impacto no circulo de seguranca",
    ]),
    ("7. Função de Liderança", [
        "função de liderança",
        "função de lideranca",
    ]),
    ("8. Dinamica Emocional (Sinek)", [
        "dinâmica emocional (sinek + neurociência)",
        "dinamica emocional (sinek + neurociencia)",
        "dinâmica emocional",
        "dinamica emocional",
    ]),
    ("9. Melhor Aproveitamento", [
        "melhor aproveitamento do líder",
        "melhor aproveitamento do lider",
        "melhor aproveitamento",
    ]),
    ("10. Riscos de Alocação", [
        "riscos se mal alocado",
        "riscos de alocação",
        "riscos de alocacao",
    ]),
    ("11. Recomendações de Desenvolvimento", [
        "recomendações de desenvolvimento",
        "recomendacoes de desenvolvimento",
    ]),
    ("12. Sintese", [
        "síntese",
        "sintese",
    ]),
]

FUNCIONARIO_HEADER_PATTERNS = [
    ("1. Visao Geral", [
        "visão geral",
        "visao geral",
    ]),
    ("2. Essencia Psicanalítica", [
        "essência psicanalítica",
        "essencia psicanalítica",
    ]),
    ("3. Motivações Inconscientes", [
        "motivações inconscientes",
        "motivacoes inconscientes",
    ]),
    ("4. Forcas", [
        "forças",
        "forcas",
    ]),
    ("5. Sombra", [
        "sombra",
    ]),
    ("6. Papeis Grupais (Bion)", [
        "tendências de papéis grupais (bion)",
        "tendencias de papeis grupais (bion)",
        "tendências de papéis grupais",
        "tendencias de papeis grupais",
        "papéis grupais (bion)",
        "papeis grupais (bion)",
    ]),
    ("7. Dinamica Emocional (Sinek)", [
        "dinâmica emocional (sinek + neurociência)",
        "dinamica emocional (sinek + neurociencia)",
        "dinâmica emocional",
        "dinamica emocional",
    ]),
    ("8. Melhor Aproveitamento", [
        "melhor aproveitamento na equipe",
        "melhor aproveitamento",
    ]),
    ("9. Riscos de Alocação", [
        "riscos se mal alocado",
        "riscos de alocação",
        "riscos de alocacao",
    ]),
    ("10. Recomendações ao Gestor", [
        "recomendações ao gestor",
        "recomendacoes ao gestor",
    ]),
    ("11. Reflexões para o Proprio Perfil", [
        "reflexões para o próprio perfil",
        "reflexoes para o proprio perfil",
    ]),
    ("12. Sintese", [
        "síntese",
        "sintese",
    ]),
]

LAUDO_SECTIONS_GESTOR = [s[0] for s in GESTOR_HEADER_PATTERNS]
LAUDO_SECTIONS_FUNCIONARIO = [s[0] for s in FUNCIONARIO_HEADER_PATTERNS]
LAUDO_SECTIONS = LAUDO_SECTIONS_GESTOR

def get_laudo_sections_for_type(respondent_type="gestor"):
    """Return the appropriate section list for the respondent type."""
    if respondent_type == "funcionario":
        return LAUDO_SECTIONS_FUNCIONARIO
    return LAUDO_SECTIONS_GESTOR

def _match_header_line(line_text, patterns):
    """Check if a line is a section header. Returns (section_key, remaining_content) or (None, None).
    A header is detected ONLY when the line starts with (optionally numbered) a known section title.
    If the header line also contains body content (separated by space after the title), that content
    is returned as 'remaining'. Short trailing words that are part of the title (like 'Principais',
    'no Grupo') are NOT treated as content."""
    clean = line_text.strip()
    if not clean:
        return None, None
    
    plain = clean.replace("***", "").replace("**", "").replace("*", "")
    test = plain.lower()
    test_no_num = test
    num_prefix_len = 0
    if test and test[0].isdigit():
        dot_idx = test.find(".")
        if dot_idx >= 0 and dot_idx <= 3:
            test_no_num = test[dot_idx + 1:].strip()
            num_prefix_len = dot_idx + 1
            while num_prefix_len < len(plain) and plain[num_prefix_len] in ' \t':
                num_prefix_len += 1
    
    for section_key, keywords in patterns:
        for kw in keywords:
            if test_no_num.startswith(kw) or test.startswith(kw):
                if test_no_num.startswith(kw):
                    after_kw_pos = num_prefix_len + len(kw)
                else:
                    after_kw_pos = len(kw)
                
                remaining_plain = plain[after_kw_pos:].strip() if after_kw_pos < len(plain) else ""
                
                if remaining_plain:
                    remaining_plain = remaining_plain.lstrip("(").lstrip(")").lstrip(":").lstrip("-").strip()
                
                if remaining_plain and len(remaining_plain) < 40:
                    remaining_plain = ""
                
                if remaining_plain:
                    rp_start = clean.find(remaining_plain[:20])
                    remaining = clean[rp_start:].strip() if rp_start >= 0 else remaining_plain
                else:
                    remaining = ""
                
                return section_key, remaining
    
    return None, None

def parse_docx_into_sections(docx_text, respondent_type="gestor"):
    """Parse docx text into structured sections. FAITHFUL extraction:
    - Everything between one header and the next belongs to that section
    - No text is cut, summarized, or modified
    - If a section is not found, it is left empty (never AI-generated)
    - Inline content after a header (on the same line) is preserved"""
    if not docx_text:
        return {}
    
    patterns = GESTOR_HEADER_PATTERNS if respondent_type == "gestor" else FUNCIONARIO_HEADER_PATTERNS
    
    all_lines = docx_text.split("\n")
    
    sections = {}
    current_section = None
    current_content = []
    title_line_seen = False
    
    for line in all_lines:
        stripped = line.strip()
        
        section_key, remaining = _match_header_line(stripped, patterns)
        
        if section_key:
            if current_section is not None:
                sections[current_section] = "\n".join(current_content).strip()
            
            current_section = section_key
            current_content = []
            title_line_seen = True
            
            if remaining:
                current_content.append(remaining)
        elif not title_line_seen and current_section is None:
            title_line_seen = True
        elif current_section is not None:
            current_content.append(line)
    
    if current_section is not None:
        sections[current_section] = "\n".join(current_content).strip()
    
    if not sections and docx_text:
        section_list = LAUDO_SECTIONS_GESTOR if respondent_type == "gestor" else LAUDO_SECTIONS_FUNCIONARIO
        sections[section_list[0]] = docx_text
    
    return sections

def parse_laudo_sections(laudo_text, respondent_type="gestor"):
    """Parse laudo text into structured sections. Pure docx extraction, no AI."""
    return parse_docx_into_sections(laudo_text, respondent_type)

def generate_docx_laudo(dominant, secondary, bion_role, respondent_name="", respondent_type="gestor"):
    """Extract laudo content directly from the corresponding .docx file.
    No AI generation - uses the exact text from the document."""
    try:
        docx_text = extract_docx_profile_text(dominant, secondary, respondent_type)
        
        if not docx_text:
            dom_sigla = ARCHETYPE_TO_SIGLA.get(dominant, "??")
            sec_sigla = ARCHETYPE_TO_SIGLA.get(secondary, "??")
            return None, f"Arquivo de perfil nao encontrado: {dom_sigla}{sec_sigla}_{respondent_type}.docx. Verifique se o arquivo existe no diretorio."
        
        return docx_text, None
    except Exception as e:
        return None, f"Erro ao extrair laudo do documento: {str(e)}"

def generate_ai_laudo(dominant, secondary, bion_role, block_sums, respondent_name="", profile_text=None, respondent_type="gestor"):
    """Legacy wrapper - now reads from .docx files instead of AI generation.
    Kept for backward compatibility with existing call sites."""
    return generate_docx_laudo(dominant, secondary, bion_role, respondent_name, respondent_type)

def save_assessment_responses(respondent_id, respondent_type, responses):
    """Save each individual response (1-5) to the database for future AI analysis."""
    conn = sqlite3.connect('lps_data.db')
    c = conn.cursor()
    
    for block_name, questions in ASSESSMENT_QUESTIONS.items():
        for q_idx, question_text in enumerate(questions):
            response_key = f"{block_name}_{q_idx}"
            response_value = responses.get(response_key, 3)
            
            response_id = str(uuid.uuid4())
            c.execute("""INSERT INTO assessment_responses 
                        (id, respondent_id, respondent_type, block_name, question_index, question_text, response_value)
                        VALUES (?, ?, ?, ?, ?, ?, ?)""",
                     (response_id, respondent_id, respondent_type, block_name, q_idx, question_text, response_value))
    
    conn.commit()
    conn.close()

def save_employee_assessment_responses(respondent_id, respondent_type, responses):
    """Save each individual employee response to the database."""
    conn = sqlite3.connect('lps_data.db')
    c = conn.cursor()
    
    for block_name, questions in EMPLOYEE_ASSESSMENT_QUESTIONS.items():
        for q_idx, question_text in enumerate(questions):
            response_key = f"{block_name}_{q_idx}"
            response_value = responses.get(response_key, 3)
            
            response_id = str(uuid.uuid4())
            c.execute("""INSERT INTO assessment_responses 
                        (id, respondent_id, respondent_type, block_name, question_index, question_text, response_value)
                        VALUES (?, ?, ?, ?, ?, ?, ?)""",
                     (response_id, respondent_id, respondent_type, block_name, q_idx, question_text, response_value))
    
    conn.commit()
    conn.close()

def calculate_cloninger_scores(responses):
    """Calculate Cloninger temperament subdimension scores from Block 8."""
    cloninger_block = "Bloco 8 – Temperamento e Carater (Cloninger)"
    scores = {}
    for code, info in CLONINGER_SUBDIMENSIONS.items():
        total = sum(responses.get(f"{cloninger_block}_{q}", 3) for q in info["questions"])
        scores[code] = {"name": info["name"], "score": total, "max": 10}
    return scores

def calculate_profile(responses):
    block_sums = {}
    for block, questions in ASSESSMENT_QUESTIONS.items():
        block_sums[block] = sum(responses.get(f"{block}_{i}", 3) for i in range(len(questions)))
    
    archetype_blocks = {k: v for k, v in block_sums.items() if k in BLOCK_TO_PROFILE}
    
    sorted_blocks = sorted(archetype_blocks.items(), key=lambda x: x[1], reverse=True)
    dom_key = sorted_blocks[0][0]
    sec_key = sorted_blocks[1][0]
    
    dominant_name = BLOCK_TO_PROFILE[dom_key]
    secondary_name = BLOCK_TO_PROFILE[sec_key]
    
    profile_text = extract_docx_profile_text(dominant_name, secondary_name, "gestor")
    if profile_text:
        details = PROFILES_DB.get(dominant_name, {}).get(secondary_name, {
            "forcas": f"Combinacao de {dominant_name} e {secondary_name}.",
            "riscos": "Necessidade de vigilia sobre dinamicas da equipe.",
            "recomendacoes": "Agende mentoria personalizada."
        })
        details["profile_text"] = profile_text
    else:
        details = PROFILES_DB.get(dominant_name, {}).get(secondary_name, {
            "forcas": f"Combinacao de {dominant_name} e {secondary_name}.",
            "riscos": "Necessidade de vigilia sobre dinamicas da equipe.",
            "recomendacoes": "Agende mentoria personalizada."
        })
    
    short_block_sums = {}
    for full_name, value in archetype_blocks.items():
        short_name = BLOCK_SHORT_NAMES.get(full_name, full_name)
        short_block_sums[short_name] = value
    
    bion_role = classify_bion_role(short_block_sums)
    
    cloninger_scores = calculate_cloninger_scores(responses)
    
    return dominant_name, secondary_name, details, bion_role, block_sums, cloninger_scores

def calculate_employee_profile(responses):
    """Calculate employee profile using employee-specific questions and Bion roles."""
    block_sums = {}
    for block, questions in EMPLOYEE_ASSESSMENT_QUESTIONS.items():
        block_sums[block] = sum(responses.get(f"{block}_{i}", 3) for i in range(len(questions)))
    
    archetype_blocks = {k: v for k, v in block_sums.items() if k in EMPLOYEE_BLOCK_TO_PROFILE}
    
    sorted_blocks = sorted(archetype_blocks.items(), key=lambda x: x[1], reverse=True)
    dom_key = sorted_blocks[0][0]
    sec_key = sorted_blocks[1][0]
    
    dominant_name = EMPLOYEE_BLOCK_TO_PROFILE[dom_key]
    secondary_name = EMPLOYEE_BLOCK_TO_PROFILE[sec_key]
    
    profile_text = extract_docx_profile_text(dominant_name, secondary_name, "funcionario")
    
    if profile_text:
        details = {
            "forcas": profile_text,
            "riscos": "",
            "recomendacoes": "",
            "profile_text": profile_text
        }
    else:
        details = PROFILES_DB.get(dominant_name, {}).get(secondary_name, {
            "forcas": f"Combinacao de {dominant_name} e {secondary_name}.",
            "riscos": "Necessidade de vigilia sobre dinamicas da equipe.",
            "recomendacoes": "Agende mentoria personalizada."
        })
    
    short_block_sums = {}
    for full_name, value in archetype_blocks.items():
        short_name = EMPLOYEE_BLOCK_SHORT_NAMES.get(full_name, full_name)
        short_block_sums[short_name] = value
    
    bion_role = classify_employee_bion_role(short_block_sums)
    
    cloninger_block = "Bloco 8 – Temperamento e Carater (Cloninger Funcionario)"
    cloninger_scores = {}
    for code, info in EMPLOYEE_CLONINGER_SUBDIMENSIONS.items():
        total = sum(responses.get(f"{cloninger_block}_{q}", 3) for q in info["questions"])
        cloninger_scores[code] = {"name": info["name"], "score": total, "max": 10}
    
    return dominant_name, secondary_name, details, bion_role, block_sums, cloninger_scores

# Render Paywall Box
def render_paywall():
    st.markdown("""
        <div class="paywall-box">
            <div class="paywall-icon">🔒</div>
            <div class="paywall-title">Conteúdo Exclusivo para Alunos</div>
            <div class="paywall-text">
                Liberação apenas após confirmação de pagamento.<br>
                Entre em contato para adquirir seu acesso.
            </div>
        </div>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Já sou aluno - Entrar", use_container_width=True, key="paywall_login"):
            st.session_state.page = "Login"
            st.rerun()
    with col2:
        st.markdown("""
            <a href="mailto:contato@lpshub.com.br" class="cta-button" style="display: block; text-align: center; background-color: #d19f09; color: #18738c;">
                Adquirir Acesso
            </a>
        """, unsafe_allow_html=True)

# PDF Generation Functions with LPS Branding
def generate_individual_pdf(employee_data, manager_name=""):
    """Generate individual employee PDF report with LPS branding."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.75*inch, bottomMargin=0.75*inch)
    
    # Colors
    navy_blue = HexColor('#18738c')
    gold = HexColor('#d19f09')
    light_gray = HexColor('#F5F5F5')
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'LPSTitle',
        parent=styles['Title'],
        fontName='Helvetica-Bold',
        fontSize=18,
        textColor=navy_blue,
        alignment=TA_CENTER,
        spaceAfter=20
    )
    subtitle_style = ParagraphStyle(
        'LPSSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=14,
        textColor=navy_blue,
        spaceAfter=10,
        spaceBefore=15
    )
    body_style = ParagraphStyle(
        'LPSBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=11,
        textColor=HexColor('#333333'),
        alignment=TA_JUSTIFY,
        spaceAfter=8
    )
    
    elements = []
    
    # Title
    elements.append(Paragraph("Liderança Psicanalítica - Laudo Individual", title_style))
    elements.append(Spacer(1, 20))
    
    # Employee Info Table
    name = employee_data.get('name', 'Não informado')
    email = employee_data.get('email', 'Não informado')
    profile_dom = employee_data.get('profile_dominant', 'Não realizado')
    profile_sec = employee_data.get('profile_secondary', 'Não realizado')
    bion_role = employee_data.get('bion_role', 'Não classificado')
    
    info_data = [
        ['Nome:', name],
        ['Email:', email],
        ['Perfil Dominante:', profile_dom],
        ['Perfil Secundário:', profile_sec],
        ['Papel de Bion:', bion_role]
    ]
    
    info_table = Table(info_data, colWidths=[2*inch, 4*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), light_gray),
        ('TEXTCOLOR', (0, 0), (0, -1), navy_blue),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, navy_blue),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 20))
    
    # Profile Details
    profile_details = employee_data.get('profile_details', {})
    if profile_details:
        elements.append(Paragraph("Análise do Perfil", subtitle_style))
        
        if 'forcas' in profile_details:
            elements.append(Paragraph("<b>Forças:</b>", body_style))
            elements.append(Paragraph(profile_details['forcas'], body_style))
        
        if 'riscos' in profile_details:
            elements.append(Paragraph("<b>Riscos:</b>", body_style))
            elements.append(Paragraph(profile_details['riscos'], body_style))
        
        if 'recomendacoes' in profile_details:
            elements.append(Paragraph("<b>Recomendações:</b>", body_style))
            elements.append(Paragraph(profile_details['recomendacoes'], body_style))
    
    # Footer
    elements.append(Spacer(1, 30))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=9,
        textColor=HexColor('#666666'),
        alignment=TA_CENTER
    )
    elements.append(Paragraph(f"Relatório gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}", footer_style))
    elements.append(Paragraph("Liderança Psicanalítica - Todos os direitos reservados", footer_style))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_team_pdf(employees_list, manager_name=""):
    """Generate team PDF report with all employees."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.75*inch, bottomMargin=0.75*inch)
    
    # Colors
    navy_blue = HexColor('#18738c')
    gold = HexColor('#d19f09')
    light_gray = HexColor('#F5F5F5')
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'LPSTitle',
        parent=styles['Title'],
        fontName='Helvetica-Bold',
        fontSize=18,
        textColor=navy_blue,
        alignment=TA_CENTER,
        spaceAfter=20
    )
    subtitle_style = ParagraphStyle(
        'LPSSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=14,
        textColor=navy_blue,
        spaceAfter=10,
        spaceBefore=15
    )
    
    elements = []
    
    # Title
    elements.append(Paragraph("Liderança Psicanalítica - Relatório da Equipe", title_style))
    if manager_name:
        elements.append(Paragraph(f"Gestor: {manager_name}", ParagraphStyle(
            'ManagerName',
            parent=styles['Normal'],
            fontSize=12,
            textColor=navy_blue,
            alignment=TA_CENTER
        )))
    elements.append(Spacer(1, 20))
    
    # Team Summary Table
    elements.append(Paragraph("Membros da Equipe", subtitle_style))
    
    # Table headers
    table_data = [['Nome', 'Perfil Dominante', 'Perfil Secundário', 'Papel de Bion']]
    
    for emp in employees_list:
        name = emp.get('name', 'Não informado')
        profile_dom = emp.get('profile_dominant', '-')
        profile_sec = emp.get('profile_secondary', '-')
        bion_role = emp.get('bion_role', '-')
        table_data.append([name, profile_dom, profile_sec, bion_role])
    
    team_table = Table(table_data, colWidths=[1.8*inch, 1.8*inch, 1.8*inch, 1.5*inch])
    team_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), navy_blue),
        ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#FFFFFF')),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BACKGROUND', (0, 1), (-1, -1), light_gray),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, navy_blue),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ]))
    elements.append(team_table)
    elements.append(Spacer(1, 20))
    
    # Bion Role Distribution
    bion_counts = {}
    for emp in employees_list:
        role = emp.get('bion_role', 'Não classificado')
        if role and role != '-':
            bion_counts[role] = bion_counts.get(role, 0) + 1
    
    if bion_counts:
        elements.append(Paragraph("Distribuição de Papéis de Bion", subtitle_style))
        dist_data = [[role, str(count)] for role, count in bion_counts.items()]
        dist_table = Table(dist_data, colWidths=[4*inch, 1*inch])
        dist_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), light_gray),
            ('TEXTCOLOR', (0, 0), (0, -1), navy_blue),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 0.5, navy_blue),
            ('ALIGN', (1, 0), (1, -1), 'CENTER'),
        ]))
        elements.append(dist_table)
    
    # Footer
    elements.append(Spacer(1, 30))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=9,
        textColor=HexColor('#666666'),
        alignment=TA_CENTER
    )
    elements.append(Paragraph(f"Relatório gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}", footer_style))
    elements.append(Paragraph("Liderança Psicanalítica - Todos os direitos reservados", footer_style))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_csv_data(employees_list):
    """Generate CSV string from employee data."""
    import csv
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Headers
    writer.writerow(['Nome', 'Email', 'Perfil Dominante', 'Perfil Secundário', 'Papel de Bion', 'Data'])
    
    for emp in employees_list:
        writer.writerow([
            emp.get('name', ''),
            emp.get('email', ''),
            emp.get('profile_dominant', ''),
            emp.get('profile_secondary', ''),
            emp.get('bion_role', ''),
            emp.get('created_at', '')
        ])
    
    return output.getvalue()

def generate_team_pdf_report(manager_name, employees):
    """Generate team PDF report from employee tuple list."""
    # Convert tuples to dict format for the PDF generator
    employees_list = []
    for emp in employees:
        if emp[10] == 1:  # completed
            employees_list.append({
                'name': emp[4] or f'Funcionario {emp[3]}',
                'email': emp[5] or '',
                'profile_dominant': emp[6] or '',
                'profile_secondary': emp[7] or '',
                'bion_role': emp[9] or ''
            })
    
    return generate_team_pdf(employees_list, manager_name).getvalue()

def generate_individual_pdf_report(employee_tuple, manager_name=""):
    """Generate individual PDF report from employee tuple."""
    employee_data = {
        'name': employee_tuple[4] or f'Funcionario {employee_tuple[3]}',
        'email': employee_tuple[5] or '',
        'profile_dominant': employee_tuple[6] or '',
        'profile_secondary': employee_tuple[7] or '',
        'bion_role': employee_tuple[9] or '',
        'profile_details': json.loads(employee_tuple[8]) if employee_tuple[8] else {}
    }
    
    return generate_individual_pdf(employee_data, manager_name).getvalue()

# GLOBAL EMPLOYEE ACCESS GUARD - Force employees to stay on EmployeeAssessment
# This runs before every page render to prevent any navigation
if is_employee_access and page != "EmployeeAssessment":
    st.session_state.page = "EmployeeAssessment"
    page = "EmployeeAssessment"

# Pages
if page == "Home":
    # Public landing page with sections
    render_sidebar_navigation()
    render_public_header()
    
    current_section = st.session_state.section
    
    # HOME SECTION - Hero
    if current_section == "home":
        BANNER1_PATH = "attached_assets/banner1_1771011645716.jpeg"
        BANNER2_PATH = "attached_assets/banner2_1771011645718.jpeg"
        if os.path.exists(BANNER1_PATH):
            st.image(BANNER1_PATH, use_container_width=True)
        if os.path.exists(BANNER2_PATH):
            st.image(BANNER2_PATH, use_container_width=True)
        
        # Video Intro (immediately after banner)
        vimeo_video("https://vimeo.com/1154882598")
        
        # Features
        st.write("")
        st.write("---")
        st.markdown("### Por Que Escolher o LPS?")
        
        feat_cols = st.columns(3)
        with feat_cols[0]:
            st.markdown("""
                <div style="text-align: center; padding: 1rem; background-color: #f8f9fa; border-radius: 10px; border: 2px solid #18738c;">
                    <div style="width: 60px; height: 60px; margin: 0 auto 1rem; background-color: #18738c; border-radius: 50%; display: flex; align-items: center; justify-content: center;">
                        <span style="color: white; font-weight: bold; font-size: 1.5rem;">N+P</span>
                    </div>
                    <h4 style="color: #18738c;">Neurociência + Psicanálise</h4>
                    <p style="color: #666;">Metodologia unica que une ciencia do cerebro com analise profunda do comportamento.</p>
                </div>
            """, unsafe_allow_html=True)
        with feat_cols[1]:
            st.markdown("""
                <div style="text-align: center; padding: 1rem; background-color: #f8f9fa; border-radius: 10px; border: 2px solid #18738c;">
                    <div style="width: 60px; height: 60px; margin: 0 auto 1rem; background-color: #d19f09; border-radius: 50%; display: flex; align-items: center; justify-content: center;">
                        <span style="color: white; font-weight: bold; font-size: 1.2rem;">LPT</span>
                    </div>
                    <h4 style="color: #18738c;">Assessment Completo</h4>
                    <p style="color: #666;">Descubra seu perfil de lideranca e mapeie sua equipe com ferramentas exclusivas.</p>
                </div>
            """, unsafe_allow_html=True)
        with feat_cols[2]:
            st.markdown("""
                <div style="text-align: center; padding: 1rem; background-color: #f8f9fa; border-radius: 10px; border: 2px solid #18738c;">
                    <div style="width: 60px; height: 60px; margin: 0 auto 1rem; background-color: #18738c; border-radius: 50%; display: flex; align-items: center; justify-content: center;">
                        <span style="color: white; font-weight: bold; font-size: 1.2rem;">IA</span>
                    </div>
                    <h4 style="color: #18738c;">IA Consultora</h4>
                    <p style="color: #666;">Converse com a LPChat para receber insights personalizados sobre sua equipe.</p>
                </div>
            """, unsafe_allow_html=True)
        
        if not st.session_state.get('authenticated'):
            st.write("")
            cta_cols = st.columns([1, 2, 1])
            with cta_cols[1]:
                st.markdown(f"""
                    <a href="{WHATSAPP_URL}" target="_blank" class="cta-button" style="display: block; text-align: center; font-size: 1.2rem;">
                        Comprar Curso / Solicitar Orcamento
                    </a>
                """, unsafe_allow_html=True)
    
    # SOBRE SECTION
    elif current_section == "sobre":
        st.markdown('<div class="section-title">Sobre</div>', unsafe_allow_html=True)
        
        st.markdown("""
            <div style="background: white; border-radius: 12px; padding: 2rem; margin-bottom: 2rem; box-shadow: 0 2px 8px rgba(0,0,0,0.08); border-left: 4px solid #c5a059;">
                <h2 style="color: #18738c; margin-top: 0;">SOBRE A AUTORA</h2>
                <p style="color: #333; line-height: 1.8; font-size: 1.05rem;">
                    Viviane Nishiura é psicóloga formada pela Universidade Mackenzie, com mais de <strong>30 anos de experiência</strong> 
                    conectando RH corporativo, gestão de projetos e clínica psicológica. Ao longo de sua trajetória, atuou em grandes 
                    organizações e desenvolveu uma visão única sobre os conflitos corporativos, identificando-os como manifestações de 
                    dinâmicas inconscientes que impactam diretamente a performance das equipes.
                </p>
                <p style="color: #333; line-height: 1.8; font-size: 1.05rem;">
                    Sua formação integra psicanálise, neurociência aplicada e gestão estratégica de pessoas, o que lhe permite traduzir 
                    conceitos complexos da psicologia profunda em ferramentas práticas para o dia a dia da liderança. Atualmente, Viviane 
                    lidera o <strong>LPS - Líder Psicanalítico</strong>, um programa inovador que ajuda gestores a sustentarem autoridade 
                    e limites com saúde mental, promovendo ambientes de trabalho mais conscientes e produtivos.
                </p>
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
            <div style="background: white; border-radius: 12px; padding: 2rem; margin-bottom: 2rem; box-shadow: 0 2px 8px rgba(0,0,0,0.08); border-left: 4px solid #18738c;">
                <h2 style="color: #18738c; margin-top: 0;">SOBRE O LPS</h2>
                <p style="color: #333; line-height: 1.8; font-size: 1.05rem;">
                    O <strong>LPS (Líder Psicanalítico)</strong> é um modelo de intervenção estruturado para o mundo corporativo, 
                    baseado no princípio de que organizações ativam regressões emocionais que afetam a tomada de decisão, os 
                    relacionamentos interpessoais e a produtividade das equipes.
                </p>
                <p style="color: #333; line-height: 1.8; font-size: 1.05rem;">
                    A metodologia integra <strong>psicanálise e neurociência</strong>, oferecendo aos líderes uma compreensão profunda 
                    dos mecanismos inconscientes que operam nas dinâmicas de grupo. O programa capacita gestores a:
                </p>
                <ul style="color: #333; line-height: 2; font-size: 1.05rem;">
                    <li><strong>Leitura Psíquica:</strong> Entender o funcionamento invisível das equipes e identificar padrões de comportamento inconscientes.</li>
                    <li><strong>Sustentação de Autoridade:</strong> Tomar decisões estratégicas sem sobrecarga emocional, mantendo limites saudáveis.</li>
                    <li><strong>Performance:</strong> Alinhar o comportamento à tarefa organizacional, reduzindo conflitos e aumentando a coesão do time.</li>
                    <li><strong>Autoconhecimento:</strong> Desenvolver a capacidade de reconhecer as próprias defesas e pontos cegos na liderança.</li>
                </ul>
                <p style="color: #333; line-height: 1.8; font-size: 1.05rem;">
                    Através de assessments exclusivos, formação teórica e mentoria individual, o LPS transforma a maneira como líderes 
                    se relacionam com suas equipes, promovendo ambientes de trabalho mais saudáveis e resultados sustentáveis.
                </p>
            </div>
        """, unsafe_allow_html=True)
        
        if st.button("VOLTAR PARA HOME", key="footer_back_home_sobre", use_container_width=True):
            st.session_state.section = "home"
            st.session_state.page = "Home"
            st.rerun()

    
    # CURSO SECTION - Module Cards with Paywall
    elif current_section == "curso":
        st.markdown('<div class="section-title">Programa de Formação LPS</div>', unsafe_allow_html=True)
        st.markdown('<p style="text-align: center; color: #000000; font-size: 1.1rem; margin-bottom: 2rem;">O essencial da psicologia diretamente aplicado no desempenho da Liderança</p>', unsafe_allow_html=True)
        
        # Premium Module Card CSS
        st.markdown("""
            <style>
            .premium-module-card {
                background-color: #18738c;
                border-radius: 12px;
                padding: 1.5rem;
                margin-bottom: 1rem;
                min-height: 280px;
                box-shadow: 0 4px 15px rgba(0,0,0,0.15);
                transition: transform 0.2s ease, box-shadow 0.2s ease;
            }
            .premium-module-card:hover {
                transform: translateY(-5px);
                box-shadow: 0 8px 25px rgba(0,0,0,0.25);
            }
            .premium-module-icon {
                font-size: 2.5rem;
                text-align: center;
                margin-bottom: 0.75rem;
            }
            .premium-module-title {
                color: #d19f09;
                font-size: 1.15rem;
                font-weight: bold;
                text-align: center;
                margin-bottom: 0.75rem;
            }
            .premium-module-desc {
                color: #FFFFFF;
                font-size: 0.9rem;
                line-height: 1.5;
                text-align: center;
            }
            </style>
        """, unsafe_allow_html=True)
        
        # Module Cards Grid - 2 rows of 4 for 8 modules
        row1 = st.columns(4)
        for idx, col in enumerate(row1):
            if idx < len(MODULES_DATA):
                mod = MODULES_DATA[idx]
                with col:
                    st.markdown(f"""
                        <div class="premium-module-card">
                            <div class="premium-module-icon">{mod['icon']}</div>
                            <div class="premium-module-title">{mod['title']}</div>
                            <div class="premium-module-desc">{mod['description']}</div>
                        </div>
                    """, unsafe_allow_html=True)
        
        st.write("")
        row2 = st.columns(4)
        for idx, col in enumerate(row2):
            mod_idx = idx + 4
            if mod_idx < len(MODULES_DATA):
                mod = MODULES_DATA[mod_idx]
                with col:
                    st.markdown(f"""
                        <div class="premium-module-card">
                            <div class="premium-module-icon">{mod['icon']}</div>
                            <div class="premium-module-title">{mod['title']}</div>
                            <div class="premium-module-desc">{mod['description']}</div>
                        </div>
                    """, unsafe_allow_html=True)
        
        # Paywall Modal for non-logged users
        if st.session_state.show_login_modal:
            st.write("")
            st.markdown("""
                <div style="background-color: #fff3cd; padding: 2rem; border-radius: 10px; border-left: 4px solid #d19f09; margin: 1rem 0;">
                    <h3 style="color: #18738c; margin-top: 0;">Conteúdo Exclusivo para Alunos</h3>
                    <p style="color: #333;">O acesso ao curso completo é liberado após a confirmação do pagamento. Entre em contato pelo e-mail contato@lpshub.com.br para adquirir seu acesso.</p>
                </div>
            """, unsafe_allow_html=True)
            if st.button("Fechar", key="close_modal"):
                st.session_state.show_login_modal = False
                st.rerun()
        
        if not st.session_state.get('authenticated'):
            st.write("---")
            cta_cols = st.columns([1, 2, 1])
            with cta_cols[1]:
                st.markdown(f"""
                    <a href="{WHATSAPP_URL}" target="_blank" class="cta-button" style="display: block; text-align: center; font-size: 1.1rem; background-color: #d19f09; color: #18738c;">
                        Comprar Curso Completo
                    </a>
                """, unsafe_allow_html=True)
    
    # LPSTEST SECTION - Preview with Paywall
    elif current_section == "lpstest":
        st.markdown('<div class="section-title">LPTest - Assessment de Liderança</div>', unsafe_allow_html=True)
        
        st.markdown("""
            <div class="about-card">
                <h3 style="color: #18738c;">Descubra Seu Perfil de Liderança</h3>
                <p>O LPTest é um assessment exclusivo, desenvolvido com bases teóricas na psicologia e neurociência 
                a partir de avaliações comportamentais e de personalidade. As questões mapeiam seu perfil através 
                de <strong>7 dimensões psicanalíticas + 1 bloco de temperamento</strong>:</p>
                <ul style="color: #666;">
                    <li><strong>Autoridade/Superego</strong> - Pressão interna e exigência consigo mesmo</li>
                    <li><strong>Contenção Emocional</strong> - Empatia e criação de segurança no grupo</li>
                    <li><strong>Narcisismo/Motivação</strong> - Sua relação com reconhecimento e validação</li>
                    <li><strong>Estrutura/Controle</strong> - Necessidade de ordem e tolerância à ambiguidade</li>
                    <li><strong>Dinâmicas Relacionais</strong> - Transferência e contratransferência com a equipe</li>
                    <li><strong>Autoconsciência</strong> - Mentalização e defesas do ego</li>
                    <li><strong>Execução/Urgência</strong> - Foco em ação e resultados</li>
                </ul>
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
            <div style="text-align: center; padding: 2rem;">
                <div style="font-size: 4rem;">📊</div>
                <h3 style="color: #18738c;">Receba Seu Perfil Híbrido</h3>
                <p style="color: #666;">Ao completar o assessment, você recebe um relatório com seu perfil dominante, 
                perfil secundário e papel grupal segundo Bion.</p>
            </div>
        """, unsafe_allow_html=True)
        
        if st.session_state.authenticated:
            if st.button("Fazer o LPTest Agora", use_container_width=True, type="primary"):
                st.session_state.page = "LPTest"
                st.rerun()
        else:
            render_paywall()
    
    # LPSCHAT SECTION - Preview with Paywall
    elif current_section == "lpschat":
        st.markdown('<div class="section-title">LPChat - Consultor de IA</div>', unsafe_allow_html=True)
        
        st.markdown("""
            <div class="about-card">
                <h3 style="color: #18738c;">Sua Consultora Psicanalítica 24/7</h3>
                <p>O LPChat é uma inteligência artificial treinada com os conceitos da metodologia LPS. 
                Ela tem acesso ao seu perfil e da sua equipe, oferecendo:</p>
                <ul style="color: #666;">
                    <li>Análise das dinâmicas grupais da sua equipe</li>
                    <li>Identificação de padrões de transferência</li>
                    <li>Sugestões de intervenções baseadas em Bion</li>
                    <li>Orientações para gestão de conflitos</li>
                </ul>
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
            <div style="background-color: #F5F5F5; padding: 1.5rem; border-radius: 10px; margin: 1rem 0;">
                <p style="color: #666; font-style: italic;">"Minha equipe está em conflito constante. O João parece 
                sempre ser culpado por tudo. Como posso intervir?"</p>
                <p style="color: #18738c; margin-top: 1rem;"><strong>LPChat:</strong> Pelo que descreve, João pode 
                estar assumindo o papel de <em>Bode Expiatório</em> do grupo - uma dinâmica comum quando há ansiedade 
                não processada. Sugiro focar na <em>Tarefa Real</em> para resgatar a racionalidade...</p>
            </div>
        """, unsafe_allow_html=True)
        
        if st.session_state.authenticated:
            if st.button("Acessar LPChat", use_container_width=True, type="primary"):
                st.session_state.page = "LPChat"
                st.rerun()
        else:
            render_paywall()
    
    # MENTORIA SECTION
    elif current_section == "mentoria":
        st.markdown('<div class="section-title">Mentoria Individual</div>', unsafe_allow_html=True)
        
        st.markdown("""
            <div class="about-card">
                <h3 style="color: #18738c;">Acompanhamento Personalizado</h3>
                <p>Sessões individuais com consultor sênior para aprofundar seu desenvolvimento como líder psicanalítico:</p>
                <ul style="color: #666;">
                    <li>Análise do seu perfil LPTest</li>
                    <li>Supervisão de casos da sua equipe</li>
                    <li>Desenvolvimento de estratégias personalizadas</li>
                    <li>Acompanhamento do seu progresso</li>
                </ul>
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"""
            <div style="text-align: center; margin-top: 2rem;">
                <a href="{WHATSAPP_URL}" target="_blank" class="cta-button" style="font-size: 1.2rem;">
                    📅 Agendar Mentoria
                </a>
            </div>
        """, unsafe_allow_html=True)
    
    # TRILHA LPS SECTION
    elif current_section == "soluções":
        st.markdown('<div class="section-title">Trilha LPS</div>', unsafe_allow_html=True)
        st.markdown('<p style="text-align: center; color: #666; font-size: 1.1rem; margin-bottom: 2rem;">Conheca as 4 etapas do programa de desenvolvimento de lideres</p>', unsafe_allow_html=True)
        
        sol_cols = st.columns(2)
        with sol_cols[0]:
            st.markdown("""
                <div class="solution-card">
                    <div class="solution-title">1 - O Curso LPS</div>
                    <p>Formação teórica completa com 8 módulos sobre psicanálise aplicada à liderança, neurociência organizacional e dinâmicas grupais de Bion.</p>
                </div>
            """, unsafe_allow_html=True)
        with sol_cols[1]:
            st.markdown("""
                <div class="solution-card">
                    <div class="solution-title">2 - LPTest</div>
                    <p>Assessment exclusivo com 70 questoes que mapeia seu perfil de lideranca atraves de dimensoes psicanalíticas, identificando seu arquetipo dominante e papel grupal.</p>
                </div>
            """, unsafe_allow_html=True)
        
        st.write("")
        sol_cols2 = st.columns(2)
        with sol_cols2[0]:
            st.markdown("""
                <div class="solution-card">
                    <div class="solution-title">3 - Mentoria</div>
                    <p>Sessões individuais com consultora sênior para aprofundar seu desenvolvimento, analisar dinâmicas da sua equipe e criar estratégias personalizadas de intervenção.</p>
                </div>
            """, unsafe_allow_html=True)
        with sol_cols2[1]:
            st.markdown("""
                <div class="solution-card">
                    <div class="solution-title">4 - LPChat</div>
                    <p>Consultor virtual de IA treinado com a metodologia LPS, que analisa o perfil da sua equipe e oferece insights psicanalíticos e de neurociência em tempo real.</p>
                </div>
            """, unsafe_allow_html=True)
        
        st.markdown(f"""
            <div style="text-align: center; margin-top: 2rem;">
                <a href="{WHATSAPP_URL}" target="_blank" class="cta-button" style="font-size: 1.2rem;">
                    Iniciar Minha Trilha LPS
                </a>
            </div>
        """, unsafe_allow_html=True)
    
    # INSIGHTS/BLOG SECTION (SEO)
    elif current_section == "insights":
        st.markdown('<div class="section-title">Insights de Liderança</div>', unsafe_allow_html=True)
        st.markdown("""
            <p style="text-align: center; color: #666; margin-bottom: 2rem;">
                Artigos e conteudos sobre Psicanálise, Neurociência e Liderança Consciente
            </p>
        """, unsafe_allow_html=True)
        
        # Featured Articles
        blog_posts = [
            {
                "title": "Os 5 Arquetipos Inconscientes que Todo Lider Possui",
                "excerpt": "Descubra como padroes ocultos influenciam suas decisoes e o comportamento da sua equipe. Aprenda a identificar seu arquetipo dominante.",
                "category": "Psicanálise",
                "date": "Janeiro 2025"
            },
            {
                "title": "Neurociência do Estresse: Como o Cortisol Afeta sua Liderança",
                "excerpt": "Entenda os mecanismos cerebrais por tras do estresse cronico e como gestores podem criar ambientes que promovem produtividade.",
                "category": "Neurociência",
                "date": "Janeiro 2025"
            },
            {
                "title": "Transferencia e Contratransferencia no Ambiente Corporativo",
                "excerpt": "Por que alguns funcionarios 'ativam' reacoes intensas em voce? A resposta esta nas dinamicas inconscientes de transferencia.",
                "category": "Psicanálise",
                "date": "Dezembro 2024"
            },
            {
                "title": "Como Reduzir Turnover com Mapeamento de Perfis",
                "excerpt": "Estudo de caso: empresa reduziu rotatividade em 40% apos aplicar LPTest e realocar funcionarios por perfil.",
                "category": "Gestão",
                "date": "Dezembro 2024"
            },
            {
                "title": "Os Papeis de Bion: Entenda as Dinamicas Ocultas do Seu Time",
                "excerpt": "Porta-voz, Bode Expiatorio, Lider de Luta-Fuga... Descubra quem esta exercendo cada papel na sua equipe.",
                "category": "Dinamica Grupal",
                "date": "Novembro 2024"
            },
            {
                "title": "Neuronios-Espelho e Empatia: A Base Neurologica da Liderança",
                "excerpt": "Como seu cerebro 'espelha' emocoes da equipe e por que isso e crucial para liderar com autenticidade.",
                "category": "Neurociência",
                "date": "Novembro 2024"
            }
        ]
        
        cols = st.columns(2)  # Initialize columns outside loop
        for i, post in enumerate(blog_posts):
            if i % 2 == 0:
                cols = st.columns(2)
            
            with cols[i % 2]:
                category_color = {
                    "Psicanálise": "#18738c",
                    "Neurociência": "#28a745",
                    "Gestão": "#d19f09",
                    "Dinamica Grupal": "#6c757d"
                }.get(post["category"], "#18738c")
                
                st.markdown(f"""
                    <div style="background: white; border-radius: 10px; padding: 1.5rem; margin-bottom: 1rem; box-shadow: 0 2px 8px rgba(0,0,0,0.1); border-left: 4px solid {category_color};">
                        <span style="background: {category_color}; color: white; padding: 3px 10px; border-radius: 15px; font-size: 0.75rem; margin-bottom: 0.5rem; display: inline-block;">{post['category']}</span>
                        <h3 style="color: #18738c; margin: 0.5rem 0; font-size: 1.1rem;">{post['title']}</h3>
                        <p style="color: #666; font-size: 0.9rem; margin-bottom: 0.5rem;">{post['excerpt']}</p>
                        <span style="color: #999; font-size: 0.8rem;">{post['date']}</span>
                    </div>
                """, unsafe_allow_html=True)
        
        st.markdown(f"""
            <div style="text-align: center; margin-top: 2rem; padding: 2rem; background: linear-gradient(135deg, #18738c 0%, #1a4f7a 100%); border-radius: 15px;">
                <h3 style="color: #d19f09; margin-bottom: 1rem;">Receba Conteudos Exclusivos</h3>
                <p style="color: white; margin-bottom: 1.5rem;">Insights exclusivos sobre lideranca psicanalítica direto no seu WhatsApp</p>
                <a href="{WHATSAPP_URL}" target="_blank" style="display: inline-block; background-color: #25D366; color: #000000; padding: 12px 30px; border-radius: 25px; text-decoration: none; font-weight: bold;">
                    Receber Insights
                </a>
                <div style="margin-top: 1.5rem; display: flex; justify-content: center; gap: 1.5rem;">
                    <a href="https://www.linkedin.com/company/lpshub" target="_blank" style="color: white; text-decoration: none; font-weight: bold;">LinkedIn</a>
                    <a href="https://www.youtube.com/@LPSHubLiderança" target="_blank" style="color: white; text-decoration: none; font-weight: bold;">YouTube</a>
                </div>
            </div>
        """, unsafe_allow_html=True)
    
    # CONTATO SECTION
    elif current_section == "contato":
        st.markdown('<div class="section-title">Entre em Contato</div>', unsafe_allow_html=True)
        
        st.markdown("""
            <div style="text-align: center; padding: 2rem;">
                <div style="font-size: 4rem;">📱</div>
                <h3 style="color: #18738c;">Fale Diretamente Conosco</h3>
                <p style="color: #666; font-size: 1.1rem;">
                    Tire suas dúvidas, solicite orçamentos ou agende sua mentoria.<br>
                    Atendimento personalizado via WhatsApp.
                </p>
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"""
            <div style="text-align: center;">
                <a href="{WHATSAPP_URL}" target="_blank" style="display: inline-block; background-color: #25D366; color: #000000; padding: 20px 50px; border-radius: 50px; font-weight: bold; font-size: 1.3rem; text-decoration: none;">
                    💬 Iniciar Conversa no WhatsApp
                </a>
            </div>
        """, unsafe_allow_html=True)
        
        st.write("")
        st.markdown("""
            <div style="text-align: center; color: #666; margin-top: 2rem;">
                <p><strong>E-mail:</strong> contato@lpshub.com.br</p>
                <p><strong>Instagram:</strong> @lpshublider</p>
                <div style="margin-top: 1rem; display: flex; justify-content: center; gap: 1.5rem;">
                    <a href="https://www.linkedin.com/company/lpshub" target="_blank" style="display: inline-flex; align-items: center; gap: 0.5rem; color: #0A66C2; text-decoration: none; font-weight: bold; font-size: 1rem;">
                        <svg width="24" height="24" viewBox="0 0 24 24" fill="#0A66C2"><path d="M20.447 20.452h-3.554v-5.569c0-1.328-.027-3.037-1.852-3.037-1.853 0-2.136 1.445-2.136 2.939v5.667H9.351V9h3.414v1.561h.046c.477-.9 1.637-1.85 3.37-1.85 3.601 0 4.267 2.37 4.267 5.455v6.286zM5.337 7.433c-1.144 0-2.063-.926-2.063-2.065 0-1.138.92-2.063 2.063-2.063 1.14 0 2.064.925 2.064 2.063 0 1.139-.925 2.065-2.064 2.065zm1.782 13.019H3.555V9h3.564v11.452zM22.225 0H1.771C.792 0 0 .774 0 1.729v20.542C0 23.227.792 24 1.771 24h20.451C23.2 24 24 23.227 24 22.271V1.729C24 .774 23.2 0 22.222 0h.003z"/></svg>
                        LinkedIn
                    </a>
                    <a href="https://www.youtube.com/@LPSHubLiderança" target="_blank" style="display: inline-flex; align-items: center; gap: 0.5rem; color: #FF0000; text-decoration: none; font-weight: bold; font-size: 1rem;">
                        <svg width="24" height="24" viewBox="0 0 24 24" fill="#FF0000"><path d="M23.498 6.186a3.016 3.016 0 0 0-2.122-2.136C19.505 3.545 12 3.545 12 3.545s-7.505 0-9.377.505A3.017 3.017 0 0 0 .502 6.186C0 8.07 0 12 0 12s0 3.93.502 5.814a3.016 3.016 0 0 0 2.122 2.136c1.871.505 9.376.505 9.376.505s7.505 0 9.377-.505a3.015 3.015 0 0 0 2.122-2.136C24 15.93 24 12 24 12s0-3.93-.502-5.814zM9.545 15.568V8.432L15.818 12l-6.273 3.568z"/></svg>
                        YouTube
                    </a>
                </div>
            </div>
        """, unsafe_allow_html=True)

elif page == "Login":
    render_sidebar_navigation()
    render_login_page()

elif page == "Dashboard":
    # Authenticated user dashboard
    if not st.session_state.authenticated:
        st.session_state.page = "Home"
        st.rerun()
    
    render_sidebar_navigation()
    render_public_header()
    
    # Get manager data
    manager_data = get_manager_by_user(st.session_state.user['id'])
    user_id = st.session_state.user['id']
    manager_id = manager_data['id'] if manager_data else None
    
    if st.session_state.get('show_welcome'):
        user_name = st.session_state.user['name']
        user_email = st.session_state.user.get('email', '')
        if user_email == 'viviane.lps':
            welcome_msg = "Bem-vinda, Viviane. Sua plataforma de Liderança Psicanalítica está pronta."
        else:
            welcome_msg = f"Bem-vindo(a), {user_name}. Sua plataforma de Liderança Psicanalítica está pronta."
        st.markdown(f"""
            <div style="background: linear-gradient(135deg, #18738c 0%, #1a4f7a 100%); color: white; padding: 1.5rem 2rem; border-radius: 12px; margin-bottom: 1.5rem; text-align: center;">
                <p style="font-family: 'Open Sans', sans-serif; font-size: 1.25rem; margin: 0; font-weight: 400;">{welcome_msg}</p>
            </div>
        """, unsafe_allow_html=True)
        st.session_state.show_welcome = False
    else:
        st.markdown(f"<h2 style='color: #18738c;'>Bem-vindo(a), {st.session_state.user['name']}!</h2>", unsafe_allow_html=True)
    
    if is_user_admin(user_id):
        if not st.session_state.get('admin_test_token'):
            st.session_state.admin_test_token = create_invite_link("equipe", user_id)
        base_url = get_app_url()
        test_link = f"{base_url}/?tipo=equipe&ref={st.session_state.admin_test_token}"
        st.info("**Link de teste (Admin):** Copie o link abaixo para testar o fluxo de convite de funcionário.")
        st.text_input("Link de convite para teste:", value=test_link, key="admin_test_link", disabled=False)
        st.caption("Use este link em uma aba anônima para simular o acesso de um funcionário.")
    
    # Dashboard CSS
    st.markdown("""
        <style>
        .dashboard-card {
            background: white;
            border-radius: 12px;
            padding: 1.5rem;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            border-left: 4px solid #d19f09;
            margin-bottom: 1rem;
        }
        .dashboard-card h3 {
            color: #18738c;
            margin: 0 0 1rem 0;
            font-size: 1.1rem;
        }
        .progress-module {
            display: flex;
            align-items: center;
            margin-bottom: 0.5rem;
        }
        .progress-module-name {
            flex: 1;
            font-size: 0.9rem;
            color: #333;
        }
        .progress-module-bar {
            flex: 2;
            background: #e0e0e0;
            border-radius: 10px;
            height: 8px;
            margin: 0 10px;
        }
        .progress-module-fill {
            background: linear-gradient(90deg, #18738c, #d19f09);
            height: 100%;
            border-radius: 10px;
            transition: width 0.3s ease;
        }
        .progress-module-percent {
            font-size: 0.85rem;
            color: #666;
            min-width: 40px;
            text-align: right;
        }
        .stat-box {
            text-align: center;
            padding: 1rem;
            background: #f8f9fa;
            border-radius: 8px;
        }
        .stat-number {
            font-size: 2rem;
            font-weight: bold;
            color: #18738c;
        }
        .stat-label {
            font-size: 0.85rem;
            color: #666;
        }
        .insight-item {
            padding: 0.75rem;
            margin-bottom: 0.5rem;
            border-radius: 8px;
            font-size: 0.9rem;
        }
        .insight-warning {
            background: #fff3cd;
            border-left: 3px solid #ffc107;
        }
        .insight-info {
            background: #cce5ff;
            border-left: 3px solid #0d6efd;
        }
        .insight-alert {
            background: #f8d7da;
            border-left: 3px solid #dc3545;
        }
        .insight-success {
            background: #d4edda;
            border-left: 3px solid #28a745;
        }
        .mentoring-btn {
            display: inline-block;
            background: linear-gradient(135deg, #18738c, #1a5490);
            color: white;
            padding: 15px 30px;
            border-radius: 25px;
            text-decoration: none;
            font-weight: bold;
            text-align: center;
            margin-top: 1rem;
        }
        .mentoring-btn:hover {
            opacity: 0.9;
            color: white;
        }
        </style>
    """, unsafe_allow_html=True)
    
    # Main Dashboard Grid
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # Course Progress Card
        st.markdown("<div class='dashboard-card'><h3>Progresso do Curso</h3>", unsafe_allow_html=True)
        
        module_status = get_module_completion_status(user_id)
        total_completed = 0
        total_lessons = 0
        
        for mod_id, status in module_status.items():
            total_completed += status['completed']
            total_lessons += status['total']
            percentage = int(status['percentage'])
            st.markdown(f"""
                <div class='progress-module'>
                    <span class='progress-module-name'>{status['name'][:25]}...</span>
                    <div class='progress-module-bar'>
                        <div class='progress-module-fill' style='width: {percentage}%'></div>
                    </div>
                    <span class='progress-module-percent'>{percentage}%</span>
                </div>
            """, unsafe_allow_html=True)
        
        overall_progress = (total_completed / total_lessons * 100) if total_lessons > 0 else 0
        st.markdown(f"<p style='text-align: center; margin-top: 1rem; color: #18738c; font-weight: bold;'>Progresso Total: {overall_progress:.0f}%</p></div>", unsafe_allow_html=True)
        
        # Radar Chart Card - Manager Profile
        manager_profile = get_manager_profile_by_user(user_id)
        if manager_profile:
            st.markdown("<div class='dashboard-card'><h3>Seu Perfil de Liderança</h3>", unsafe_allow_html=True)
            
            # Try to get block sums from assessment responses
            block_sums = get_assessment_block_sums(user_id, "manager")
            if block_sums:
                profile_name = f"{manager_profile['dominant']} + {manager_profile['secondary']}"
                radar_chart = generate_radar_chart(block_sums, profile_name)
                if radar_chart:
                    st.image(radar_chart, use_container_width=True)
                    st.markdown(f"""
                        <div style='text-align: center; padding: 1rem; background: #f8f9fa; border-radius: 8px; margin-top: 1rem;'>
                            <strong style='color: #18738c;'>Perfil Dominante:</strong> {manager_profile['dominant']}<br>
                            <strong style='color: #18738c;'>Perfil Secundario:</strong> {manager_profile['secondary']}
                        </div>
                    """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                    <div style='text-align: center; padding: 1rem;'>
                        <strong style='color: #18738c;'>{manager_profile['dominant']} + {manager_profile['secondary']}</strong>
                        <p style='color: #666; font-size: 0.9rem; margin-top: 0.5rem;'>Refaca o LPTest para visualizar seu grafico radar</p>
                    </div>
                """, unsafe_allow_html=True)
            
            st.markdown("</div>", unsafe_allow_html=True)
        
    
    with col2:
        # Assessment Stats Card
        if manager_id:
            st.markdown("<div class='dashboard-card'><h3>Gestão de LPTest</h3>", unsafe_allow_html=True)
            
            stats = get_assessment_stats(manager_id)
            
            stat_cols = st.columns(2)
            with stat_cols[0]:
                st.markdown(f"""
                    <div class='stat-box'>
                        <div class='stat-number'>{stats['applied']}</div>
                        <div class='stat-label'>Aplicados</div>
                    </div>
                """, unsafe_allow_html=True)
            with stat_cols[1]:
                st.markdown(f"""
                    <div class='stat-box'>
                        <div class='stat-number'>{stats['remaining']}</div>
                        <div class='stat-label'>Restantes</div>
                    </div>
                """, unsafe_allow_html=True)
            
            st.markdown("</div>", unsafe_allow_html=True)
        
        # Mentoring Card
        st.markdown("<div class='dashboard-card'><h3>Mentoria Individual</h3>", unsafe_allow_html=True)
        st.markdown("<p style='font-size: 0.9rem; color: #666;'>Agende uma sessão exclusiva com consultor sênior para aprofundar seus insights de liderança.</p>", unsafe_allow_html=True)
        st.markdown(f"""
            <a href='{WHATSAPP_URL}' target='_blank' class='mentoring-btn' data-testid='button-agendar-mentoria'>
                Agendar Mentoria
            </a>
        """, unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Team Chart Preview (if team is mapped)
    if manager_id:
        employees = get_secure_manager_employees(user_id, manager_id)
        completed_employees = [e for e in employees if e[10] == 1]
        if len(completed_employees) >= 2:
            st.markdown("<div class='dashboard-card'><h3>Mapeamento da Equipe</h3>", unsafe_allow_html=True)
            manager_name = st.session_state.user['name'] if st.session_state.user else "Gestor"
            chart_data = generate_team_chart(employees, manager_name)
            if chart_data:
                st.image(chart_data, use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)
    
    st.write("---")
    
    # Quick Access Buttons
    st.markdown("<h4 style='color: #18738c;'>Acesso Rápido</h4>", unsafe_allow_html=True)
    dash_cols = st.columns(4)
    with dash_cols[0]:
        if st.button("Curso", key="btn-dash-curso", use_container_width=True):
            st.session_state.page = "LPS Curso"
            st.rerun()
    with dash_cols[1]:
        if st.button("LPTest", key="btn-dash-test", use_container_width=True):
            st.session_state.page = "LPTest"
            st.rerun()
    with dash_cols[2]:
        if st.button("Equipe", key="btn-dash-equipe", use_container_width=True):
            st.session_state.page = "TeamManagement"
            st.rerun()
    with dash_cols[3]:
        if st.button("LPChat", key="btn-dash-chat", use_container_width=True):
            st.session_state.page = "LPChat"
            st.rerun()
    
    st.write("---")
    
    # Admin section
    if is_user_admin(user_id):
        admin_col1, admin_col2, admin_col3 = st.columns([2, 1, 1])
        with admin_col1:
            if st.button("Sair", key="btn-logout", use_container_width=True):
                st.session_state.authenticated = False
                st.session_state.user = None
                st.session_state.manager_data = None
                st.session_state.page = "Home"
                st.rerun()
        with admin_col2:
            if st.button("Gestão LPS", key="btn-gestao", use_container_width=True, type="primary"):
                st.session_state.page = "GestaoLPS"
                st.rerun()
        with admin_col3:
            if st.button("Admin E-mail", key="btn-admin", use_container_width=True):
                st.session_state.page = "AdminEmail"
                st.rerun()
    else:
        if st.button("Sair", key="btn-logout", use_container_width=True):
            st.session_state.authenticated = False
            st.session_state.user = None
            st.session_state.manager_data = None
            st.session_state.page = "Home"
            st.rerun()

elif page == "LPS Curso":
    if not st.session_state.authenticated:
        st.session_state.page = "Login"
        st.rerun()
    
    user_premium = is_user_premium(st.session_state.user['id'])
    
    st.title("Programa LPS")
    
    st.markdown("""
        <style>
        [data-testid="stExpander"] summary {
            color: #555555 !important;
            font-weight: 600 !important;
        }
        [data-testid="stExpander"] {
            border-color: #555555 !important;
        }
        </style>
    """, unsafe_allow_html=True)
    
    if user_premium:
        user_id = st.session_state.user['id']
        db_progress = get_course_progress(user_id)
        if db_progress:
            st.session_state.progress = db_progress
        
        total_lessons = sum(len(m['videos']) for m in MODULES_DATA)
        completed_lessons = sum(1 for v in st.session_state.progress.values() if v)
        
        st.markdown(f"<p style='color: #18738c; font-weight: bold;'>Progresso Geral: {completed_lessons}/{total_lessons} aulas concluídas</p>", unsafe_allow_html=True)
        st.progress(completed_lessons / total_lessons if total_lessons > 0 else 0)
    
    st.write("---")
    
    for mod in MODULES_DATA:
        with st.expander(mod['name']):
            st.markdown(f"<p style='color: #18738c;'>{mod.get('description', 'Conteúdo do módulo')}</p>", unsafe_allow_html=True)
            if user_premium:
                user_id = st.session_state.user['id']
                for v_idx, v_url in enumerate(mod['videos']):
                    lesson_id = f"m{mod['id']}_v{v_idx}"
                    vimeo_video(v_url)
                    new_value = st.checkbox("Concluí esta aula", value=st.session_state.progress.get(lesson_id, False), key=lesson_id)
                    if new_value != st.session_state.progress.get(lesson_id, False):
                        st.session_state.progress[lesson_id] = new_value
                        save_course_progress(user_id, st.session_state.progress)
            else:
                render_premium_gate()

elif page == "LPTest":
    if not st.session_state.authenticated:
        st.session_state.page = "Login"
        st.rerun()
    
    user_premium = is_user_premium(st.session_state.user['id'])
    
    st.title("📝 LPTest Assessment - Seu Perfil")
    
    st.markdown("""
        <div style="background: linear-gradient(135deg, #18738c 0%, #1a4f7a 100%); color: white; padding: 1.5rem 2rem; border-radius: 12px; margin-bottom: 1.5rem;">
            <h3 style="color: #FFFFFF !important; margin-top: 0;">Assessment de Liderança em 7 Eixos</h3>
            <p style="margin-bottom: 0.5rem; color: #FFFFFF !important;">O LPTest mapeia seu perfil de liderança através de 70 questões distribuídas em 7 dimensões psicanalíticas:</p>
            <p style="font-size: 0.9rem; opacity: 0.9; color: #FFFFFF !important;">
                <strong>1.</strong> Autoridade Interna &nbsp;|&nbsp;
                <strong>2.</strong> Vínculo e Empatia &nbsp;|&nbsp;
                <strong>3.</strong> Potência Realizadora &nbsp;|&nbsp;
                <strong>4.</strong> Resiliência &nbsp;|&nbsp;
                <strong>5.</strong> Pensamento Estratégico &nbsp;|&nbsp;
                <strong>6.</strong> Expressão e Influência &nbsp;|&nbsp;
                <strong>7.</strong> Integração e Sabedoria
            </p>
        </div>
    """, unsafe_allow_html=True)
    
    if not user_premium:
        render_premium_gate()
        st.stop()
    
    saved_profile = get_manager_profile_by_user(st.session_state.user['id'])
    
    if saved_profile:
        st.success("Seu perfil já está salvo! Você pode refazer o teste a qualquer momento.")
        st.markdown(f"""
            <div class="result-card">
                <div class="profile-title">Seu Perfil Atual: {saved_profile['dominant']} + {saved_profile['secondary']}</div>
            </div>
        """, unsafe_allow_html=True)
        saved_pdf = generate_laudo_pdf(
            "", st.session_state.user.get('name', ''),
            saved_profile['dominant'], saved_profile['secondary'],
            saved_profile['details'].get('bion_role', ''),
            respondent_type="gestor"
        )
        safe_nm = (st.session_state.user.get('name', '') or "gestor").replace(" ", "_").lower()
        st.download_button(
            "BAIXAR LAUDO COMPLETO (PDF)",
            data=saved_pdf,
            file_name=f"laudo_lps_{safe_nm}.pdf",
            mime="application/pdf",
            key="dl_saved_profile_pdf",
            use_container_width=True,
            type="primary"
        )
        st.write("---")
        if st.button("Refazer LPTest"):
            st.session_state.show_test_form = True
            st.rerun()
    
    # Show form if no saved profile OR user wants to redo
    if not saved_profile or st.session_state.get('show_test_form', False):
        st.write("Responda às afirmações. (1 = Discordo Totalmente, 5 = Concordo Totalmente)")
        
        with st.form("manager_assessment"):
            responses = render_assessment_form("manager")
            submit = st.form_submit_button("Gerar Meu Perfil de Liderança")
            
            if submit and st.session_state.user:
                dominant, secondary, details, bion_role, block_sums, cloninger_scores = calculate_profile(responses)
                user_id = st.session_state.user['id']
                save_manager_profile(user_id, dominant, secondary, details)
                
                save_assessment_responses(user_id, "manager", responses)
                
                st.session_state.manager_data = get_manager_by_user(user_id)
                st.session_state.show_test_form = False
                st.session_state.assessment_results = {
                    "dominant": dominant,
                    "secondary": secondary,
                    "details": details,
                    "bion_role": bion_role,
                    "block_sums": block_sums,
                    "cloninger_scores": cloninger_scores
                }
                st.session_state.ai_laudo = None
                st.session_state.laudo_requested = True
                st.rerun()
        
        if st.session_state.assessment_results:
            res = st.session_state.assessment_results

            if res.get('block_sums'):
                profile_name = f"{res['dominant']} + {res['secondary']}"
                radar_chart = generate_radar_chart(res['block_sums'], profile_name)
                if radar_chart:
                    st.image(radar_chart, use_container_width=True)

                manager_name = st.session_state.user.get('name', '') if st.session_state.user else ''
                docx_text = extract_docx_profile_text(res['dominant'], res['secondary'], "gestor")
                if docx_text:
                    pdf_data = generate_laudo_pdf(
                        docx_text, manager_name,
                        res['dominant'], res['secondary'], res['bion_role'],
                        respondent_type="gestor"
                    )
                    safe_name = (manager_name or "gestor").replace(" ", "_").lower()
                    st.download_button(
                        "Baixar Laudo Completo em PDF",
                        data=pdf_data,
                        file_name=f"laudo_lps_{safe_name}_{datetime.now().strftime('%Y%m%d')}.pdf",
                        mime="application/pdf",
                        key="download_laudo_pdf",
                        use_container_width=True,
                        type="primary"
                    )

                    if not st.session_state.get('ai_laudo'):
                        st.session_state.ai_laudo = docx_text
                        mgr_user_id = st.session_state.user.get('id', '') if st.session_state.user else ''
                        if mgr_user_id:
                            save_laudo(mgr_user_id, "gestor", manager_name, res['dominant'], res['secondary'], res['bion_role'], docx_text)
                else:
                    st.warning("Arquivo .docx do perfil não encontrado. Verifique se os documentos estão no diretório attached_assets/.")

elif page == "TeamManagement":
    if not st.session_state.authenticated:
        st.session_state.page = "Login"
        st.rerun()
    st.title("Gestão de Equipe")
    
    manager_data = st.session_state.manager_data
    if not manager_data:
        st.error("Dados do gestor não encontrados. Por favor, faça login novamente.")
    else:
        manager_id = manager_data['id']
        user_id = st.session_state.user['id']
        manager_profile = manager_data if manager_data.get('dominant') else None
        
        # Get employees data first with multitenancy validation
        employees = get_secure_manager_employees(user_id, manager_id)
        
        # Initialize session state for showing links
        if 'show_employee_links' not in st.session_state:
            st.session_state.show_employee_links = False
        
        has_existing_links = len(employees) > 0
        
        is_admin = is_user_admin(user_id)
        if is_admin:
            tab_admin_emails, tab_resultados, tab_admin_monitoramento = st.tabs(["Cadastro de E-mails", "Resultados da Equipe", "Monitoramento"])
        else:
            tab_admin_emails, tab_resultados = st.tabs(["Cadastro de E-mails", "Resultados da Equipe"])
        
        with tab_resultados:
            st.subheader("Resultados da Equipe")
            
            if employees:
                completed_count = sum(1 for e in employees if e[10] == 1)
                manager_name = st.session_state.user['name'] if st.session_state.user else "Gestor"
                
                st.metric("Respostas Recebidas", f"{completed_count}/4")
                
                if completed_count > 0:
                    st.markdown("#### Exportar Relatorios")
                    
                    col_csv, col_pdf, col_chart = st.columns(3)
                    
                    with col_csv:
                        csv_data = "Nome,E-mail,Perfil Dominante,Perfil Secundario,Papel de Bion\n"
                        for emp in employees:
                            if emp[10] == 1:
                                emp_name = emp[4] or f'Funcionario {emp[3]}'
                                csv_data += f'"{emp_name}","{emp[5]}","{emp[6]}","{emp[7]}","{emp[9]}"\n'
                        
                        st.download_button(
                            label="Baixar CSV",
                            data=csv_data,
                            file_name="resultados_equipe_lps.csv",
                            mime="text/csv",
                            use_container_width=True
                        )
                    
                    with col_pdf:
                        pdf_data = generate_team_pdf_report(manager_name, employees)
                        st.download_button(
                            label="Baixar PDF",
                            data=pdf_data,
                            file_name="relatorio_equipe_lps.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                    
                    with col_chart:
                        chart_data = generate_team_chart(employees, manager_name)
                        if chart_data:
                            st.download_button(
                                label="Baixar Grafico",
                                data=chart_data,
                                file_name="grafico_perfil_equipe.png",
                                mime="image/png",
                                use_container_width=True
                            )
                    
                    st.write("---")
                    
                    # Display team chart preview
                    st.markdown("#### Grafico do Perfil da Equipe")
                    chart_preview = generate_team_chart(employees, manager_name)
                    if chart_preview:
                        st.image(chart_preview, use_container_width=True)
                
                if manager_profile and completed_count > 0:
                    st.markdown(f"""
                        <div style="background-color: #18738c; color: white; padding: 10px; border-radius: 8px; margin: 20px 0 10px 0;">
                            <strong>Comparacao:</strong> Seu perfil ({manager_profile['dominant']}) vs Equipe
                        </div>
                    """, unsafe_allow_html=True)
                
                # Pre-load saved laudos from database
                for emp_check in employees:
                    if emp_check[10] == 1:
                        emp_check_id = emp_check[0]
                        emp_laudo_key = f"emp_laudo_{emp_check_id}"
                        if not st.session_state.get(emp_laudo_key):
                            saved = get_laudo(emp_check_id, "funcionario")
                            if saved:
                                st.session_state[emp_laudo_key] = saved[7]
                
                # Only show completed employees in results tab
                completed_employees = [emp for emp in employees if emp[10] == 1]
                
                if completed_employees:
                    for emp in completed_employees:
                        emp_name = emp[4] or f'Funcionario {emp[3]}'
                        emp_id = emp[0]
                        safe_name = emp_name.replace(" ", "_").lower()[:20]
                        
                        emp_docx_text = extract_docx_profile_text(
                            emp[6] or "O Idealista Exigente",
                            emp[7] or "O Contenedor Empatico",
                            "funcionario"
                        )
                        
                        with st.container():
                            name_col, pdf_col = st.columns([4, 1])
                            with name_col:
                                st.markdown(f"<h4 style='color: #18738c; margin:0; padding-top:6px;'>{emp_name}</h4>", unsafe_allow_html=True)
                            with pdf_col:
                                if emp_docx_text:
                                    pdf_emp = generate_laudo_pdf(
                                        emp_docx_text, emp_name,
                                        emp[6] or "", emp[7] or "", emp[9] or "",
                                        respondent_type="funcionario"
                                    )
                                    st.download_button(
                                        label="PDF",
                                        data=pdf_emp,
                                        file_name=f"laudo_lps_{safe_name}.pdf",
                                        mime="application/pdf",
                                        key=f"download_pdf_{emp_id}"
                                    )
                            st.markdown(f"""
                                <p style='margin: 2px 0 8px 0; font-size: 0.9rem; color: #555;'>
                                    <strong>Perfil:</strong> {emp[6]} + {emp[7]} | <span class="bion-badge">{emp[9]}</span>
                                </p>
                            """, unsafe_allow_html=True)
                            st.write("---")
                else:
                    st.info("Nenhum funcionário respondeu ainda. Os resultados aparecerão aqui assim que completarem o assessment.")
            else:
                st.info("Nenhum colaborador cadastrado ainda. Use a aba 'Cadastro de E-mails' para gerar links de convite.")
        
        with tab_admin_emails:
                st.markdown("""
                    <style>
                    .gestao-card {
                        background: white; border-radius: 12px; padding: 1.5rem;
                        box-shadow: 0 2px 8px rgba(0,0,0,0.1); border-left: 4px solid #d19f09; margin-bottom: 1.5rem;
                    }
                    .gestao-card h3 { color: #18738c; margin: 0 0 1rem 0; font-size: 1.1rem; }
                    .status-badge { display: inline-block; padding: 4px 12px; border-radius: 12px; font-size: 0.8rem; font-weight: bold; }
                    .status-pendente { background: #fff3cd; color: #856404; }
                    .status-andamento { background: #cce5ff; color: #004085; }
                    .status-concluido { background: #d4edda; color: #155724; }
                    .auth-user-row { display: flex; align-items: center; justify-content: space-between; padding: 0.75rem; border-bottom: 1px solid #eee; }
                    .auth-user-row:last-child { border-bottom: none; }
                    </style>
                """, unsafe_allow_html=True)
                st.markdown("<div class='gestao-card'><h3>Cadastro de E-mails Autorizados</h3>", unsafe_allow_html=True)
                st.markdown("<p style='color: #666; font-size: 0.9rem;'>Cadastre os colaboradores e compartilhe o link de convite gerado. Ao clicar no link, eles acessam diretamente o assessment.</p>", unsafe_allow_html=True)
                
                with st.form("tm_add_auth_email_form", clear_on_submit=True):
                    col_name, col_email, col_type = st.columns([2, 2, 1])
                    with col_name:
                        auth_name = st.text_input("Nome do Colaborador", placeholder="Ex: Maria Silva", key="tm_auth_name_input")
                    with col_email:
                        auth_email = st.text_input("E-mail do Colaborador", placeholder="maria@empresa.com", key="tm_auth_email_input")
                    with col_type:
                        auth_type = st.selectbox("Tipo", ["equipe", "lider"], format_func=lambda x: "Equipe" if x == "equipe" else "Líder", key="tm_auth_type_select")
                    
                    submit_auth = st.form_submit_button("Cadastrar E-mail", use_container_width=True, type="primary")
                    if submit_auth:
                        if auth_name and auth_email:
                            success, error = add_authorized_user(auth_email, auth_name, auth_type, user_id)
                            if success:
                                st.success(f"E-mail {auth_email} cadastrado com sucesso!")
                                st.rerun()
                            else:
                                st.warning(error)
                        else:
                            st.error("Preencha o nome e o e-mail do colaborador.")
                
                st.markdown("</div>", unsafe_allow_html=True)
                
                st.markdown("<div class='gestao-card'><h3>E-mails Cadastrados</h3>", unsafe_allow_html=True)
                st.markdown("<p style='color: #666; font-size: 0.85rem;'>Copie o link ao lado de cada colaborador e envie por WhatsApp ou e-mail.</p>", unsafe_allow_html=True)
                auth_users = get_authorized_users(user_id)
                if auth_users:
                    base_url = get_app_url()
                    for au in auth_users:
                        au_id = au[0]
                        au_email = au[1]
                        au_name = au[2]
                        au_type = au[3]
                        au_status = au[4]
                        au_date = au[7][:16] if au[7] else "-"
                        au_token = au[8] if len(au) > 8 else None
                        
                        status_class = "status-concluido" if au_status == "concluido" else ("status-andamento" if au_status == "em_andamento" else "status-pendente")
                        status_label = "Concluído" if au_status == "concluido" else ("Em Andamento" if au_status == "em_andamento" else "Pendente")
                        type_label = "Equipe" if au_type == "equipe" else "Líder"
                        
                        st.markdown(f"""
                            <div style='padding: 0.6rem 0; border-bottom: 1px solid #eee;'>
                                <div style='display:flex; align-items:center; gap:8px; flex-wrap:wrap;'>
                                    <strong style='color:#18738c;'>{au_name}</strong>
                                    <span style='color:#666; font-size:0.85rem;'>{au_email}</span>
                                    <span style='color:#999; font-size:0.8rem;'>({type_label})</span>
                                    <span class='status-badge {status_class}' style='margin-left:auto;'>{status_label}</span>
                                </div>
                            </div>
                        """, unsafe_allow_html=True)
                        
                        if au_status != "concluido":
                            if au_token:
                                tipo_param = "lider" if au_type == "lider" else "equipe"
                                full_link = f"{base_url}/?tipo={tipo_param}&ref={au_token}"
                                col_link, col_del = st.columns([5, 1])
                                with col_link:
                                    st.text_input(
                                        f"Link de convite — {au_name}",
                                        value=full_link,
                                        key=f"au_link_{au_id}",
                                        label_visibility="collapsed",
                                        help="Selecione e copie este link. Envie por WhatsApp ou e-mail."
                                    )
                                with col_del:
                                    if st.button("🗑", key=f"tm_remove_auth_{au_id}", help="Remover cadastro"):
                                        delete_authorized_user(au_id)
                                        st.rerun()
                            else:
                                col_regen, col_del = st.columns([5, 1])
                                with col_regen:
                                    if st.button(f"Gerar link para {au_name}", key=f"tm_regen_{au_id}", type="primary"):
                                        new_token = str(uuid.uuid4())[:12]
                                        link_id = str(uuid.uuid4())
                                        conn_t = get_db()
                                        c_t = conn_t.cursor()
                                        c_t.execute("UPDATE authorized_users SET invite_token = ? WHERE id = ?", (new_token, au_id))
                                        c_t.execute("INSERT INTO invite_links (id, token, invite_type, created_by) VALUES (?, ?, ?, ?)", (link_id, new_token, au_type, user_id))
                                        conn_t.commit()
                                        conn_t.close()
                                        st.rerun()
                                with col_del:
                                    if st.button("🗑", key=f"tm_remove_auth_{au_id}", help="Remover cadastro"):
                                        delete_authorized_user(au_id)
                                        st.rerun()
                        else:
                            st.caption(f"✅ Assessment concluído em {au_date}")
                        
                        st.write("")
                else:
                    st.info("Nenhum e-mail cadastrado ainda.")
                st.markdown("</div>", unsafe_allow_html=True)
        
        if is_admin:
            with tab_admin_monitoramento:
                st.markdown("<div class='gestao-card'><h3>Monitoramento em Tempo Real</h3>", unsafe_allow_html=True)
                st.markdown("<p style='color: #666; font-size: 0.9rem;'>Acompanhe os resultados de todos os assessments realizados na plataforma.</p>", unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)
                
                leaders_data = get_all_leaders_results()
                employees_data = get_all_employees_results()
                monitoring_data = get_admin_monitoring_data(user_id)
                
                st.markdown("""
                    <div class='gestao-card'>
                        <h3 style='color: #18738c; border-bottom: 2px solid #d19f09; padding-bottom: 0.5rem; margin-bottom: 1rem;'>Resultados do LPTest Líder</h3>
                        <p style='color: #666; font-size: 0.85rem; margin-bottom: 1rem;'>Gestores que completaram o assessment de liderança.</p>
                    </div>
                """, unsafe_allow_html=True)
                
                if leaders_data:
                    ldr_header = st.columns([2, 2, 2, 2, 2])
                    with ldr_header[0]: st.markdown("<strong style='color: #18738c;'>Nome</strong>", unsafe_allow_html=True)
                    with ldr_header[1]: st.markdown("<strong style='color: #18738c;'>E-mail</strong>", unsafe_allow_html=True)
                    with ldr_header[2]: st.markdown("<strong style='color: #18738c;'>Perfil Dominante</strong>", unsafe_allow_html=True)
                    with ldr_header[3]: st.markdown("<strong style='color: #18738c;'>Perfil Secundário</strong>", unsafe_allow_html=True)
                    with ldr_header[4]: st.markdown("<strong style='color: #18738c;'>Ações</strong>", unsafe_allow_html=True)
                    st.markdown("<hr style='margin: 0.5rem 0; border-color: #eee;'>", unsafe_allow_html=True)
                    
                    csv_leaders = "Nome,E-mail,Perfil Dominante,Perfil Secundário,Papel Bion,Data\n"
                    for ldr in leaders_data:
                        l_mgr_id = ldr[0]; l_user_id = ldr[1]; l_name = ldr[2] or "-"; l_email = ldr[3] or "-"
                        l_dominant = ldr[4] or "-"; l_secondary = ldr[5] or "-"
                        l_details = json.loads(ldr[6]) if ldr[6] else {}
                        l_bion = l_details.get("bion_role", "-") if isinstance(l_details, dict) else "-"
                        l_date = ldr[7][:16] if ldr[7] else "-"
                        csv_leaders += f'"{l_name}","{l_email}","{l_dominant}","{l_secondary}","{l_bion}","{l_date}"\n'
                        
                        ldr_cols = st.columns([2, 2, 2, 2, 2])
                        with ldr_cols[0]: st.write(l_name)
                        with ldr_cols[1]: st.write(l_email)
                        with ldr_cols[2]: st.markdown(f"<span class='status-badge status-concluido'>{l_dominant}</span>", unsafe_allow_html=True)
                        with ldr_cols[3]: st.write(l_secondary)
                        with ldr_cols[4]:
                            pdf_data_ldr = generate_laudo_pdf("", l_name, l_dominant or "O Idealista Exigente", l_secondary or "O Contenedor Empático", l_bion, respondent_type="gestor")
                            st.download_button("Baixar Laudo (PDF)", data=pdf_data_ldr, file_name=f"laudo_lider_{l_name.replace(' ','_').lower()}.pdf", mime="application/pdf", key=f"tm_dl_laudo_ldr_{l_mgr_id}", use_container_width=True, type="primary")
                    
                    st.write("---")
                    st.download_button("Exportar Líderes (CSV)", data=csv_leaders, file_name="relatorio_lideres_lps.csv", mime="text/csv", key="tm_export_leaders_csv", use_container_width=True)
                else:
                    st.info("Nenhum gestor completou o LPTest Líder ainda.")
                
                st.write("")
                st.markdown("""
                    <div class='gestao-card'>
                        <h3 style='color: #18738c; border-bottom: 2px solid #d19f09; padding-bottom: 0.5rem; margin-bottom: 1rem;'>Resultados do LPTest Equipe</h3>
                        <p style='color: #666; font-size: 0.85rem; margin-bottom: 1rem;'>Colaboradores que completaram o assessment de equipe.</p>
                    </div>
                """, unsafe_allow_html=True)
                
                if employees_data:
                    emp_header = st.columns([2, 2, 2, 1, 2])
                    with emp_header[0]: st.markdown("<strong style='color: #18738c;'>Nome</strong>", unsafe_allow_html=True)
                    with emp_header[1]: st.markdown("<strong style='color: #18738c;'>E-mail</strong>", unsafe_allow_html=True)
                    with emp_header[2]: st.markdown("<strong style='color: #18738c;'>Perfil</strong>", unsafe_allow_html=True)
                    with emp_header[3]: st.markdown("<strong style='color: #18738c;'>Gestor</strong>", unsafe_allow_html=True)
                    with emp_header[4]: st.markdown("<strong style='color: #18738c;'>Ações</strong>", unsafe_allow_html=True)
                    st.markdown("<hr style='margin: 0.5rem 0; border-color: #eee;'>", unsafe_allow_html=True)
                    
                    csv_employees = "Nome,E-mail,Perfil Dominante,Perfil Secundário,Papel Bion,Gestor,Data\n"
                    for emp_row in employees_data:
                        e_id = emp_row[0]; e_name = emp_row[1] or "-"; e_email = emp_row[2] or "-"
                        e_dominant = emp_row[3] or "-"; e_secondary = emp_row[4] or "-"; e_bion = emp_row[5] or "-"
                        e_manager_name = emp_row[9] or "-"; e_date = emp_row[10][:16] if emp_row[10] else "-"
                        csv_employees += f'"{e_name}","{e_email}","{e_dominant}","{e_secondary}","{e_bion}","{e_manager_name}","{e_date}"\n'
                        
                        emp_cols = st.columns([2, 2, 2, 1, 2])
                        with emp_cols[0]: st.write(e_name)
                        with emp_cols[1]: st.write(e_email)
                        with emp_cols[2]: st.markdown(f"<span class='status-badge status-concluido'>{e_dominant}</span>", unsafe_allow_html=True)
                        with emp_cols[3]: st.write(e_manager_name)
                        with emp_cols[4]:
                            pdf_data_emp = generate_laudo_pdf("", e_name, e_dominant or "O Idealista Exigente", e_secondary or "O Contenedor Empático", e_bion, respondent_type="funcionario")
                            st.download_button("Baixar Laudo (PDF)", data=pdf_data_emp, file_name=f"laudo_equipe_{e_name.replace(' ','_').lower()}.pdf", mime="application/pdf", key=f"tm_dl_laudo_emp_{e_id}", use_container_width=True, type="primary")
                    
                    st.write("---")
                    st.download_button("Exportar Equipes (CSV)", data=csv_employees, file_name="relatorio_equipes_lps.csv", mime="text/csv", key="tm_export_employees_csv", use_container_width=True)
                else:
                    st.info("Nenhum colaborador completou o LPTest Equipe ainda.")
                
                st.markdown("<div class='gestao-card'><h3>Resumo Geral</h3>", unsafe_allow_html=True)
                total_leaders = len(leaders_data) if leaders_data else 0
                total_employees = len(employees_data) if employees_data else 0
                total_auth = len(monitoring_data) if monitoring_data else 0
                total_completed = len([r for r in monitoring_data if r[3] == "concluido"]) if monitoring_data else 0
                sum_cols = st.columns(4)
                with sum_cols[0]: st.metric("Líderes Avaliados", total_leaders)
                with sum_cols[1]: st.metric("Colaboradores Avaliados", total_employees)
                with sum_cols[2]: st.metric("Total Convidados", total_auth)
                with sum_cols[3]: st.metric("Convites Concluídos", total_completed)
                st.markdown("</div>", unsafe_allow_html=True)

elif page == "InviteWelcome":
    invite_ref = st.session_state.get('invite_ref', '')
    invite_tipo = st.session_state.get('invite_tipo', 'equipe')
    
    invite_data = get_invite_by_token(invite_ref) if invite_ref else None
    
    st.markdown("""
        <style>
        .welcome-container {
            max-width: 500px;
            margin: 2rem auto;
            text-align: center;
        }
        .welcome-card {
            background: white;
            border-radius: 16px;
            padding: 2.5rem;
            box-shadow: 0 4px 20px rgba(0,0,0,0.1);
            border-top: 4px solid #18738c;
        }
        .welcome-title {
            color: #18738c;
            font-size: 1.5rem;
            font-weight: bold;
            margin: 1rem 0;
        }
        .welcome-subtitle {
            color: #666;
            font-size: 1rem;
            margin-bottom: 1.5rem;
            line-height: 1.5;
        }
        .welcome-badge {
            display: inline-block;
            background: linear-gradient(135deg, #18738c, #1a8ba6);
            color: white;
            padding: 6px 16px;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: bold;
            margin-bottom: 1rem;
        }
        </style>
    """, unsafe_allow_html=True)
    
    st.markdown("<div class='welcome-container'>", unsafe_allow_html=True)
    
    if os.path.exists(LOGO_PATH):
        col_logo = st.columns([1, 2, 1])
        with col_logo[1]:
            st.image(LOGO_PATH, use_container_width=True)
    
    tipo_label = "Assessment de Equipe" if invite_tipo == "equipe" else "Assessment de Lider"
    
    if not invite_data:
        st.markdown(f"""
            <div class='welcome-card'>
                <div class='welcome-title'>Link Invalido</div>
                <p class='welcome-subtitle'>
                    Este link de convite nao e valido ou ja foi utilizado. 
                    Solicite um novo link ao seu gestor.
                </p>
            </div>
        """, unsafe_allow_html=True)
        if st.button("Ir para a pagina inicial", key="invite_go_home", use_container_width=True):
            st.session_state.invite_ref = None
            st.session_state.invite_tipo = None
            st.session_state.page = "Home"
            st.rerun()
    elif invite_data[5] == 1:
        st.markdown(f"""
            <div class='welcome-card'>
                <div class='welcome-title'>Convite já Utilizado</div>
                <p class='welcome-subtitle'>
                    Este link de convite já foi utilizado por <strong>{invite_data[4] or 'um participante'}</strong>.
                    Solicite um novo link ao seu gestor se necessário.
                </p>
            </div>
        """, unsafe_allow_html=True)
        if st.button("Ir para a página inicial", key="invite_used_home", use_container_width=True):
            st.session_state.invite_ref = None
            st.session_state.invite_tipo = None
            st.session_state.page = "Home"
            st.rerun()
    else:
        invite_creator_user_id = invite_data[3] if invite_data else None

        st.markdown(f"""
            <div class='welcome-card'>
                <span class='welcome-badge'>{tipo_label}</span>
                <div class='welcome-title'>Bem-vindo(a) à Plataforma LPS</div>
                <p class='welcome-subtitle'>
                    Você foi convidado(a) para participar do {tipo_label} da Liderança Psicanalítica.
                    Preencha seus dados abaixo para começar.
                </p>
            </div>
        """, unsafe_allow_html=True)

        st.write("")

        with st.form("invite_start_form", clear_on_submit=False):
            participant_name = st.text_input("Seu nome completo", placeholder="Ex: Maria Silva", key="invite_name_input")
            participant_email = st.text_input("Seu e-mail", placeholder="maria@empresa.com", key="invite_email_input")
            submit_start = st.form_submit_button("Iniciar Assessment", use_container_width=True, type="primary")

            if submit_start:
                if not participant_name or not participant_email:
                    st.error("Preencha seu nome e e-mail para continuar.")
                else:
                    manager_id_for_invite = None
                    if invite_creator_user_id:
                        conn = get_db()
                        c = conn.cursor()
                        c.execute("SELECT id FROM managers WHERE user_id = ?", (invite_creator_user_id,))
                        mgr = c.fetchone()
                        conn.close()
                        if mgr:
                            manager_id_for_invite = mgr[0]

                    if manager_id_for_invite:
                        existing_slots = get_manager_employees(manager_id_for_invite)
                        next_slot = len(existing_slots) + 1
                        new_token = generate_employee_link(manager_id_for_invite, next_slot)
                        mark_invite_used(invite_ref, participant_email.lower().strip())
                        auth_check = check_email_authorized(participant_email)
                        if auth_check:
                            update_authorized_user_status(participant_email, "em_andamento")
                        st.session_state.employee_token = new_token
                        st.session_state.invite_verified_name = participant_name.strip()
                        st.session_state.invite_verified_email = participant_email.lower().strip()
                        st.session_state.page = "EmployeeAssessment"
                        st.rerun()
                    else:
                        st.error("Não foi possível identificar o gestor associado a este convite. Entre em contato: contato@lpshub.com.br")
    
    st.markdown("</div>", unsafe_allow_html=True)

elif page == "EmployeeAssessment":
    token = st.session_state.employee_token
    employee = get_employee_by_token(token)
    
    if not employee:
        st.error("Link invalido ou expirado.")
        st.markdown("""
            <div style='text-align: center; padding: 2rem;'>
                <p>Se você recebeu este link do seu gestor, entre em contato para solicitar um novo link.</p>
            </div>
        """, unsafe_allow_html=True)
    elif employee[10] == 1:  # already completed - Results page with radar + PDF
        st.markdown("""
            <div style='text-align: center; padding: 2rem;'>
                <h1 style='color: #18738c;'>Obrigado pela participacao!</h1>
            </div>
        """, unsafe_allow_html=True)
        st.markdown(f"""
            <div class="result-card" style="max-width: 600px; margin: 0 auto;">
                <div class="profile-title">Seu Perfil foi Registrado</div>
                <p style="text-align: center; font-size: 1.4rem;"><strong>{employee[6]} + {employee[7]}</strong></p>
                <div style="background-color: #d19f09; padding: 10px; border-radius: 20px; text-align: center; margin: 15px auto; max-width: 200px;">
                    <strong style="color: #18738c;">{employee[9]}</strong>
                </div>
            </div>
        """, unsafe_allow_html=True)

        emp_id = employee[0]
        emp_block_sums = get_assessment_block_sums(emp_id, "employee")
        if emp_block_sums:
            emp_profile_name = f"{employee[6]} + {employee[7]}"
            emp_radar = generate_radar_chart(emp_block_sums, emp_profile_name)
            if emp_radar:
                st.image(emp_radar, use_container_width=True)

        st.write("")
        pdf_col1, pdf_col2, pdf_col3 = st.columns([1, 2, 1])
        with pdf_col2:
            emp_pdf = generate_laudo_pdf(
                "", employee[4] or "Funcionario",
                employee[6] or "O Idealista Exigente",
                employee[7] or "O Contenedor Empatico",
                employee[9] or "", respondent_type="funcionario"
            )
            emp_safe_name = (employee[4] or "funcionario").replace(" ", "_").lower()[:20]
            st.download_button(
                "BAIXAR LAUDO COMPLETO (PDF)",
                data=emp_pdf,
                file_name=f"laudo_lps_{emp_safe_name}.pdf",
                mime="application/pdf",
                key="dl_emp_self_laudo",
                use_container_width=True,
                type="primary"
            )

        st.markdown("""
            <div style='text-align: center; margin-top: 2rem; padding: 1rem; background-color: #f5f5f5; border-radius: 10px;'>
                <p style='color: #666; margin: 0;'>Voce pode fechar esta pagina. Obrigado por participar do LPS!</p>
            </div>
        """, unsafe_allow_html=True)
    else:
        # Assessment Form - clean page for employees
        st.markdown("""
            <div style='text-align: center; margin-bottom: 1rem;'>
                <h1 style='color: #18738c;'>LPTest - Assessment de Equipe</h1>
                <p style='color: #666;'>Responda as afirmacoes. (1 = Discordo Totalmente, 5 = Concordo Totalmente)</p>
            </div>
        """, unsafe_allow_html=True)
        
        # Check consent from database (persistent storage)
        consent_given = get_employee_consent(token)
        
        # Single unified form with consent + assessment
        # Consent is validated and saved atomically with assessment results
        with st.form("employee_assessment"):
            col1, col2 = st.columns(2)
            with col1:
                name = st.text_input("Seu Nome Completo", placeholder="Maria Silva")
            with col2:
                email = st.text_input("Seu E-mail (recebera seu resultado)", placeholder="maria@email.com")
            
            st.write("---")
            
            # Show consent terms if not yet consented
            if not consent_given:
                st.markdown("""
                    <div style="background-color: #f8f9fa; padding: 1.5rem; border-radius: 10px; border-left: 4px solid #18738c; margin-bottom: 1.5rem;">
                        <h4 style="color: #18738c; margin-top: 0;">Termo de Consentimento - LGPD</h4>
                        <p style="color: #333; font-size: 0.9rem;">Ao concluir esta avaliacao, voce aceita que:</p>
                        <ul style="color: #333; font-size: 0.85rem;">
                            <li>Seus dados serao utilizados exclusivamente para fins de desenvolvimento profissional</li>
                            <li>Os resultados serao compartilhados com seu gestor para analise de perfil de equipe</li>
                            <li>Seus dados sao protegidos por sigilo e nao serao compartilhados com terceiros</li>
                            <li>Voce pode solicitar a exclusao dos seus dados a qualquer momento</li>
                        </ul>
                    </div>
                """, unsafe_allow_html=True)
                consent = st.checkbox("Aceito que meus dados sejam processados para este mapeamento de perfil profissional", key="consent_checkbox")
            else:
                consent = True  # Already consented
            
            st.write("---")
            responses = render_assessment_form("employee", is_employee=True)
            submit = st.form_submit_button("Concluir Avaliação", use_container_width=True)
            
            if submit:
                # Validation
                if not name or not email:
                    st.error("Preencha seu nome e e-mail.")
                elif not consent:
                    st.error("Voce precisa aceitar os termos de consentimento para continuar.")
                else:
                    # Atomic operation: save consent (if new) + save results
                    if not consent_given:
                        save_employee_consent(token)
                    
                    # Calculate and save results
                    dominant, secondary, details, bion_role, block_sums, cloninger_scores = calculate_employee_profile(responses)
                    save_employee_result(token, name, email, dominant, secondary, details, bion_role)
                    
                    # Save individual responses for future AI analysis
                    employee_id = employee[0]  # Get employee ID
                    save_employee_assessment_responses(employee_id, "employee", responses)
                    
                    # Send email notifications (if SMTP configured)
                    # Get manager info for notification
                    manager_id = employee[1]
                    manager_info = get_manager_by_id(manager_id)
                    if manager_info:
                        manager_name = manager_info.get('name', 'Gestor')
                        manager_email = manager_info.get('email', '')
                        
                        # Send result to employee
                        send_employee_result_email(name, email, dominant, secondary, bion_role, manager_name)
                        
                        # Notify manager
                        if manager_email:
                            send_manager_notification_email(manager_email, manager_name, name, f"{dominant} + {secondary}", bion_role)
                    
                    st.balloons()
                    st.rerun()

elif page == "LPChat":
    if not st.session_state.authenticated:
        st.session_state.page = "Login"
        st.rerun()
    
    user_premium = is_user_premium(st.session_state.user['id'])
    
    user_id = st.session_state.user['id']
    access_status = can_access_premium_features(user_id)
    course_completed = access_status['course_completed']
    
    # Custom chat styling with brand colors
    st.markdown("""
        <style>
        .chat-header {
            background: linear-gradient(135deg, #18738c 0%, #1a4f7a 100%);
            color: white;
            padding: 2rem;
            border-radius: 15px;
            margin-bottom: 1.5rem;
            text-align: center;
            box-shadow: 0 4px 15px rgba(13, 59, 102, 0.3);
        }
        .chat-header h1 {
            color: #FFFFFF !important;
            margin: 0;
            font-size: 2rem;
            -webkit-text-fill-color: #FFFFFF !important;
        }
        .chat-header p {
            color: rgba(255,255,255,0.9);
            margin: 0.5rem 0 0 0;
        }
        .chat-container {
            background: #f8f9fa;
            border-radius: 15px;
            padding: 1.5rem;
            border: 2px solid #e9ecef;
        }
        .example-questions {
            background: linear-gradient(135deg, rgba(85, 85, 85, 0.08) 0%, rgba(85, 85, 85, 0.03) 100%);
            border: 1px solid #555555;
            border-radius: 12px;
            padding: 1.5rem;
            margin-bottom: 1.5rem;
        }
        .example-questions h4 {
            color: #18738c;
            margin: 0 0 1rem 0;
        }
        .example-q {
            background: white;
            padding: 0.75rem 1rem;
            border-radius: 8px;
            margin: 0.5rem 0;
            border-left: 3px solid #555555;
            color: #333;
            font-style: italic;
            cursor: pointer;
            transition: all 0.2s;
        }
        .example-q:hover {
            background: #f0f0f0;
            border-left-color: #555555;
            transform: translateX(5px);
        }
        .team-context-card {
            background: white;
            border-radius: 12px;
            padding: 1.5rem;
            margin-bottom: 1.5rem;
            border: 1px solid #e0e0e0;
            box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        }
        .team-context-card h4 {
            color: #18738c;
            margin: 0 0 1rem 0;
            padding-bottom: 0.5rem;
            border-bottom: 2px solid #d19f09;
        }
        .team-member {
            display: flex;
            align-items: center;
            padding: 0.75rem;
            background: #f8f9fa;
            border-radius: 8px;
            margin: 0.5rem 0;
        }
        .team-member-icon {
            width: 40px;
            height: 40px;
            background: #18738c;
            color: #d19f09;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            margin-right: 1rem;
            font-weight: bold;
        }
        .bion-badge {
            background: #d19f09;
            color: #18738c;
            padding: 0.25rem 0.75rem;
            border-radius: 20px;
            font-size: 0.8rem;
            font-weight: 600;
            margin-left: auto;
        }
        </style>
    """, unsafe_allow_html=True)
    
    # Chat header with brand styling
    st.markdown("""
        <div class="chat-header">
            <h1>LPChat</h1>
            <p>Consultora de IA em Psicanálise e Neurociência aplicada a Liderança</p>
        </div>
    """, unsafe_allow_html=True)
    
    if not user_premium:
        st.markdown("""
            <div class="example-questions">
                <h4>O que o LPChat pode fazer por você:</h4>
                <div class="example-q">Análise de dinâmicas inconscientes na sua equipe</div>
                <div class="example-q">Mapeamento de conflitos e pontos de tensão grupal</div>
                <div class="example-q">Recomendações baseadas em Neurociência e Psicanálise</div>
                <div class="example-q">Estratégias de liderança personalizadas para seu perfil</div>
            </div>
        """, unsafe_allow_html=True)
        render_premium_gate()
        st.stop()
    
    if True:
        # Initialize chat history
        if 'chat_messages' not in st.session_state:
            st.session_state.chat_messages = []
        if 'last_chat_request_time' not in st.session_state:
            st.session_state.last_chat_request_time = 0
        
        # Get manager and employee data for context
        manager_data = st.session_state.manager_data
        user_name = st.session_state.user['name'] if st.session_state.user else "Gestor"
        
        # Fetch complete employees data from database
        employees_context = ""
        employees_list_display = []
        team_dynamics_analysis = ""
        
        if manager_data:
            user_id = st.session_state.user['id']
            employees = get_secure_manager_employees(user_id, manager_data['id'])
            if employees:
                employees_list = []
                bion_roles_count = {}
                
                for emp in employees:
                    if emp[10] == 1:  # completed
                        emp_name = emp[4] or f'Funcionario {emp[3]}'
                        dominant = emp[6] or "N/A"
                        secondary = emp[7] or "N/A"
                        bion_role = emp[9] or "N/A"
                        
                        # Track Bion roles for team dynamics
                        if bion_role and bion_role != "N/A":
                            bion_roles_count[bion_role] = bion_roles_count.get(bion_role, 0) + 1
                        
                        emp_info = f"""- Nome: {emp_name}
  Perfil Dominante: {dominant}
  Perfil Secundario: {secondary}
  Papel de Bion (dinamica grupal): {bion_role}
  Email: {emp[5] or 'N/A'}"""
                        employees_list.append(emp_info)
                        employees_list_display.append({
                            'name': emp_name,
                            'dominant': dominant,
                            'secondary': secondary,
                            'bion': bion_role
                        })
                
                if employees_list:
                    employees_context = "\n\n".join(employees_list)
                    
                    # Generate team dynamics summary
                    dynamics_parts = []
                    if bion_roles_count:
                        dynamics_parts.append(f"Distribuicao de Papeis de Bion na equipe: {bion_roles_count}")
                        
                        # Identify potential risks
                        if bion_roles_count.get('Bode Expiatorio', 0) > 0:
                            dynamics_parts.append("ALERTA: Ha um Bode Expiatorio na equipe - risco de projecoes negativas")
                        if bion_roles_count.get('Sabotador Silencioso', 0) > 0:
                            dynamics_parts.append("ATENCAO: Ha um Sabotador Silencioso - resistencia passiva as mudancas")
                        if bion_roles_count.get('Lider de Luta-Fuga', 0) > 1:
                            dynamics_parts.append("RISCO: Multiplos Lideres de Luta-Fuga podem gerar conflitos de poder")
                    
                    team_dynamics_analysis = "\n".join(dynamics_parts)
        
        # Manager profile context with details
        manager_profile = ""
        if manager_data and manager_data.get('dominant'):
            manager_profile = f"""Perfil do Gestor:
- Perfil Dominante: {manager_data['dominant']}
- Perfil Secundario: {manager_data['secondary']}
- Implicacoes: O gestor com perfil {manager_data['dominant']} tende a liderar com foco em {get_profile_tendency(manager_data['dominant'])}"""
        
        # Enhanced system prompt with Psychoanalysis and Neuroscience
        system_prompt = f"""Consultora LPS (Liderança Psicanalítica) de Viviane Nishiura. Especialista em psicanálise de grupos (Bion, Kernberg) e neurociência organizacional. Insights EXCLUSIVOS para o gestor.

EQUIPE:
Gestor: {user_name}
{manager_profile if manager_profile else "Gestor ainda não completou o LPTest."}

{employees_context if employees_context else "Nenhum funcionário completou o assessment."}

{team_dynamics_analysis if team_dynamics_analysis else ""}

CONCEITOS-CHAVE:
Bion: Porta-voz (expressa tensões), Bode Expiatório (absorve projeções), Dependente (evita autonomia), Luta-Fuga (reativo), Sabotador Silencioso (resistência passiva). Pressupostos Básicos: Dependência, Luta-Fuga, Acasalamento.
Neurociência: Córtisol alto=paralisia/conflito; Oxitocina alta=cooperação/confiança. Neurônios-espelho: líder regula emocionalmente a equipe. Amígdala dispara antes da consciência.
Transferência: funcionários projetam figuras parentais no líder. Contratransferência: reações emocionais do líder.
Sinek (EDSO): Círculo de Segurança forte=inovação/colaboração; fraco=proteção/competição.
Perfis: Idealista Exigente, Contenedor Empático, Buscador de Reconhecimento, Estruturador Cauteloso, Relacional Reativo, Observador Consciente, Executor Decidido.

INSTRUÇÕES:
- Ao iniciar uma conversa, apresente-se: "Olá! Sou o assistente especializado da Plataforma LPS. Como posso ajudar em sua jornada de Liderança, Psicanálise e Neurociência hoje?"
- Analise dados reais da equipe. Identifique papéis inconscientes, conflitos e sinergias entre perfis.
- Use neurociência para explicar comportamentos. Quando usar termos técnicos, explique-os brevemente entre parênteses.
- Sugira foco na TAREFA REAL para resolver regressões.
- Seja empática, profunda mas acessível. Português brasileiro."""

        # Display team context card
        if employees_list_display:
            st.markdown('<div class="team-context-card">', unsafe_allow_html=True)
            st.markdown('<h4>Sua Equipe Mapeada</h4>', unsafe_allow_html=True)
            for emp in employees_list_display:
                initial = emp['name'][0].upper() if emp['name'] else "?"
                st.markdown(f"""
                    <div class="team-member">
                        <div class="team-member-icon">{initial}</div>
                        <div>
                            <strong>{emp['name']}</strong><br>
                            <small style="color: #666;">{emp['dominant']} + {emp['secondary']}</small>
                        </div>
                        <span class="bion-badge">{emp['bion']}</span>
                    </div>
                """, unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.info("Nenhum funcionário completou o assessment ainda. Gere links de convite na área de Equipe.")
        
        # Example questions section
        st.markdown("""
            <div class="example-questions">
                <h4>Exemplos de perguntas que voce pode fazer:</h4>
                <div class="example-q">"Quais sao as principais dinamicas de transferencia no meu time?"</div>
                <div class="example-q">"Como posso melhorar a produtividade deste grupo?"</div>
                <div class="example-q">"Existe algum Bode Expiatorio na minha equipe?"</div>
                <div class="example-q">"Qual funcionario seria mais adequado para liderar o novo projeto?"</div>
                <div class="example-q">"Por que me sinto tao irritado com o Joao? (contratransferencia)"</div>
                <div class="example-q">"Como lidar com a resistencia passiva da Maria?"</div>
            </div>
        """, unsafe_allow_html=True)
        
        # Chat container
        st.markdown('<div class="chat-container">', unsafe_allow_html=True)
        
        # Display chat history with custom styling
        for message in st.session_state.chat_messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Chat input
        if prompt := st.chat_input("Descreva a situacao da sua equipe ou faca uma pergunta..."):
            now = time.time()
            time_since_last = now - st.session_state.last_chat_request_time
            if time_since_last < 5:
                st.info("Aguarde alguns segundos antes de enviar uma nova mensagem.")
            else:
                st.session_state.last_chat_request_time = now
                st.session_state.chat_messages.append({"role": "user", "content": prompt})
                with st.chat_message("user"):
                    st.markdown(prompt)
                
                with st.chat_message("assistant"):
                    with st.spinner("Analisando dinâmicas da equipe..."):
                        try:
                            client = get_openai_client()
                            if not client:
                                st.error("A chave OPENAI_API_KEY nos Secrets está inválida ou não configurada. Acesse os Secrets do Replit e insira uma chave válida da OpenAI.")
                            else:
                                recent_messages = st.session_state.chat_messages[-6:]
                                messages = [{"role": "system", "content": system_prompt}]
                                for msg in recent_messages:
                                    messages.append({"role": msg["role"], "content": msg["content"]})
                                
                                response = client.chat.completions.create(
                                    model="gpt-4o-mini",
                                    messages=messages,
                                    max_tokens=2048,
                                    temperature=0.7
                                )
                                
                                if response.choices and response.choices[0].message.content:
                                    assistant_message = response.choices[0].message.content
                                    st.markdown(assistant_message)
                                    st.session_state.chat_messages.append({"role": "assistant", "content": assistant_message})
                        
                        except Exception as e:
                            st.error(f"Erro técnico: {str(e)}")
        
        # Export and Clear buttons with styling
        if st.session_state.chat_messages:
            # Get the last AI response for export
            last_ai_response = ""
            for msg in reversed(st.session_state.chat_messages):
                if msg["role"] == "assistant":
                    last_ai_response = msg["content"]
                    break
            
            col1, col2, col3 = st.columns([1, 1, 1])
            
            with col1:
                if last_ai_response and employees_list_display:
                    # Generate PDF from last AI analysis
                    analysis_pdf = generate_ai_analysis_pdf(
                        user_name,
                        last_ai_response,
                        employees if manager_data else []
                    )
                    st.download_button(
                        label="Exportar Analise (PDF)",
                        data=analysis_pdf,
                        file_name=f"analise_ia_lps_{datetime.now().strftime('%Y%m%d')}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
            
            with col2:
                if st.button("Limpar Conversa", key="btn-clear-chat", use_container_width=True):
                    st.session_state.chat_messages = []
                    st.rerun()

elif page == "Mentoria":
    if not st.session_state.authenticated:
        st.session_state.page = "Login"
        st.rerun()
    
    user_premium = is_user_premium(st.session_state.user['id'])
    
    render_sidebar_navigation()
    render_public_header()
    
    user_id = st.session_state.user['id']
    access_status = can_access_premium_features(user_id)
    
    st.markdown("""
        <div style="text-align: center; margin-bottom: 2rem;">
            <h1 style="color: #18738c;">Mentoria Executiva LPS</h1>
            <p style="color: #666;">Sessões individuais para aprofundar sua jornada de liderança consciente</p>
        </div>
    """, unsafe_allow_html=True)
    
    if not user_premium:
        st.markdown("""
            <div style="background: #f8f9fa; border-radius: 12px; padding: 1.5rem; margin-bottom: 1.5rem; border: 1px solid #e0e0e0;">
                <h4 style="color: #18738c; margin-top: 0;">O que inclui a Mentoria Executiva:</h4>
                <ul style="color: #444; line-height: 2;">
                    <li>Sessões individuais com Viviane Nishiura</li>
                    <li>Análise aprofundada de dinâmicas de liderança</li>
                    <li>Estratégias personalizadas baseadas em Psicanálise e Neurociência</li>
                    <li>Acompanhamento de exercícios práticos e padrões comportamentais</li>
                </ul>
            </div>
        """, unsafe_allow_html=True)
        render_premium_gate()
        st.stop()
    
    if not access_status['can_access']:
        # Show paywall
        if not access_status['payment_active']:
            st.markdown("""
                <div style='background-color: #f8d7da; padding: 2rem; border-radius: 10px; border-left: 4px solid #dc3545; margin-bottom: 2rem;'>
                    <h3 style='color: #721c24; margin-top: 0;'>Conteúdo Exclusivo para Alunos</h3>
                    <p style='color: #721c24;'>
                        O agendamento de mentoria é liberado após a confirmação do pagamento.
                        Entre em contato pelo e-mail contato@lpshub.com.br para adquirir seu acesso.
                    </p>
                </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
                <div style='background-color: #fff3cd; padding: 2rem; border-radius: 10px; border-left: 4px solid #ffc107; margin-bottom: 2rem;'>
                    <h3 style='color: #856404; margin-top: 0;'>Complete o curso para liberar a Mentoria</h3>
                    <p style='color: #856404;'>
                        O acesso à mentoria é liberado após a conclusão dos módulos teóricos do curso.
                        Isso garante uma base sólida para aproveitar ao máximo sua sessão.
                    </p>
                </div>
            """, unsafe_allow_html=True)
            if st.button("Ir para o Curso", key="mentoria-goto-curso"):
                st.session_state.page = "LPS Curso"
                st.rerun()
    else:
        # Full mentoria content for paying users
        st.markdown("""
            <div class="about-card">
                <h3 style="color: #18738c;">O que esperar da Mentoria</h3>
                <ul>
                    <li><strong>Sessão individual de 1 hora</strong> com consultor sênior</li>
                    <li>Análise aprofundada do seu perfil de liderança</li>
                    <li>Discussão sobre dinâmicas de equipe e conflitos inconscientes</li>
                    <li>Estratégias práticas baseadas em psicanálise e neurociência</li>
                    <li>Acompanhamento personalizado para seu desenvolvimento</li>
                </ul>
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
            <div class="about-card">
                <h3 style="color: #18738c;">Como se preparar</h3>
                <ol>
                    <li>Revise seu perfil de lideranca no Dashboard</li>
                    <li>Analise os resultados do mapeamento de equipe</li>
                    <li>Identifique 2-3 desafios especificos que deseja abordar</li>
                    <li>Anote situacoes concretas para discutir</li>
                </ol>
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"""
            <div style="background: linear-gradient(135deg, #18738c 0%, #1a4f7a 100%); padding: 2rem; border-radius: 15px; text-align: center; margin-top: 2rem;">
                <h2 style="color: #d19f09; margin-bottom: 1rem;">Agende sua Sessao</h2>
                <p style="color: white; margin-bottom: 1.5rem;">Entre em contato via WhatsApp para agendar sua mentoria individual.</p>
                <a href="{WHATSAPP_URL}" target="_blank" style="display: inline-block; background-color: #25D366; color: #000000; padding: 1rem 3rem; border-radius: 25px; text-decoration: none; font-weight: bold; font-size: 1.1rem;">
                    Agendar Mentoria
                </a>
            </div>
        """, unsafe_allow_html=True)

elif page == "Sobre":
    if not st.session_state.authenticated:
        st.session_state.page = "Login"
        st.rerun()
    st.title("Sobre a Plataforma LPS")
    
    st.subheader("Sobre a Autora")
    st.markdown("""Viviane Nishiura é psicóloga (Mackenzie) com mais de 30 anos de trajetória conectando RH corporativo, gestão de projetos e clínica psicológica. Sua experiência permite identificar conflitos corporativos como manifestações de dinâmicas inconscientes. Atualmente lidera o LPS – Líder Psicanalítico, integrando teoria psicanalítica e visão sistêmica para ajudar líderes a sustentarem autoridade e limites com saúde mental.""")
    
    st.subheader("Sobre o LPS (Líder Psicanalítico)")
    st.markdown("""O LPS é um modelo de intervenção estruturado para o mundo corporativo. Baseia-se no princípio de que organizações ativam regressões emocionais. O programa foca em:

- **Leitura Psíquica:** Entender o funcionamento invisível das equipes.
- **Sustentação de Autoridade:** Tomar decisões estratégicas sem sobrecarga emocional.
- **Performance:** Alinhamento do comportamento à tarefa organizacional.""")
    
    st.markdown(f"""
        <div style="text-align: center; margin-top: 2rem;">
            <a href="{WHATSAPP_URL}" target="_blank" style="display: inline-block; background-color: #25D366; color: white; padding: 12px 30px; border-radius: 25px; text-decoration: none; font-weight: bold;">
                Fale Conosco no WhatsApp
            </a>
        </div>
    """, unsafe_allow_html=True)
    
    if st.button("VOLTAR PARA HOME", key="footer_back_home_contato", use_container_width=True):
        st.session_state.section = "home"
        st.session_state.page = "Home"
        st.rerun()

elif page == "GestaoLPS":
    st.session_state.page = "TeamManagement"
    st.rerun()

elif page == "GuiaSuporte":
    # Guide and Support page for managers - Authentication required
    if not st.session_state.authenticated:
        st.session_state.page = "Login"
        st.rerun()
    
    render_sidebar_navigation()
    render_public_header()
    
    st.markdown("""
        <div style="text-align: center; margin-bottom: 2rem;">
            <h1 style="color: #18738c;">Guia e Suporte</h1>
            <p style="color: #666;">Materiais de apoio para maximizar sua experiencia com a plataforma LPS</p>
        </div>
    """, unsafe_allow_html=True)
    
    # Download Manual Button
    st.markdown("""
        <div style="background: linear-gradient(135deg, #18738c 0%, #1a4f7a 100%); padding: 2rem; border-radius: 15px; margin-bottom: 2rem; text-align: center;">
            <h2 style="color: #d19f09; margin-bottom: 1rem;">Manual do Gestor LPS</h2>
            <p style="color: white; margin-bottom: 1.5rem;">Guia completo com tudo que voce precisa saber para aplicar a Liderança Psicanalítica na sua equipe.</p>
        </div>
    """, unsafe_allow_html=True)
    
    # Generate and offer PDF download
    pdf_data = generate_manager_guide_pdf()
    st.download_button(
        label="📥 Baixar Manual do Gestor (PDF)",
        data=pdf_data,
        file_name="Manual_do_Gestor_LPS.pdf",
        mime="application/pdf",
        use_container_width=True,
        type="primary"
    )
    
    st.write("---")
    
    # Section 1: Como interpretar seu Perfil
    st.markdown("""
        <div class="about-card">
            <h3 style="color: #18738c;">1. Como Interpretar seu Perfil</h3>
            <p>Seu perfil de lideranca revela seus <strong>arquetipos inconscientes dominantes</strong> - padroes de comportamento 
            que operam abaixo da consciencia e influenciam como voce lidera sua equipe.</p>
            <p>Os 7 perfis de lideranca do LPS sao:</p>
            <ul>
                <li><strong>O Idealista Exigente:</strong> Excelencia e padroes elevados de lideranca</li>
                <li><strong>O Contenedor Empatico:</strong> Estabilidade emocional e seguranca do grupo</li>
                <li><strong>O Buscador de Reconhecimento:</strong> Inspiracao e motivacao atraves de carisma</li>
                <li><strong>O Estruturador Cauteloso:</strong> Organizacao, controle e previsibilidade</li>
                <li><strong>O Relacional Reativo:</strong> Sensibilidade as dinamicas interpessoais</li>
                <li><strong>O Observador Consciente:</strong> Autoconsciencia e analise profunda</li>
                <li><strong>O Executor Decidido:</strong> Acao decisiva e foco em resultados</li>
            </ul>
            <p><em>Cada perfil tem seus pontos fortes e areas de atenção. O LPTest mostra seu perfil dominante e secundario.</em></p>
        </div>
    """, unsafe_allow_html=True)
    
    # Section 2: Mapeamento de Equipe
    st.markdown("""
        <div class="about-card">
            <h3 style="color: #18738c;">2. Mapeamento de Equipe</h3>
            <p>O LPTest mapeia os perfis inconscientes de sua equipe, permitindo entender dinamicas grupais 
            que impactam produtividade, conflitos e turnover.</p>
            <h4 style="color: #18738c;">Como usar para reduzir turnover:</h4>
            <ul>
                <li>Identifique funcionarios cujo perfil nao se adequa ao cargo atual</li>
                <li>Realoque pessoas com base em suas tendencias naturais</li>
                <li>Crie pares complementares (ex: Estruturador + Criativo)</li>
            </ul>
            <h4 style="color: #18738c;">Como usar para reduzir conflitos:</h4>
            <ul>
                <li>Identifique <strong>Bodes Expiatorios</strong> e proteja-os de projecoes negativas</li>
                <li>Reconheca <strong>Porta-Vozes</strong> como sensores do clima organizacional</li>
                <li>Transforme <strong>Lideres de Luta-Fuga</strong> em agentes de mudanca construtiva</li>
            </ul>
        </div>
    """, unsafe_allow_html=True)
    
    # Section 3: Uso Estratégico do LPChat
    st.markdown("""
        <div class="about-card">
            <h3 style="color: #18738c;">3. Uso Estratégico do LPChat</h3>
            <p>O LPChat é sua <strong>consultora de IA especializada</strong> em psicanálise e neurociência aplicada à liderança. 
            Para obter os melhores insights, faça perguntas específicas sobre sua equipe.</p>
            <h4 style="color: #18738c;">Perguntas recomendadas:</h4>
            <ul>
                <li>"Quais conflitos inconscientes podem surgir entre [Funcionário A] e [Funcionário B]?"</li>
                <li>"Qual funcionário seria ideal para liderar o projeto X, considerando os perfis mapeados?"</li>
                <li>"Como posso dar feedback construtivo para um Dependente sem gerar mais dependência?"</li>
                <li>"Quais dinâmicas de transferência podem estar afetando minha relação com a equipe?"</li>
                <li>"Como aplicar conceitos de neurociência para reduzir o estresse no time?"</li>
            </ul>
            <p><em>Dica: Quanto mais específica a pergunta, melhor será a resposta da IA.</em></p>
        </div>
    """, unsafe_allow_html=True)
    
    # Section 4: Passo a Passo da Mentoria
    st.markdown(f"""
        <div class="about-card">
            <h3 style="color: #18738c;">4. Passo a Passo da Mentoria</h3>
            <p>A Mentoria Executiva LPS é o momento de aprofundar sua jornada de liderança consciente. 
            Prepare-se adequadamente para maximizar os resultados.</p>
            <h4 style="color: #18738c;">Antes da sessão:</h4>
            <ol>
                <li>Revise seu perfil de liderança no Dashboard</li>
                <li>Analise os resultados do mapeamento de equipe</li>
                <li>Identifique 2-3 desafios específicos que deseja abordar</li>
                <li>Anote situações concretas para discutir</li>
            </ol>
            <h4 style="color: #18738c;">Durante a sessão:</h4>
            <ul>
                <li>Compartilhe abertamente seus desafios</li>
                <li>Pergunte sobre padrões inconscientes que não consegue ver</li>
                <li>Solicite exercícios práticos para aplicar no dia a dia</li>
            </ul>
            <h4 style="color: #18738c;">Como agendar:</h4>
            <p>Acesse o menu <strong>Mentoria</strong> no Dashboard ou envie e-mail para 
            <a href="mailto:contato@lpshub.com.br" style="color: #18738c;">contato@lpshub.com.br</a> para agendar sua sessão.</p>
        </div>
    """, unsafe_allow_html=True)
    
    # Contact Support
    st.markdown(f"""
        <div style="background-color: #f5f5f5; padding: 1.5rem; border-radius: 10px; text-align: center; margin-top: 2rem;">
            <h3 style="color: #18738c;">Precisa de Ajuda?</h3>
            <p>Entre em contato com nossa equipe via WhatsApp para suporte personalizado.</p>
            <a href="{WHATSAPP_URL}" target="_blank" style="display: inline-block; background-color: #25D366; color: #000000; padding: 0.75rem 2rem; border-radius: 25px; text-decoration: none; font-weight: bold;">
                Falar com Suporte
            </a>
        </div>
    """, unsafe_allow_html=True)

elif page == "Privacy":
    # Privacy and Terms page
    st.markdown('<div class="section-title">Privacidade e Termos de Uso</div>', unsafe_allow_html=True)
    
    st.markdown("""
        <div class="about-card">
            <h3 style="color: #18738c;">Politica de Privacidade - LGPD</h3>
            <p>A Plataforma LPS (Liderança Psicanalítica) esta comprometida com a proteção dos dados pessoais de todos os usuarios, 
            em conformidade com a Lei Geral de Protecao de Dados (LGPD - Lei 13.709/2018).</p>
        </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
        <div class="about-card">
            <h3 style="color: #18738c;">Dados Coletados</h3>
            <p><strong>Para Gestores:</strong></p>
            <ul style="color: #333;">
                <li>Nome completo e e-mail para autenticacao</li>
                <li>Resultados do assessment de perfil de lideranca</li>
                <li>Progresso no curso e modulos acessados</li>
            </ul>
            <p><strong>Para Funcionarios:</strong></p>
            <ul style="color: #333;">
                <li>Nome e e-mail (opcional)</li>
                <li>Respostas do assessment de perfil</li>
                <li>Resultado do mapeamento de perfil e papel de Bion</li>
            </ul>
        </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
        <div class="about-card">
            <h3 style="color: #18738c;">Finalidade do Tratamento</h3>
            <p>Os dados coletados sao utilizados <strong>exclusivamente</strong> para:</p>
            <ul style="color: #333;">
                <li>Desenvolvimento profissional e autoconhecimento dos participantes</li>
                <li>Geracao de insights sobre dinamicas de equipe para gestores</li>
                <li>Personalizacao da experiencia de aprendizagem</li>
                <li>Comunicacoes relacionadas ao programa LPS</li>
            </ul>
            <p style="color: #18738c;"><strong>Os dados NUNCA serao vendidos, compartilhados com terceiros ou utilizados para fins comerciais alem do programa.</strong></p>
        </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
        <div class="about-card">
            <h3 style="color: #18738c;">Sigilo e Confidencialidade</h3>
            <p>Todos os resultados de assessment sao tratados com <strong>sigilo absoluto</strong>:</p>
            <ul style="color: #333;">
                <li>Funcionarios: Seus resultados sao visiveis apenas para o gestor que enviou o convite</li>
                <li>Gestores: Tem acesso somente aos dados de seus proprios funcionarios</li>
                <li>Isolamento de dados: Cada gestor tem seu ambiente isolado, sem acesso a dados de outros gestores</li>
                <li>Criptografia: Senhas sao armazenadas com criptografia bcrypt de alta seguranca</li>
            </ul>
        </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
        <div class="about-card">
            <h3 style="color: #18738c;">Direitos do Usuario (LGPD)</h3>
            <p>Voce tem direito a:</p>
            <ul style="color: #333;">
                <li><strong>Acesso:</strong> Solicitar copia de todos os seus dados</li>
                <li><strong>Correcao:</strong> Retificar dados incorretos ou desatualizados</li>
                <li><strong>Exclusao:</strong> Solicitar a remocao de seus dados do sistema</li>
                <li><strong>Portabilidade:</strong> Receber seus dados em formato estruturado</li>
                <li><strong>Revogacao:</strong> Retirar seu consentimento a qualquer momento</li>
            </ul>
            <p>Para exercer esses direitos, entre em contato via WhatsApp.</p>
        </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
        <div class="about-card">
            <h3 style="color: #18738c;">Termos de Uso</h3>
            <p>Ao utilizar a Plataforma LPS, voce concorda que:</p>
            <ul style="color: #333;">
                <li>Os conteudos do curso sao protegidos por direitos autorais</li>
                <li>O acesso e pessoal e intransferivel</li>
                <li>Nao e permitido compartilhar materiais ou credenciais de acesso</li>
                <li>Os resultados de assessment sao para uso interno de desenvolvimento profissional</li>
                <li>A plataforma reserva-se o direito de suspender acessos em caso de violacao</li>
            </ul>
        </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
        <div class="about-card">
            <h3 style="color: #18738c;">Contato</h3>
            <p>Para duvidas sobre privacidade e proteção de dados:</p>
            <p><strong>Responsavel:</strong> Viviane Nishiura</p>
            <p><strong>E-mail:</strong> contato@liderancapsicanalítica.com.br</p>
        </div>
    """, unsafe_allow_html=True)
    
    st.markdown(f"""
        <div style="text-align: center; margin-top: 2rem;">
            <p style="color: #666; font-size: 0.9rem;">Ultima atualizacao: Janeiro/2026</p>
        </div>
    """, unsafe_allow_html=True)
    
    if st.button("Voltar para Home", key="btn-privacy-home", use_container_width=True):
        st.session_state.page = "Home"
        st.rerun()

# Footer for all pages (except employee assessment)
if not is_employee_access:
    st.markdown("""
        <div style="margin-top: 3rem; padding: 1.5rem; background-color: #f5f5f5; border-radius: 8px; text-align: center;">
            <p style="margin: 0; color: #666; font-size: 0.9rem;">
                Liderança Psicanalítica - Viviane Nishiura & Equipe
            </p>
        </div>
    """, unsafe_allow_html=True)
    
    if st.button("VOLTAR PARA HOME", key="footer_back_home_global", use_container_width=True):
        st.session_state.section = "home"
        st.session_state.page = "Home"
        st.rerun()
